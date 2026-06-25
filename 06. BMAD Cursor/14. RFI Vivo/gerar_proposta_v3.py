from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

BLUE = RGBColor(0x00, 0x47, 0x7B)
DARK_BLUE = RGBColor(0x00, 0x2B, 0x5C)
ACCENT = RGBColor(0x00, 0x78, 0xD4)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x33, 0x33, 0x33)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        set_cell_shading(cell, "00477B")
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            p = cell.paragraphs[0]
            run = p.runs[0] if p.runs else p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            if c_idx > 0 and any(ch.isdigit() for ch in str(val)):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F7FC")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLUE if level <= 2 else DARK_BLUE
        run.font.name = 'Calibri'
    return h


def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    return p


def para(text, bold=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def note(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run.italic = True


# ================================================================
# CAPA
# ================================================================
for _ in range(6):
    doc.add_paragraph()

para("PROPOSTA TÉCNICA", bold=True, size=28, color=BLUE,
     align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para("COE Data & AI Tools", bold=True, size=18, color=DARK_BLUE,
     align=WD_ALIGN_PARAGRAPH.CENTER)
para("Evolução do Modelo Operacional, Governança e\nTecnologia de Modelagem de Dados",
     size=14, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(4):
    doc.add_paragraph()

para("Pilar 3: Arquitetura, DevOps & Industrialização",
     bold=True, size=12, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(3):
    doc.add_paragraph()

para("Keyrus Brasil", bold=True, size=12, color=BLACK,
     align=WD_ALIGN_PARAGRAPH.CENTER)
para("Junho/2026 | Versão 3.0", size=10, color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ================================================================
# SUMÁRIO
# ================================================================
heading("Sumário", 1)
for item in [
    "1. Contexto e Escopo",
    "2. Perfis Profissionais",
    "3. Estimativa de Horas por Objetivo",
    "   3.1 Assessment do Ambiente Atual e Definição de Arquitetura Alvo",
    "   3.2 Definição do Framework de DataOps e Modelo de Ciclo de Vida",
    "   3.3 Arquitetura de Ingestão & Integração de Dados",
    "   3.4 Implementação de Pipelines DevOps (CI/CD & IaC)",
    "   3.5 Ativação e Transição para Operação",
    "4. Resumo Consolidado de Horas",
    "5. Cronograma Macro",
    "6. Alocação de Equipe",
    "7. Premissas e Riscos",
]:
    p = doc.add_paragraph(item)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

doc.add_page_break()

# ================================================================
# 1. CONTEXTO E ESCOPO
# ================================================================
heading("1. Contexto e Escopo", 1)

para("Esta proposta técnica apresenta o planejamento detalhado para a execução do "
     "Pilar 3 (Arquitetura, DevOps & Industrialização) do programa COE Data & AI Tools "
     "da Telefônica/Vivo, visando a evolução do modelo operacional, governança e "
     "tecnologia de modelagem de dados.")

para("O escopo abrange 5 objetivos focados na definição arquitetural, padronização de "
     "processos de engenharia de dados e implementação de pipelines automatizados, "
     "culminando na ativação e transição para operação contínua. As estimativas "
     "consideram horas de trabalho realizadas integralmente por profissionais humanos "
     "especializados, sem utilização de ferramentas de IA generativa.")

heading("Escopo Abrangido", 2)

add_styled_table(
    ["Pilar", "Descrição", "Objetivos"],
    [["Pilar 3", "Arquitetura, DevOps & Industrialização", "5"]],
    col_widths=[3, 10, 3])

para("")
p = doc.add_paragraph()
run = p.add_run("Nota: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run(
    "Os Pilares 1 (Governança), 2 (Modelagem de Dados) e 4 (Operação & Sustentação) "
    "não estão no escopo desta proposta. Pressupõe-se que serão tratados separadamente "
    "conforme necessidade da Telefônica/Vivo.")
run.font.size = Pt(10)
run.font.color.rgb = GRAY

doc.add_page_break()

# ================================================================
# 2. PERFIS PROFISSIONAIS
# ================================================================
heading("2. Perfis Profissionais", 1)

para("A equipe proposta é composta por 5 perfis especializados, dimensionados para "
     "o escopo de arquitetura, engenharia de dados e DevOps em ambiente corporativo "
     "de grande porte no setor de telecomunicações.")

add_styled_table(
    ["Sigla", "Perfil", "Responsabilidades Principais"],
    [
        ["GP", "Gerente de Projeto",
         "Coordenação, planejamento, reporte, gestão de riscos e stakeholders"],
        ["ADS", "Arquiteto de Dados/Soluções Sênior",
         "Design de arquitetura, liderança técnica, decisões estruturais"],
        ["EDS", "Engenheiro de Dados Sênior",
         "Desenvolvimento de pipelines, engenharia de dados, templates"],
        ["EDO", "Engenheiro DevOps/DataOps",
         "CI/CD, IaC, automação, monitoramento, observabilidade"],
        ["ADP", "Analista de Dados/Processos",
         "Documentação, mapeamento de processos, análise de gaps"],
    ],
    col_widths=[2, 5, 9])

doc.add_page_break()

# ================================================================
# 3. ESTIMATIVA DE HORAS POR OBJETIVO
# ================================================================
heading("3. Estimativa de Horas por Objetivo", 1)

# --- 3.1 ---
heading("3.1 – Assessment do Ambiente Atual e Definição de Arquitetura Alvo", 2)
para("Objetivo: Avaliar o ambiente atual e definir a arquitetura alvo de dados, "
     "garantindo alinhamento com governança, ferramentas e necessidades do negócio.",
     size=10)

para("Atividades:", bold=True, size=10)
for a in [
    "Assessment do landscape atual: plataformas de dados, ingestão, integração e pipelines",
    "Identificação de gaps técnicos e funcionais",
    "Levantamento de necessidades: disponibilidade, escalabilidade, integração com ferramenta de modelagem",
    "Definição da arquitetura alvo: ingestão, processamento, armazenamento e consumo",
    "Definição de padrões arquiteturais (lakehouse, data mesh quando aplicável)",
]:
    bullet(a)

para("Entregáveis:", bold=True, size=10)
for e in ["Assessment arquitetural (AS-IS)", "Arquitetura alvo (TO-BE)",
          "Mapa de gaps e recomendações", "Diretrizes arquiteturais"]:
    bullet(e)

para("Premissas:", bold=True, size=10)
for p_text in [
    "Acesso ao ambiente e documentação existente",
    "Envolvimento da arquitetura corporativa da Vivo",
    "Direcionamento estratégico prévio disponível",
]:
    bullet(p_text)

para("Estimativa de Horas:", bold=True, size=10)
add_styled_table(
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "28",
         "Coordenação de workshops e alinhamento com stakeholders"],
        ["ADS – Arquiteto de Dados Sr.", "140",
         "Assessment profundo, desenho da arquitetura alvo, padrões"],
        ["EDS – Engenheiro de Dados Sr.", "72",
         "Mapeamento técnico de pipelines, fontes e integrações"],
        ["ADP – Analista de Processos", "40",
         "Documentação, gap analysis, consolidação de entrevistas"],
    ],
    col_widths=[5, 2, 9])
para("Total Objetivo 3.1: 280 horas | Duração: 5 semanas",
     bold=True, size=10, color=BLUE)

doc.add_page_break()

# --- 3.2 ---
heading("3.2 – Definição do Framework de DataOps e Modelo de Ciclo de Vida", 2)
para("Objetivo: Definir como a modelagem de dados será integrada ao ciclo de "
     "engenharia, garantindo automação e consistência.", size=10)

para("Atividades:", bold=True, size=10)
for a in [
    "Definição do ciclo de vida de modelos: criação, validação, versionamento e publicação",
    "Definição de padrões de DataOps: CI/CD para dados, versionamento de modelos",
    "Definição de integração entre ferramenta de modelagem e pipelines",
    "Definição de padrões de automação: geração de DDLs, sincronização de modelos",
    "Definição de governança técnica (enforcement automático)",
]:
    bullet(a)

para("Entregáveis:", bold=True, size=10)
for e in ["Framework de DataOps", "Modelo de lifecycle de dados e modelos",
          "Diretrizes de versionamento e deploy", "Arquitetura lógica de pipelines"]:
    bullet(e)

para("Premissas:", bold=True, size=10)
for p_text in [
    "Tooling compatível com automação",
    "Sincronia na execução com o Pilar 1 para aderência aos padrões de governança",
]:
    bullet(p_text)

para("Estimativa de Horas:", bold=True, size=10)
add_styled_table(
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "24",
         "Alinhamento entre pilares e gestão de dependências"],
        ["ADS – Arquiteto de Dados Sr.", "96",
         "Definição do framework e governança técnica"],
        ["EDO – Engenheiro DevOps", "72",
         "Padrões CI/CD para dados, automação de DDLs"],
        ["EDS – Engenheiro de Dados Sr.", "64",
         "Modelo de lifecycle, integração ferramenta-pipelines"],
    ],
    col_widths=[5, 2, 9])
para("Total Objetivo 3.2: 256 horas | Duração: 4 semanas",
     bold=True, size=10, color=BLUE)

doc.add_page_break()

# --- 3.3 ---
heading("3.3 – Arquitetura de Ingestão & Integração de Dados", 2)
para("Objetivo: Estruturar a camada de ingestão e integração, garantindo "
     "escalabilidade e padronização.", size=10)

para("Atividades:", bold=True, size=10)
for a in [
    "Definição da arquitetura de ingestão: batch e streaming (quando aplicável)",
    "Padronização de pipelines de ingestão",
    "Definição de integração com múltiplas fontes e sistemas críticos",
    "Definição de padrões de transformação por camada",
    "Implementação de mecanismos de padronização de ingestão",
]:
    bullet(a)

para("Entregáveis:", bold=True, size=10)
for e in ["Arquitetura de ingestão", "Padrões de integração",
          "Templates de pipelines reutilizáveis", "Guidelines de transformação"]:
    bullet(e)

para("Premissas:", bold=True, size=10)
for p_text in [
    "Acesso às fontes de dados",
    "Disponibilidade da infraestrutura necessária",
]:
    bullet(p_text)

para("Estimativa de Horas:", bold=True, size=10)
add_styled_table(
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "28",
         "Gestão de acessos às fontes e coordenação com infraestrutura"],
        ["ADS – Arquiteto de Dados Sr.", "112",
         "Design de arquitetura de ingestão e padrões arquiteturais"],
        ["EDS – Engenheiro de Dados Sr.", "124",
         "Desenvolvimento de templates, definição de transformações"],
        ["EDO – Engenheiro DevOps", "48",
         "Mecanismos de padronização e automação de ingestão"],
    ],
    col_widths=[5, 2, 9])
para("Total Objetivo 3.3: 312 horas | Duração: 5 semanas",
     bold=True, size=10, color=BLUE)

doc.add_page_break()

# --- 3.4 ---
heading("3.4 – Implementação de Pipelines DevOps (CI/CD & IaC)", 2)
para("Objetivo: Implementar pipelines automatizados para garantir deploy contínuo "
     "e consistência entre ambientes.", size=10)

para("Atividades:", bold=True, size=10)
for a in [
    "Implementação de pipelines CI/CD: versionamento, build, deploy automatizado",
    "Implementação de Infrastructure as Code (IaC)",
    "Configuração de ambientes dev / test / prod",
    "Integração com ferramenta de modelagem e repositórios (git)",
    "Automação de deploy de artefatos de dados",
]:
    bullet(a)

para("Entregáveis:", bold=True, size=10)
for e in ["Pipelines de CI/CD implementados", "Scripts IaC",
          "Ambientes configurados (dev / test / prod)",
          "Fluxos automatizados de deploy"]:
    bullet(e)

para("Premissas:", bold=True, size=10)
for p_text in [
    "Existência de plataforma de CI/CD no ambiente Vivo",
    "Acesso a ambientes dev, test e prod conforme necessário",
]:
    bullet(p_text)

para("Estimativa de Horas:", bold=True, size=10)
add_styled_table(
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "32",
         "Gestão de entregas e coordenação entre ambientes"],
        ["ADS – Arquiteto de Dados Sr.", "56",
         "Revisão arquitetural e aderência à arquitetura alvo"],
        ["EDO – Engenheiro DevOps", "176",
         "Implementação CI/CD, scripts IaC, configuração de ambientes"],
        ["EDS – Engenheiro de Dados Sr.", "104",
         "Integração com ferramenta de modelagem e automação"],
        ["ADP – Analista de Processos", "32",
         "Documentação técnica e evidências de teste"],
    ],
    col_widths=[5, 2, 9])
para("Total Objetivo 3.4: 400 horas | Duração: 6 semanas",
     bold=True, size=10, color=BLUE)

doc.add_page_break()

# --- 3.5 ---
heading("3.5 – Ativação e Transição para Operação", 2)
para("Objetivo: Verificar que a arquitetura e pipelines implementados operem de "
     "forma estável e realizar a transferência estruturada para os times de "
     "operação da Telefônica/Vivo.", size=10)

para("Atividades:", bold=True, size=10)
for a in [
    "Execução assistida dos pipelines em ambiente de produção",
    "Configuração e validação do monitoramento inicial",
    "Ajustes de performance em pipelines e infraestrutura",
    "Elaboração do backlog de evolução priorizado",
    "Entrega de documentação técnica de transição",
    "Transferência formal para times de operação",
]:
    bullet(a)

para("Entregáveis:", bold=True, size=10)
for e in ["Pipelines operacionais validados em produção",
          "Ambiente estabilizado com monitoramento ativo",
          "Backlog de evolução priorizado",
          "Documentação técnica de transição"]:
    bullet(e)

para("Premissas:", bold=True, size=10)
for p_text in [
    "Objetivos 3.1 a 3.4 concluídos e validados",
    "Equipe de operação da Vivo identificada e disponível para handover",
]:
    bullet(p_text)

para("Estimativa de Horas:", bold=True, size=10)
add_styled_table(
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "24",
         "Gestão da transição e aceite formal"],
        ["ADS – Arquiteto de Dados Sr.", "24",
         "Validação arquitetural final e revisão de backlog"],
        ["EDS – Engenheiro de Dados Sr.", "48",
         "Execução assistida e ajustes de pipelines"],
        ["EDO – Engenheiro DevOps", "56",
         "Monitoramento, ajustes de performance e estabilização"],
        ["ADP – Analista de Processos", "28",
         "Documentação de transição e backlog de evolução"],
    ],
    col_widths=[5, 2, 9])
para("Total Objetivo 3.5: 180 horas | Duração: 3 semanas",
     bold=True, size=10, color=BLUE)

doc.add_page_break()

# ================================================================
# 4. RESUMO CONSOLIDADO
# ================================================================
heading("4. Resumo Consolidado de Horas", 1)

heading("4.1 Por Objetivo", 2)

add_styled_table(
    ["#", "Objetivo", "Horas", "Duração"],
    [
        ["3.1", "Assessment do Ambiente e Arquitetura Alvo", "280", "5 sem."],
        ["3.2", "Framework de DataOps e Ciclo de Vida", "256", "4 sem."],
        ["3.3", "Arquitetura de Ingestão & Integração", "312", "5 sem."],
        ["3.4", "Pipelines DevOps (CI/CD & IaC)", "400", "6 sem."],
        ["3.5", "Ativação e Transição para Operação", "180", "3 sem."],
        ["", "SUBTOTAL BASE", "1.428", ""],
        ["", "Buffer de Contingência (10%)", "144", "—"],
        ["", "TOTAL GERAL DA PROPOSTA", "1.572", "5 meses"],
    ],
    col_widths=[1.5, 8, 2, 2.5])

para("")

heading("4.2 Por Perfil Profissional", 2)

add_styled_table(
    ["Sigla", "Perfil", "Horas", "Meses FT*"],
    [
        ["GP", "Gerente de Projeto", "136", "0,9"],
        ["ADS", "Arquiteto de Dados/Soluções Sr.", "428", "2,7"],
        ["EDS", "Engenheiro de Dados Sênior", "412", "2,6"],
        ["EDO", "Engenheiro DevOps/DataOps", "352", "2,2"],
        ["ADP", "Analista de Dados/Processos", "100", "0,6"],
        ["", "SUBTOTAL BASE", "1.428", ""],
        ["", "Buffer de Contingência (10%)", "144", ""],
        ["", "TOTAL GERAL", "1.572", ""],
    ],
    col_widths=[2, 5.5, 2, 2])

note("* FT = Full-time equivalente, considerando 160h/mês por profissional")

heading("4.3 Distribuição Percentual", 2)

add_styled_table(
    ["Categoria", "Perfis", "Horas", "% do Base"],
    [
        ["Perfis Técnicos", "ADS + EDS + EDO", "1.192", "83%"],
        ["Gestão e Documentação", "GP + ADP", "236", "17%"],
        ["Total", "", "1.428", "100%"],
    ],
    col_widths=[4, 4, 3, 3])

doc.add_page_break()

# ================================================================
# 5. CRONOGRAMA MACRO
# ================================================================
heading("5. Cronograma Macro", 1)

para("O cronograma considera um período de 5 meses, com início previsto para "
     "Agosto/2026 e conclusão em Dezembro/2026.")

heading("5.1 Distribuição Mensal", 2)

add_styled_table(
    ["Período", "Atividades Principais", "Equipe Ativa"],
    [
        ["Mês 1\nAgo/2026",
         "Kick-off do projeto\n3.1 Assessment do Ambiente e Arquitetura Alvo",
         "GP, ADS, EDS, ADP"],
        ["Mês 2\nSet/2026",
         "3.2 Framework DataOps\n3.3 Arquitetura de Ingestão (início)",
         "GP, ADS, EDS, EDO"],
        ["Mês 3\nOut/2026",
         "3.3 Ingestão (conclusão)\n3.4 Pipelines DevOps (início)",
         "GP, ADS, EDS,\nEDO, ADP"],
        ["Mês 4\nNov/2026",
         "3.4 Pipelines DevOps (conclusão)",
         "GP, ADS, EDS,\nEDO, ADP"],
        ["Mês 5\nDez/2026",
         "3.5 Ativação e Transição\nEncerramento do projeto",
         "GP, ADS, EDS,\nEDO, ADP"],
    ],
    col_widths=[3, 7, 4])

para("")

heading("5.2 Dependências entre Objetivos", 2)

add_styled_table(
    ["Dependência", "Justificativa"],
    [
        ["3.1 → 3.2, 3.3",
         "O assessment é pré-requisito para definição do framework e da arquitetura de ingestão"],
        ["3.2 → 3.4",
         "O framework de DataOps define os padrões que serão implementados nos pipelines CI/CD"],
        ["3.3 → 3.4",
         "A arquitetura de ingestão fornece os templates que serão automatizados via CI/CD"],
        ["3.4 → 3.5",
         "Os pipelines devem estar implementados para a fase de ativação e transição"],
    ],
    col_widths=[4, 12])

doc.add_page_break()

# ================================================================
# 6. ALOCAÇÃO DE EQUIPE
# ================================================================
heading("6. Alocação de Equipe", 1)

add_styled_table(
    ["Período", "GP", "ADS", "EDS", "EDO", "ADP", "Total"],
    [
        ["Mês 1 (Ago)", "1", "1", "1", "—", "1", "4"],
        ["Mês 2 (Set)", "1", "1", "1", "1", "—", "4"],
        ["Mês 3 (Out)", "1", "1", "1", "1", "1", "5"],
        ["Mês 4 (Nov)", "1", "1", "1", "1", "1", "5"],
        ["Mês 5 (Dez)", "1", "1", "1", "1", "1", "5"],
    ],
    col_widths=[3, 2, 2, 2, 2, 2, 2])

para("")
para("Equipe média ao longo do projeto: 4-5 profissionais simultâneos",
     bold=True, size=10)
para("Equipe máxima: 5 profissionais (a partir do mês 3)",
     bold=True, size=10)

doc.add_page_break()

# ================================================================
# 7. PREMISSAS E RISCOS
# ================================================================
heading("7. Premissas e Riscos", 1)

heading("7.1 Premissas Gerais", 2)

for pt in [
    "Horas calculadas para profissionais humanos sem uso de IA generativa",
    "Jornada de 8h/dia, 20 dias úteis/mês = 160h/mês por profissional",
    "Disponibilidade dos stakeholders Vivo para workshops e validações",
    "Acesso aos ambientes e documentação existente garantidos desde o início",
    "Pilares 1 (Governança) e 2 (Modelagem de Dados) sendo tratados separadamente",
    "Ferramentas de CI/CD já licenciadas e disponíveis no ambiente",
    "Infraestrutura de cloud/on-premises disponível conforme necessidades",
    "Equipe de operação da Vivo identificada para receber a transição",
    "Não inclui horas de gestão comercial, pré-venda ou revisões executivas",
    "Operação, sustentação e observabilidade não estão no escopo desta proposta",
]:
    bullet(pt)

heading("7.2 Premissas por Objetivo", 2)

add_styled_table(
    ["Objetivo", "Premissas Específicas"],
    [
        ["3.1", "Acesso ao ambiente e documentação; envolvimento da arquitetura "
                "corporativa; direcionamento estratégico prévio"],
        ["3.2", "Tooling compatível com automação; sincronia com Pilar 1 "
                "para aderência aos padrões de governança"],
        ["3.3", "Acesso às fontes de dados; disponibilidade da infraestrutura"],
        ["3.4", "Existência de plataforma de CI/CD; acesso a ambientes "
                "dev/test/prod"],
        ["3.5", "Objetivos anteriores concluídos; equipe de operação da Vivo "
                "disponível para handover"],
    ],
    col_widths=[2.5, 13.5])

para("")

heading("7.3 Principais Riscos", 2)

add_styled_table(
    ["Severidade", "Risco", "Mitigação"],
    [
        ["Alto", "Indisponibilidade de stakeholders para validações",
         "Calendário fixo de workshops definido no kick-off"],
        ["Alto", "Dependência de Pilares 1 e 2 não concluídos no prazo",
         "Checkpoints quinzenais de alinhamento entre pilares"],
        ["Médio", "Complexidade do ambiente maior que a estimada",
         "Buffer de contingência de 10% incluído na proposta"],
        ["Médio", "Mudanças de escopo durante a execução",
         "Processo formal de change request com análise de impacto"],
        ["Médio", "Incompatibilidade de ferramentas com automação",
         "PoC técnica no início de cada fase de implementação"],
        ["Baixo", "Rotatividade da equipe Keyrus",
         "Documentação contínua e knowledge base do projeto"],
    ],
    col_widths=[2.5, 6, 7.5])

# ================================================================
# CONSIDERAÇÕES FINAIS
# ================================================================
doc.add_page_break()

heading("Considerações Finais", 1)

para("Esta proposta foi elaborada com base nas necessidades identificadas no "
     "documento de RFI, focando exclusivamente no Pilar 3 (Arquitetura, DevOps & "
     "Industrialização) do programa COE Data & AI Tools.")

para("O escopo contempla desde o assessment do ambiente atual até a entrega de "
     "pipelines operacionais automatizados, passando pela definição de frameworks "
     "de DataOps, arquitetura de ingestão e implementação completa de CI/CD com "
     "Infrastructure as Code.")

para("As estimativas refletem o esforço necessário para uma entrega de qualidade "
     "em ambiente corporativo de grande porte no setor de telecomunicações, com "
     "83% das horas alocadas em perfis técnicos especializados.")

para("A Keyrus está à disposição para apresentar esta proposta em detalhes e "
     "discutir eventuais ajustes de escopo, cronograma ou equipe conforme necessário.")

for _ in range(3):
    doc.add_paragraph()

para("_______________________________________________",
     align=WD_ALIGN_PARAGRAPH.CENTER)
para("Keyrus Brasil", bold=True, size=11,
     align=WD_ALIGN_PARAGRAPH.CENTER)
para("Junho/2026", size=10, color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER)

# ================================================================
# SALVAR
# ================================================================
output = r"C:\01. Foursys\06. BMAD Cursor\14. RFI Vivo\Proposta Tecnica COE Data AI Tools - Keyrus v3.docx"
doc.save(output)
print(f"Documento gerado com sucesso: {output}")

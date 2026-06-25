from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

BLUE = RGBColor(0x00, 0x47, 0x7B)
DARK_BLUE = RGBColor(0x00, 0x2B, 0x5C)
ACCENT = RGBColor(0x00, 0x78, 0xD4)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x33, 0x33, 0x33)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        set_cell_shading(cell, "00477B")

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            p = cell.paragraphs[0]
            run = p.runs[0] if p.runs else p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            if c_idx > 0 and any(ch.isdigit() for ch in str(val)):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F7FC")

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLUE if level <= 2 else DARK_BLUE
        run.font.name = 'Calibri'
    return h


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    return p


def add_para(doc, text, bold=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


# ================================================================
# CAPA
# ================================================================
for _ in range(6):
    doc.add_paragraph()

add_para(doc, "PROPOSTA TÉCNICA", bold=True, size=28, color=BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para(doc, "COE Data & AI Tools", bold=True, size=18, color=DARK_BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc,
         "Evolução do Modelo Operacional, Governança e\nTecnologia de Modelagem de Dados",
         bold=False, size=14, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(4):
    doc.add_paragraph()

add_para(doc,
         "Pilares 3 e 4: Arquitetura, DevOps, Industrialização,\nOperação & Sustentação",
         bold=True, size=12, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(3):
    doc.add_paragraph()

add_para(doc, "Foursys Informática Ltda.", bold=True, size=12, color=BLACK,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Junho/2026 | Versão 2.0", bold=False, size=10, color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ================================================================
# SUMÁRIO
# ================================================================
add_heading_styled(doc, "Sumário", 1)
for item in [
    "1. Contexto e Escopo",
    "2. Perfis Profissionais",
    "3. Estimativa de Horas por Objetivo",
    "   3.1 Pilar 3 – Arquitetura, DevOps & Industrialização",
    "   3.2 Pilar 4 – Operação & Sustentação",
    "   3.3 Fase Unificada – Ativação, Validação e Transição",
    "4. Resumo Consolidado de Horas",
    "5. Cronograma Macro",
    "6. Alocação de Equipe",
    "7. Premissas e Riscos",
]:
    p = doc.add_paragraph(item)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

doc.add_page_break()

# ================================================================
# 1. CONTEXTO E ESCOPO
# ================================================================
add_heading_styled(doc, "1. Contexto e Escopo", 1)

add_para(doc,
         "Esta proposta técnica apresenta o planejamento detalhado para a execução dos "
         "Pilares 3 (Arquitetura, DevOps & Industrialização) e 4 (Operação & Sustentação) "
         "do programa COE Data & AI Tools da Telefônica/Vivo, visando a evolução do modelo "
         "operacional, governança e tecnologia de modelagem de dados.")

add_para(doc,
         "O escopo abrange 8 objetivos distribuídos nos dois pilares, incluindo uma fase "
         "unificada de ativação, validação e transição que consolida as entregas de ambos os "
         "pilares. As estimativas consideram horas de trabalho realizadas integralmente por "
         "profissionais humanos especializados, sem utilização de ferramentas de IA generativa, "
         "garantindo horas realistas para planejamento e orçamento.")

add_heading_styled(doc, "Pilares Abrangidos", 2)

add_styled_table(doc,
    ["Pilar", "Descrição", "Objetivos"],
    [
        ["Pilar 3", "Arquitetura, DevOps & Industrialização", "5 + fase unificada"],
        ["Pilar 4", "Operação & Sustentação", "2 + fase unificada"],
    ],
    col_widths=[3, 10, 3])

add_para(doc, "")
p = doc.add_paragraph()
run = p.add_run("Nota: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run(
    "Os Pilares 1 (Governança) e 2 (Modelagem de Dados) são referenciados como "
    "dependências mas não estão no escopo desta proposta. Pressupõe-se que serão "
    "executados em paralelo por equipes dedicadas.")
run.font.size = Pt(10)
run.font.color.rgb = GRAY

doc.add_page_break()

# ================================================================
# 2. PERFIS PROFISSIONAIS
# ================================================================
add_heading_styled(doc, "2. Perfis Profissionais", 1)

add_para(doc,
         "A equipe proposta é composta por 7 perfis especializados, dimensionados para "
         "atender a complexidade de um ambiente corporativo de grande porte no setor de "
         "telecomunicações.")

add_styled_table(doc,
    ["Sigla", "Perfil", "Responsabilidades Principais"],
    [
        ["GP", "Gerente de Projeto",
         "Coordenação, planejamento, reporte, gestão de riscos e stakeholders"],
        ["ADS", "Arquiteto de Dados/Soluções Sênior",
         "Design de arquitetura, liderança técnica, decisões estruturais"],
        ["EDS", "Engenheiro de Dados Sênior",
         "Desenvolvimento de pipelines, engenharia de dados, templates"],
        ["EDO", "Engenheiro DevOps/DataOps",
         "CI/CD, IaC, automação, monitoramento, observabilidade"],
        ["CIO", "Consultor ITSM/Operações Sênior",
         "Modelo operacional, processos ITSM, SLAs, boas práticas ITIL"],
        ["ESD", "Especialista em Segurança de Dados",
         "Framework de segurança, compliance, LGPD, controle de acesso"],
        ["ADP", "Analista de Dados/Processos",
         "Documentação, mapeamento de processos, análise de gaps"],
    ],
    col_widths=[2, 5, 9])

doc.add_page_break()

# ================================================================
# 3. ESTIMATIVA DE HORAS POR OBJETIVO
# ================================================================
add_heading_styled(doc, "3. Estimativa de Horas por Objetivo", 1)

# ---- PILAR 3 ----
add_heading_styled(doc, "3.1  Pilar 3 – Arquitetura, DevOps & Industrialização", 2)

# --- Obj 3.1 ---
add_heading_styled(doc,
    "Objetivo 3.1 – Assessment do Ambiente Atual e Definição de Arquitetura Alvo", 3)
add_para(doc,
    "Avaliar o ambiente atual e definir a arquitetura alvo de dados, garantindo "
    "alinhamento com governança, ferramentas e necessidades do negócio.", size=10)

add_para(doc, "Atividades:", bold=True, size=10)
for a in [
    "Assessment do landscape atual: plataformas de dados, ingestão, integração e pipelines",
    "Identificação de gaps técnicos e funcionais",
    "Levantamento de necessidades: disponibilidade, escalabilidade, integração com ferramenta de modelagem",
    "Definição da arquitetura alvo: ingestão, processamento, armazenamento e consumo",
    "Definição de padrões arquiteturais (lakehouse, data mesh quando aplicável)",
]:
    add_bullet(doc, a)

add_para(doc, "Entregáveis:", bold=True, size=10)
for e in ["Assessment arquitetural (AS-IS)", "Arquitetura alvo (TO-BE)",
          "Mapa de gaps e recomendações", "Diretrizes arquiteturais"]:
    add_bullet(doc, e)

add_para(doc, "Estimativa de Horas:", bold=True, size=10)
add_styled_table(doc,
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "28", "Coordenação de workshops e stakeholders"],
        ["ADS – Arquiteto de Dados Sr.", "140",
         "Assessment profundo, desenho da arquitetura alvo"],
        ["EDS – Engenheiro de Dados Sr.", "72",
         "Mapeamento técnico de pipelines e integrações"],
        ["ADP – Analista de Processos", "40",
         "Documentação, gap analysis, consolidação de entrevistas"],
    ],
    col_widths=[5, 2, 9])
add_para(doc, "Total Objetivo 3.1: 280 horas | Duração: 5 semanas",
         bold=True, size=10, color=BLUE)
add_para(doc, "")

# --- Obj 3.2 ---
add_heading_styled(doc,
    "Objetivo 3.2 – Definição do Framework de DataOps e Modelo de Ciclo de Vida", 3)
add_para(doc,
    "Definir como a modelagem de dados será integrada ao ciclo de engenharia, "
    "garantindo automação e consistência.", size=10)

add_para(doc, "Atividades:", bold=True, size=10)
for a in [
    "Definição do ciclo de vida de modelos: criação, validação, versionamento e publicação",
    "Definição de padrões de DataOps: CI/CD para dados, versionamento de modelos",
    "Definição de integração entre ferramenta de modelagem e pipelines",
    "Definição de padrões de automação: geração de DDLs, sincronização de modelos",
    "Definição de governança técnica (enforcement automático)",
]:
    add_bullet(doc, a)

add_para(doc, "Entregáveis:", bold=True, size=10)
for e in ["Framework de DataOps", "Modelo de lifecycle de dados e modelos",
          "Diretrizes de versionamento e deploy", "Arquitetura lógica de pipelines"]:
    add_bullet(doc, e)

add_para(doc, "Estimativa de Horas:", bold=True, size=10)
add_styled_table(doc,
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "24",
         "Alinhamento entre pilares e gestão de dependências"],
        ["ADS – Arquiteto de Dados Sr.", "96",
         "Definição do framework e governança técnica"],
        ["EDO – Engenheiro DevOps", "72",
         "Padrões CI/CD para dados, automação de DDLs"],
        ["EDS – Engenheiro de Dados Sr.", "64",
         "Modelo de lifecycle, integração ferramenta-pipelines"],
    ],
    col_widths=[5, 2, 9])
add_para(doc, "Total Objetivo 3.2: 256 horas | Duração: 4 semanas",
         bold=True, size=10, color=BLUE)
add_para(doc, "")

# --- Obj 3.3 ---
add_heading_styled(doc,
    "Objetivo 3.3 – Arquitetura de Ingestão & Integração de Dados", 3)
add_para(doc,
    "Estruturar a camada de ingestão e integração, garantindo "
    "escalabilidade e padronização.", size=10)

add_para(doc, "Atividades:", bold=True, size=10)
for a in [
    "Definição da arquitetura de ingestão: batch e streaming (quando aplicável)",
    "Padronização de pipelines de ingestão",
    "Definição de integração com múltiplas fontes e sistemas críticos",
    "Definição de padrões de transformação",
    "Implementação de mecanismos de padronização de ingestão",
]:
    add_bullet(doc, a)

add_para(doc, "Entregáveis:", bold=True, size=10)
for e in ["Arquitetura de ingestão", "Padrões de integração",
          "Templates de pipelines", "Guidelines de transformação"]:
    add_bullet(doc, e)

add_para(doc, "Estimativa de Horas:", bold=True, size=10)
add_styled_table(doc,
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "28",
         "Gestão de acessos às fontes e coordenação"],
        ["ADS – Arquiteto de Dados Sr.", "112",
         "Design de arquitetura de ingestão"],
        ["EDS – Engenheiro de Dados Sr.", "124",
         "Desenvolvimento de templates e transformações"],
        ["EDO – Engenheiro DevOps", "48",
         "Mecanismos de padronização e automação"],
    ],
    col_widths=[5, 2, 9])
add_para(doc, "Total Objetivo 3.3: 312 horas | Duração: 5 semanas",
         bold=True, size=10, color=BLUE)
add_para(doc, "")

# --- Obj 3.4 ---
add_heading_styled(doc,
    "Objetivo 3.4 – Implementação de Pipelines DevOps (CI/CD & IaC)", 3)
add_para(doc,
    "Implementar pipelines automatizados para garantir deploy contínuo e "
    "consistência entre ambientes.", size=10)

add_para(doc, "Atividades:", bold=True, size=10)
for a in [
    "Implementação de pipelines CI/CD: versionamento, build, deploy automatizado",
    "Implementação de Infrastructure as Code (IaC)",
    "Configuração de ambientes dev / test / prod",
    "Integração com ferramenta de modelagem e repositórios (git)",
    "Automação de deploy de artefatos de dados",
]:
    add_bullet(doc, a)

add_para(doc, "Entregáveis:", bold=True, size=10)
for e in ["Pipelines de CI/CD implementados", "Scripts IaC",
          "Ambientes configurados", "Fluxos automatizados"]:
    add_bullet(doc, e)

add_para(doc, "Estimativa de Horas:", bold=True, size=10)
add_styled_table(doc,
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "32",
         "Gestão de entregas e coordenação entre ambientes"],
        ["ADS – Arquiteto de Dados Sr.", "56",
         "Revisão arquitetural e aderência à arquitetura alvo"],
        ["EDO – Engenheiro DevOps", "176",
         "Implementação CI/CD, scripts IaC, configuração de ambientes"],
        ["EDS – Engenheiro de Dados Sr.", "104",
         "Integração com ferramenta de modelagem e automação"],
        ["ADP – Analista de Processos", "32",
         "Documentação técnica e evidências de teste"],
    ],
    col_widths=[5, 2, 9])
add_para(doc, "Total Objetivo 3.4: 400 horas | Duração: 6 semanas",
         bold=True, size=10, color=BLUE)
add_para(doc, "")

# --- Obj 3.5 ---
add_heading_styled(doc,
    "Objetivo 3.5 – Observabilidade, Segurança, Performance e Escalabilidade", 3)
add_para(doc,
    "Garantir confiabilidade, qualidade e eficiência do ambiente de dados, "
    "preparando a arquitetura para crescimento sustentável e novos casos de uso "
    "(IA, near real-time, etc.).", size=10)

add_para(doc, "Atividades:", bold=True, size=10)
for a in [
    "Implementação de monitores: qualidade de dados e consistência de modelos",
    "Definição de métricas para performance e disponibilidade",
    "Implementação de alertas e dashboards operacionais",
    "Implementação de arquitetura escalável e múltiplas camadas de armazenamento",
    "Preparação para streaming e workloads de IA",
    "Implementação de segurança por domínio, compliance e detecção de anomalias",
    "Implementação de health checks",
]:
    add_bullet(doc, a)

add_para(doc, "Entregáveis:", bold=True, size=10)
for e in ["Arquitetura escalável implementada", "Framework de segurança",
          "Diretrizes de evolução arquitetural", "Framework de observabilidade",
          "Monitores de qualidade", "Dashboards operacionais",
          "Alertas configurados"]:
    add_bullet(doc, e)

add_para(doc, "Estimativa de Horas:", bold=True, size=10)
add_styled_table(doc,
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "32",
         "Alinhamento com políticas corporativas"],
        ["ADS – Arquiteto de Dados Sr.", "88",
         "Arquitetura escalável, diretrizes de evolução, preparação para IA"],
        ["EDS – Engenheiro de Dados Sr.", "72",
         "Monitores de qualidade de dados e consistência"],
        ["EDO – Engenheiro DevOps", "96",
         "Dashboards, alertas, health checks, métricas"],
        ["ESD – Especialista em Segurança", "96",
         "Framework de segurança, compliance, LGPD, detecção de anomalias"],
        ["ADP – Analista de Processos", "32",
         "Documentação de KPIs e procedimentos"],
    ],
    col_widths=[5, 2, 9])
add_para(doc, "Total Objetivo 3.5: 416 horas | Duração: 6 semanas",
         bold=True, size=10, color=BLUE)

doc.add_page_break()

# ---- PILAR 4 ----
add_heading_styled(doc, "3.2  Pilar 4 – Operação & Sustentação", 2)

# --- Obj 4.1 ---
add_heading_styled(doc, "Objetivo 4.1 – Desenho do Modelo Operacional", 3)
add_para(doc,
    "Definir como a operação da plataforma e da ferramenta será estruturada, "
    "alinhada à governança (Pilar 1) e à arquitetura (Pilar 3).", size=10)

add_para(doc, "Atividades:", bold=True, size=10)
for a in [
    "Definição de níveis de suporte N1, N2 e N3",
    "Definição de papéis e responsabilidades: operação da ferramenta, pipelines e incidentes",
    "Definição de processos ITSM (incidentes, problemas, requisições, mudanças, onboarding – ITIL)",
    "Definição de SLAs operacionais por severidade",
]:
    add_bullet(doc, a)

add_para(doc, "Entregáveis:", bold=True, size=10)
for e in ["Operating model de Run", "Matriz RACI para a operação",
          "Processos ITSM definidos", "SLAs definidos"]:
    add_bullet(doc, e)

add_para(doc, "Estimativa de Horas:", bold=True, size=10)
add_styled_table(doc,
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "32",
         "Coordenação com governança e alinhamento organizacional"],
        ["CIO – Consultor ITSM/Operações", "128",
         "Operating model, RACI, processos ITSM completos, SLAs"],
        ["ADS – Arquiteto de Dados Sr.", "32",
         "Alinhamento técnico com arquitetura definida"],
        ["ADP – Analista de Processos", "64",
         "Mapeamento de processos e documentação"],
    ],
    col_widths=[5, 2, 9])
add_para(doc, "Total Objetivo 4.1: 256 horas | Duração: 4 semanas",
         bold=True, size=10, color=BLUE)
add_para(doc, "")

# --- Obj 4.2 ---
add_heading_styled(doc,
    "Objetivo 4.2 – Configuração do Modelo de Operação", 3)
add_para(doc,
    "Configurar e estruturar os mecanismos operacionais necessários "
    "para suportar o Run.", size=10)

add_para(doc, "Atividades:", bold=True, size=10)
for a in [
    "Configuração de processos (workflows de chamados, gestão de incidentes e acessos)",
    "Configuração da gestão da ferramenta de modelagem: perfis, permissões, licenças, onboarding",
    "Configuração técnica de suporte: monitoramento, logs, alertas, integração com ITSM",
    "Definição de políticas de backup e DR",
]:
    add_bullet(doc, a)

add_para(doc, "Entregáveis:", bold=True, size=10)
for e in ["Plataforma preparada para operação",
          "Ferramenta configurada para uso em escala",
          "Fluxos operacionais configurados", "Integração com ITSM"]:
    add_bullet(doc, e)

add_para(doc, "Estimativa de Horas:", bold=True, size=10)
add_styled_table(doc,
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "28",
         "Gestão de entregas e coordenação"],
        ["CIO – Consultor ITSM/Operações", "112",
         "Configuração de workflows e integração ITSM"],
        ["EDO – Engenheiro DevOps", "88",
         "Monitoramento, logs, alertas, backup/DR"],
        ["EDS – Engenheiro de Dados Sr.", "48",
         "Configuração técnica da ferramenta de modelagem"],
        ["ESD – Especialista em Segurança", "48",
         "Políticas de backup, DR, segurança operacional"],
        ["ADP – Analista de Processos", "44",
         "Documentação de configurações e fluxos"],
    ],
    col_widths=[5, 2, 9])
add_para(doc, "Total Objetivo 4.2: 368 horas | Duração: 5 semanas",
         bold=True, size=10, color=BLUE)

doc.add_page_break()

# ---- FASE UNIFICADA ----
add_heading_styled(doc,
    "3.3  Fase Unificada – Ativação, Validação e Transição para Operação", 2)

add_para(doc,
    "Esta fase consolida os antigos objetivos 3.6 (Ativação e Transição – Pilar 3) e "
    "4.3 (Validação Operacional e Transição – Pilar 4) em uma única fase integrada. "
    "A unificação elimina duplicidade de esforços de gestão, documentação e handover, "
    "resultando em uma transição mais coesa e eficiente.", size=10)

add_para(doc, "Atividades:", bold=True, size=10)
for a in [
    "Execução assistida dos pipelines em produção",
    "Configuração e validação do monitoramento inicial",
    "Ajustes de performance em pipelines e infraestrutura",
    "Validação de SLAs, tempos de resposta alvo e fluxos operacionais",
    "Capacitação prática da equipe responsável pela operação (processos e ITSM)",
    "Entrega de runbooks operacionais, playbooks e procedimentos de suporte",
    "Transferência formal para times de operação e sustentação",
]:
    add_bullet(doc, a)

add_para(doc, "Entregáveis:", bold=True, size=10)
for e in [
    "Pipelines operacionais validados em produção",
    "Ambiente estabilizado com monitoramento ativo",
    "Backlog de evolução priorizado",
    "Equipe de operação capacitada",
    "Runbooks e Playbooks operacionais",
    "Documentação operacional oficialmente transferida",
    "Transição para AMS concluída com aceite formal",
]:
    add_bullet(doc, e)

add_para(doc, "Estimativa de Horas:", bold=True, size=10)
add_styled_table(doc,
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "40",
         "Gestão da transição, aceite formal e coordenação de handover"],
        ["ADS – Arquiteto de Dados Sr.", "32",
         "Validação arquitetural final e revisão de backlog"],
        ["EDS – Engenheiro de Dados Sr.", "56",
         "Execução assistida e ajustes de pipelines"],
        ["EDO – Engenheiro DevOps", "80",
         "Monitoramento, ajustes de performance e estabilização"],
        ["CIO – Consultor ITSM/Operações", "72",
         "Validação de SLAs, treinamento da equipe de operação"],
        ["ADP – Analista de Processos", "60",
         "Runbooks, playbooks e documentação completa de transição"],
    ],
    col_widths=[5, 2, 9])
add_para(doc, "Total Fase Unificada: 340 horas | Duração: 5 semanas",
         bold=True, size=10, color=BLUE)

p = doc.add_paragraph()
run = p.add_run(
    "Nota: Esta fase unificada substitui os objetivos 3.6 e 4.3 originais. "
    "A economia de 180 horas resulta da eliminação de duplicidade em gestão de "
    "projeto, documentação e atividades de handover que eram comuns a ambos.")
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.italic = True

doc.add_page_break()

# ================================================================
# 4. RESUMO CONSOLIDADO
# ================================================================
add_heading_styled(doc, "4. Resumo Consolidado de Horas", 1)

add_heading_styled(doc, "4.1 Por Objetivo", 2)

add_styled_table(doc,
    ["#", "Objetivo", "Horas", "Duração"],
    [
        ["3.1", "Assessment do Ambiente e Arquitetura Alvo", "280", "5 sem."],
        ["3.2", "Framework de DataOps e Ciclo de Vida", "256", "4 sem."],
        ["3.3", "Arquitetura de Ingestão & Integração", "312", "5 sem."],
        ["3.4", "Pipelines DevOps (CI/CD & IaC)", "400", "6 sem."],
        ["3.5", "Observabilidade, Segurança, Performance", "416", "6 sem."],
        ["", "Subtotal Pilar 3 (excl. transição)", "1.664", ""],
        ["4.1", "Desenho do Modelo Operacional", "256", "4 sem."],
        ["4.2", "Configuração do Modelo de Operação", "368", "5 sem."],
        ["", "Subtotal Pilar 4 (excl. transição)", "624", ""],
        ["U", "Ativação, Validação e Transição (unificada)", "340", "5 sem."],
        ["", "SUBTOTAL BASE", "2.628", ""],
        ["", "Buffer de Contingência (10%)", "272", "—"],
        ["", "TOTAL GERAL DA PROPOSTA", "2.900", "8 meses"],
    ],
    col_widths=[1.5, 8, 2, 2.5])

add_para(doc, "")

add_heading_styled(doc, "4.2 Por Perfil Profissional", 2)

add_styled_table(doc,
    ["Sigla", "Perfil", "Horas", "Meses FT*"],
    [
        ["GP", "Gerente de Projeto", "244", "1,5"],
        ["ADS", "Arquiteto de Dados/Soluções Sr.", "556", "3,5"],
        ["EDS", "Engenheiro de Dados Sênior", "540", "3,4"],
        ["EDO", "Engenheiro DevOps/DataOps", "560", "3,5"],
        ["CIO", "Consultor ITSM/Operações Sr.", "312", "2,0"],
        ["ESD", "Especialista em Segurança", "144", "0,9"],
        ["ADP", "Analista de Dados/Processos", "272", "1,7"],
        ["", "SUBTOTAL BASE", "2.628", ""],
        ["", "Buffer de Contingência (10%)", "272", ""],
        ["", "TOTAL GERAL", "2.900", ""],
    ],
    col_widths=[2, 5.5, 2, 2])

p = doc.add_paragraph()
run = p.add_run("* FT = Full-time equivalente, considerando 160h/mês por profissional")
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.italic = True

doc.add_page_break()

# ================================================================
# 5. CRONOGRAMA MACRO
# ================================================================
add_heading_styled(doc, "5. Cronograma Macro", 1)

add_para(doc,
    "O cronograma considera um período de 8 meses, com início previsto para "
    "Agosto/2026 e conclusão em Março/2027. As fases possuem dependências e "
    "sobreposições planejadas para otimizar o prazo total.")

add_heading_styled(doc, "5.1 Distribuição Mensal", 2)

add_styled_table(doc,
    ["Período", "Atividades Principais", "Equipe Ativa"],
    [
        ["Mês 1\nAgo/2026",
         "Kick-off do projeto\n3.1 Assessment completo",
         "GP, ADS, EDS, ADP"],
        ["Mês 2\nSet/2026",
         "3.2 Framework DataOps\n3.3 Ingestão (início)",
         "GP, ADS, EDS, EDO"],
        ["Mês 3\nOut/2026",
         "3.3 Ingestão (conclusão)\n3.4 Pipelines DevOps (início)\n4.1 Modelo Operacional (início)",
         "GP, ADS, EDS, EDO,\nCIO, ADP"],
        ["Mês 4\nNov/2026",
         "3.4 Pipelines (conclusão)\n3.5 Observabilidade (início)\n4.1 Modelo Operacional (conclusão)",
         "GP, ADS, EDS, EDO,\nCIO, ADP"],
        ["Mês 5\nDez/2026",
         "3.5 Observabilidade (conclusão)\n4.2 Configuração Operação (início)",
         "GP, ADS, EDS, EDO,\nCIO, ESD, ADP"],
        ["Mês 6\nJan/2027",
         "4.2 Configuração (conclusão)\nPreparação para ativação",
         "GP, EDO, CIO,\nESD, ADP"],
        ["Mês 7\nFev/2027",
         "Fase Unificada: Ativação, Validação\ne Transição (início)",
         "GP, ADS, EDS, EDO,\nCIO, ADP"],
        ["Mês 8\nMar/2027",
         "Fase Unificada (conclusão)\nEncerramento do projeto",
         "GP, EDO, CIO, ADP"],
    ],
    col_widths=[3, 7, 4])

add_para(doc, "")

add_heading_styled(doc, "5.2 Dependências entre Objetivos", 2)

add_styled_table(doc,
    ["Dependência", "Justificativa"],
    [
        ["3.1 → 3.2, 3.3",
         "O assessment é pré-requisito para framework e arquitetura de ingestão"],
        ["3.2 → 3.4",
         "O framework de DataOps define padrões implementados nos pipelines CI/CD"],
        ["3.3 → 3.5",
         "A arquitetura de ingestão deve estar definida para observabilidade"],
        ["3.4, 3.5 → Fase Unificada",
         "Pipelines e observabilidade devem estar prontos para ativação"],
        ["Pilar 1 → 4.1",
         "O modelo de governança (Pilar 1) precede o modelo operacional"],
        ["4.1 → 4.2",
         "O desenho do modelo operacional precede sua configuração"],
        ["4.2 → Fase Unificada",
         "A configuração operacional precede a validação e transição"],
    ],
    col_widths=[4, 12])

doc.add_page_break()

# ================================================================
# 6. ALOCAÇÃO DE EQUIPE
# ================================================================
add_heading_styled(doc, "6. Alocação de Equipe", 1)

add_styled_table(doc,
    ["Período", "GP", "ADS", "EDS", "EDO", "CIO", "ESD", "ADP", "Total"],
    [
        ["Mês 1 (Ago)", "1", "1", "1", "—", "—", "—", "1", "4"],
        ["Mês 2 (Set)", "1", "1", "1", "1", "—", "—", "—", "4"],
        ["Mês 3 (Out)", "1", "1", "1", "1", "1", "—", "1", "6"],
        ["Mês 4 (Nov)", "1", "1", "1", "1", "1", "—", "1", "6"],
        ["Mês 5 (Dez)", "1", "1", "1", "1", "1", "1", "1", "7"],
        ["Mês 6 (Jan)", "1", "—", "—", "1", "1", "1", "1", "5"],
        ["Mês 7 (Fev)", "1", "1", "1", "1", "1", "—", "1", "6"],
        ["Mês 8 (Mar)", "1", "—", "—", "1", "1", "—", "1", "4"],
    ],
    col_widths=[3, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5])

add_para(doc, "")
add_para(doc, "Equipe média ao longo do projeto: 5 profissionais simultâneos",
         bold=True, size=10)
add_para(doc, "Equipe máxima (pico em Dezembro/2026): 7 profissionais",
         bold=True, size=10)

doc.add_page_break()

# ================================================================
# 7. PREMISSAS E RISCOS
# ================================================================
add_heading_styled(doc, "7. Premissas e Riscos", 1)

add_heading_styled(doc, "7.1 Premissas Gerais", 2)

for pt in [
    "Horas calculadas para profissionais humanos sem uso de IA generativa",
    "Jornada de 8h/dia, 20 dias úteis/mês = 160h/mês por profissional",
    "Disponibilidade dos stakeholders Vivo para workshops e validações",
    "Acesso aos ambientes e documentação existente garantidos desde o início",
    "Pilares 1 e 2 sendo executados em paralelo por equipes dedicadas",
    "Ferramentas de ITSM, CI/CD e monitoramento já licenciadas e disponíveis",
    "A integração com a ferramenta de ITSM pressupõe implementação nativa",
    "Infraestrutura de cloud/on-premises disponível conforme necessidades",
    "Equipe de operação da Telefônica definida e disponível para transição",
    "Não inclui horas de gestão comercial, pré-venda ou revisões executivas",
]:
    add_bullet(doc, pt)

add_heading_styled(doc, "7.2 Premissas por Objetivo", 2)

add_styled_table(doc,
    ["Objetivo", "Premissas Específicas"],
    [
        ["3.1", "Acesso ao ambiente e documentação; envolvimento da arquitetura corporativa; "
                "direcionamento estratégico prévio"],
        ["3.2", "Tooling compatível com automação; sincronia com Pilar 1"],
        ["3.3", "Acesso às fontes de dados; disponibilidade da infraestrutura"],
        ["3.4", "Existência de plataforma de CI/CD; acesso a ambientes dev/test/prod"],
        ["3.5", "Ferramentas de monitoramento disponíveis; sincronia com Pilar 1 para KPIs; "
                "aderência a políticas corporativas; capacidade de infraestrutura"],
        ["4.1", "Modelo de governança (Pilar 1) definido; arquitetura e ferramentas especificadas"],
        ["4.2", "Ferramentas disponíveis; infraestrutura pronta; integração ITSM nativa"],
        ["Fase Unificada", "Equipe de operação definida; aprovação do modelo pela Telefônica; "
                          "sincronia com conclusão dos Pilares 3 e 4"],
    ],
    col_widths=[2.5, 13.5])

add_para(doc, "")

add_heading_styled(doc, "7.3 Principais Riscos", 2)

add_styled_table(doc,
    ["Severidade", "Risco", "Mitigação"],
    [
        ["Alto", "Indisponibilidade de stakeholders para validações",
         "Calendário fixo de workshops definido no kick-off"],
        ["Alto", "Dependência de Pilares 1 e 2 não concluídos no prazo",
         "Checkpoints quinzenais de alinhamento entre pilares"],
        ["Médio", "Complexidade do ambiente maior que a estimada",
         "Buffer de contingência de 10% incluído na proposta"],
        ["Médio", "Mudanças de escopo durante a execução",
         "Processo formal de change request com análise de impacto"],
        ["Médio", "Incompatibilidade de ferramentas com automação",
         "PoC técnica no início de cada fase de implementação"],
        ["Baixo", "Rotatividade da equipe Foursys",
         "Documentação contínua e knowledge base do projeto"],
    ],
    col_widths=[2.5, 6, 7.5])

# ================================================================
# CONSIDERAÇÕES FINAIS
# ================================================================
doc.add_page_break()

add_heading_styled(doc, "Considerações Finais", 1)

add_para(doc,
    "Esta proposta foi elaborada com base nas necessidades identificadas no documento "
    "de RFI e nas boas práticas de mercado para projetos de transformação de dados em "
    "ambientes corporativos de grande porte no setor de telecomunicações.")

add_para(doc,
    "As estimativas de horas refletem o esforço necessário para uma entrega de qualidade, "
    "considerando a complexidade do ambiente Telefônica/Vivo, as múltiplas integrações "
    "envolvidas e a necessidade de transição estruturada para operação contínua.")

add_para(doc,
    "A presente versão incorpora otimizações na alocação de perfis profissionais e a "
    "unificação das fases de transição, resultando em uma proposta mais enxuta sem "
    "comprometer a qualidade e a completude das entregas.")

add_para(doc,
    "Estamos à disposição para apresentar esta proposta em detalhes e discutir "
    "eventuais ajustes de escopo, cronograma ou equipe conforme necessário.")

for _ in range(3):
    doc.add_paragraph()

add_para(doc, "_______________________________________________",
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Foursys Informática Ltda.", bold=True, size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Junho/2026", size=10, color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER)

# ================================================================
# SALVAR
# ================================================================
output_path = r"C:\01. Foursys\06. BMAD Cursor\14. RFI Vivo\Proposta Tecnica COE Data AI Tools - Foursys v2.docx"
doc.save(output_path)
print(f"Documento gerado com sucesso: {output_path}")

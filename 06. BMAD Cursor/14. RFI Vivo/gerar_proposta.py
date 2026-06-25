from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

BLUE = RGBColor(0x00, 0x47, 0x7B)
DARK_BLUE = RGBColor(0x00, 0x2B, 0x5C)
LIGHT_BLUE = RGBColor(0xE8, 0xF0, 0xFE)
ACCENT = RGBColor(0x00, 0x78, 0xD4)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x33, 0x33, 0x33)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        set_cell_shading(cell, "00477B")

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            p = cell.paragraphs[0]
            run = p.runs[0] if p.runs else p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            if c_idx > 0 and any(ch.isdigit() for ch in str(val)):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F7FC")

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLUE if level <= 2 else DARK_BLUE
        run.font.name = 'Calibri'
    return h

def add_bullet(doc, text, bold_prefix=None, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    if indent_level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * indent_level)
    return p

def add_para(doc, text, bold=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p

# ============================================================
# CAPA
# ============================================================
for _ in range(6):
    doc.add_paragraph()

add_para(doc, "PROPOSTA TÉCNICA", bold=True, size=28, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para(doc, "COE Data & AI Tools", bold=True, size=18, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Evolução do Modelo Operacional, Governança e\nTecnologia de Modelagem de Dados", 
         bold=False, size=14, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(4):
    doc.add_paragraph()

add_para(doc, "Pilares 3 e 4: Arquitetura, DevOps, Industrialização,\nOperação, Sustentação & FinOps",
         bold=True, size=12, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(3):
    doc.add_paragraph()

add_para(doc, "Foursys Informática Ltda.", bold=True, size=12, color=BLACK, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Junho/2026 | Versão 1.0", bold=False, size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ============================================================
# SUMÁRIO
# ============================================================
add_heading_styled(doc, "Sumário", 1)
sumario_items = [
    "1. Contexto e Escopo",
    "2. Perfis Profissionais",
    "3. Estimativa de Horas por Objetivo",
    "   3.1 Pilar 3 – Arquitetura, DevOps & Industrialização",
    "   3.2 Pilar 4 – Operação, Sustentação & FinOps",
    "4. Itens Adicionais Recomendados",
    "5. Resumo Consolidado de Horas",
    "6. Cronograma Macro",
    "7. Alocação de Equipe",
    "8. Premissas e Riscos"
]
for item in sumario_items:
    p = doc.add_paragraph(item)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

doc.add_page_break()

# ============================================================
# 1. CONTEXTO E ESCOPO
# ============================================================
add_heading_styled(doc, "1. Contexto e Escopo", 1)

add_para(doc, 
    "Esta proposta técnica apresenta o planejamento detalhado para a execução dos "
    "Pilares 3 (Arquitetura, DevOps & Industrialização) e 4 (Operação, Sustentação & FinOps) "
    "do programa COE Data & AI Tools da Telefônica/Vivo, visando a evolução do modelo operacional, "
    "governança e tecnologia de modelagem de dados.")

add_para(doc,
    "O escopo abrange 9 objetivos distribuídos nos dois pilares, desde o assessment do ambiente "
    "atual até a transição completa para operação sustentada. As estimativas consideram horas de "
    "trabalho realizadas integralmente por profissionais humanos especializados, sem utilização de "
    "ferramentas de IA generativa, garantindo horas realistas para planejamento e orçamento.")

add_heading_styled(doc, "Pilares Abrangidos", 2)

add_styled_table(doc,
    ["Pilar", "Descrição", "Qtd. Objetivos"],
    [
        ["Pilar 3", "Arquitetura, DevOps & Industrialização", "6"],
        ["Pilar 4", "Operação, Sustentação & FinOps", "3"],
    ],
    col_widths=[3, 10, 3]
)

add_para(doc, "")
p = doc.add_paragraph()
run = p.add_run("Nota: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("Os Pilares 1 (Governança) e 2 (Modelagem de Dados) são referenciados como "
                "dependências mas não estão no escopo desta proposta. Pressupõe-se que serão "
                "executados em paralelo por equipes dedicadas.")
run.font.size = Pt(10)
run.font.color.rgb = GRAY

doc.add_page_break()

# ============================================================
# 2. PERFIS PROFISSIONAIS
# ============================================================
add_heading_styled(doc, "2. Perfis Profissionais", 1)

add_para(doc, "A equipe proposta é composta por 7 perfis especializados, dimensionados para "
              "atender a complexidade de um ambiente corporativo de grande porte no setor de telecomunicações.")

profiles = [
    ["GP", "Gerente de Projeto", "Coordenação, planejamento, reporte, gestão de riscos e stakeholders"],
    ["ADS", "Arquiteto de Dados/Soluções Sênior", "Design de arquitetura, liderança técnica, decisões estruturais"],
    ["EDS", "Engenheiro de Dados Sênior", "Desenvolvimento de pipelines, engenharia de dados, templates"],
    ["EDO", "Engenheiro DevOps/DataOps", "CI/CD, IaC, automação, monitoramento, observabilidade"],
    ["CIO", "Consultor ITSM/Operações Sênior", "Modelo operacional, processos ITSM, SLAs, boas práticas ITIL"],
    ["ESD", "Especialista em Segurança de Dados", "Framework de segurança, compliance, LGPD, controle de acesso"],
    ["ADP", "Analista de Dados/Processos", "Documentação, mapeamento de processos, análise de gaps"],
]

add_styled_table(doc,
    ["Sigla", "Perfil", "Responsabilidades Principais"],
    profiles,
    col_widths=[2, 5, 9]
)

doc.add_page_break()

# ============================================================
# 3. ESTIMATIVA DE HORAS POR OBJETIVO
# ============================================================
add_heading_styled(doc, "3. Estimativa de Horas por Objetivo", 1)

# ---------- PILAR 3 ----------
add_heading_styled(doc, "3.1  Pilar 3 – Arquitetura, DevOps & Industrialização", 2)

# --- Objetivo 3.1 ---
add_heading_styled(doc, "Objetivo 3.1 – Assessment do Ambiente Atual e Definição de Arquitetura Alvo", 3)
add_para(doc, "Objetivo: Avaliar o ambiente atual e definir a arquitetura alvo de dados, garantindo "
              "alinhamento com governança, ferramentas e necessidades do negócio.", size=10)

add_para(doc, "Atividades:", bold=True, size=10)
activities_31 = [
    "Assessment do landscape atual: plataformas de dados, ingestão, integração e pipelines existentes",
    "Identificação de gaps técnicos e funcionais",
    "Levantamento de necessidades: disponibilidade, escalabilidade, integração com ferramenta de modelagem",
    "Definição da arquitetura alvo: ingestão, processamento, armazenamento e consumo",
    "Definição de padrões arquiteturais (lakehouse, data mesh quando aplicável)",
]
for a in activities_31:
    add_bullet(doc, a)

add_para(doc, "Entregáveis:", bold=True, size=10)
for e in ["Assessment arquitetural (AS-IS)", "Arquitetura alvo (TO-BE)", "Mapa de gaps e recomendações", "Diretrizes arquiteturais"]:
    add_bullet(doc, e)

add_para(doc, "Estimativa de Horas:", bold=True, size=10)
add_styled_table(doc,
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "32", "Coordenação de workshops e alinhamento com stakeholders"],
        ["ADS – Arquiteto de Dados Sr.", "152", "Assessment profundo, desenho da arquitetura alvo"],
        ["EDS – Engenheiro de Dados Sr.", "88", "Mapeamento técnico de pipelines e integrações"],
        ["ADP – Analista de Processos", "48", "Documentação, gap analysis, consolidação de entrevistas"],
    ],
    col_widths=[5, 2, 9]
)
add_para(doc, "Total Objetivo 3.1: 320 horas | Duração: 5 semanas", bold=True, size=10, color=BLUE)
add_para(doc, "")

# --- Objetivo 3.2 ---
add_heading_styled(doc, "Objetivo 3.2 – Definição do Framework de DataOps e Modelo de Ciclo de Vida", 3)
add_para(doc, "Objetivo: Definir como a modelagem de dados será integrada ao ciclo de engenharia, "
              "garantindo automação e consistência.", size=10)

add_para(doc, "Atividades:", bold=True, size=10)
for a in [
    "Definição do ciclo de vida de modelos: criação, validação, versionamento e publicação",
    "Definição de padrões de DataOps: CI/CD para dados, versionamento de modelos",
    "Definição de integração entre ferramenta de modelagem e pipelines",
    "Definição de padrões de automação: geração de DDLs, sincronização de modelos",
    "Definição de governança técnica (enforcement automático)",
]:
    add_bullet(doc, a)

add_para(doc, "Entregáveis:", bold=True, size=10)
for e in ["Framework de DataOps", "Modelo de lifecycle de dados e modelos", 
          "Diretrizes de versionamento e deploy", "Arquitetura lógica de pipelines"]:
    add_bullet(doc, e)

add_para(doc, "Estimativa de Horas:", bold=True, size=10)
add_styled_table(doc,
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "24", "Alinhamento entre pilares e gestão de dependências"],
        ["ADS – Arquiteto de Dados Sr.", "112", "Definição do framework e governança técnica"],
        ["EDO – Engenheiro DevOps", "88", "Padrões CI/CD para dados, automação de DDLs"],
        ["EDS – Engenheiro de Dados Sr.", "72", "Modelo de lifecycle, integração ferramenta-pipelines"],
    ],
    col_widths=[5, 2, 9]
)
add_para(doc, "Total Objetivo 3.2: 296 horas | Duração: 4 semanas", bold=True, size=10, color=BLUE)
add_para(doc, "")

# --- Objetivo 3.3 ---
add_heading_styled(doc, "Objetivo 3.3 – Arquitetura de Ingestão & Integração de Dados", 3)
add_para(doc, "Objetivo: Estruturar a camada de ingestão e integração, garantindo "
              "escalabilidade e padronização.", size=10)

add_para(doc, "Atividades:", bold=True, size=10)
for a in [
    "Definição da arquitetura de ingestão: batch e streaming (quando aplicável)",
    "Padronização de pipelines de ingestão",
    "Definição de integração com múltiplas fontes e sistemas críticos",
    "Definição de padrões de transformação",
    "Implementação de mecanismos de padronização de ingestão",
]:
    add_bullet(doc, a)

add_para(doc, "Entregáveis:", bold=True, size=10)
for e in ["Arquitetura de ingestão", "Padrões de integração", 
          "Templates de pipelines", "Guidelines de transformação"]:
    add_bullet(doc, e)

add_para(doc, "Estimativa de Horas:", bold=True, size=10)
add_styled_table(doc,
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "32", "Gestão de acessos às fontes e coordenação"],
        ["ADS – Arquiteto de Dados Sr.", "128", "Design de arquitetura de ingestão"],
        ["EDS – Engenheiro de Dados Sr.", "144", "Desenvolvimento de templates e transformações"],
        ["EDO – Engenheiro DevOps", "56", "Mecanismos de padronização e automação"],
    ],
    col_widths=[5, 2, 9]
)
add_para(doc, "Total Objetivo 3.3: 360 horas | Duração: 5 semanas", bold=True, size=10, color=BLUE)
add_para(doc, "")

# --- Objetivo 3.4 ---
add_heading_styled(doc, "Objetivo 3.4 – Implementação de Pipelines DevOps (CI/CD & IaC)", 3)
add_para(doc, "Objetivo: Implementar pipelines automatizados para garantir deploy contínuo e "
              "consistência entre ambientes.", size=10)

add_para(doc, "Atividades:", bold=True, size=10)
for a in [
    "Implementação de pipelines CI/CD: versionamento, build, deploy automatizado",
    "Implementação de Infrastructure as Code (IaC)",
    "Configuração de ambientes dev / test / prod",
    "Integração com ferramenta de modelagem e repositórios (git)",
    "Automação de deploy de artefatos de dados",
]:
    add_bullet(doc, a)

add_para(doc, "Entregáveis:", bold=True, size=10)
for e in ["Pipelines de CI/CD implementados", "Scripts IaC", 
          "Ambientes configurados", "Fluxos automatizados"]:
    add_bullet(doc, e)

add_para(doc, "Estimativa de Horas:", bold=True, size=10)
add_styled_table(doc,
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "36", "Gestão de entregas e coordenação entre ambientes"],
        ["ADS – Arquiteto de Dados Sr.", "72", "Revisão arquitetural e aderência à arquitetura alvo"],
        ["EDO – Engenheiro DevOps", "184", "Implementação CI/CD, scripts IaC, ambientes"],
        ["EDS – Engenheiro de Dados Sr.", "112", "Integração com ferramenta de modelagem"],
        ["ADP – Analista de Processos", "36", "Documentação técnica e evidências de teste"],
    ],
    col_widths=[5, 2, 9]
)
add_para(doc, "Total Objetivo 3.4: 440 horas | Duração: 6 semanas", bold=True, size=10, color=BLUE)
add_para(doc, "")

# --- Objetivo 3.5 ---
add_heading_styled(doc, "Objetivo 3.5 – Observabilidade, Segurança, Performance e Escalabilidade", 3)
add_para(doc, "Objetivo: Garantir confiabilidade, qualidade e eficiência do ambiente de dados, "
              "preparando a arquitetura para crescimento sustentável e novos casos de uso "
              "(IA, near real-time, etc.).", size=10)

add_para(doc, "Atividades:", bold=True, size=10)
for a in [
    "Implementação de monitores: qualidade de dados e consistência de modelos",
    "Definição de métricas para performance e disponibilidade",
    "Implementação de alertas e dashboards operacionais",
    "Implementação de arquitetura escalável e múltiplas camadas de armazenamento",
    "Preparação para streaming e workloads de IA",
    "Implementação de segurança por domínio, compliance e detecção de anomalias",
    "Implementação de health checks",
]:
    add_bullet(doc, a)

add_para(doc, "Entregáveis:", bold=True, size=10)
for e in ["Arquitetura escalável implementada", "Framework de segurança", 
          "Diretrizes de evolução arquitetural", "Framework de observabilidade",
          "Monitores de qualidade", "Dashboards operacionais", "Alertas configurados"]:
    add_bullet(doc, e)

add_para(doc, "Estimativa de Horas:", bold=True, size=10)
add_styled_table(doc,
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "36", "Alinhamento com políticas corporativas"],
        ["ADS – Arquiteto de Dados Sr.", "112", "Arquitetura escalável, preparação para IA"],
        ["EDS – Engenheiro de Dados Sr.", "88", "Monitores de qualidade de dados"],
        ["EDO – Engenheiro DevOps", "108", "Dashboards, alertas, health checks, métricas"],
        ["ESD – Especialista em Segurança", "96", "Framework de segurança, compliance, LGPD"],
        ["ADP – Analista de Processos", "40", "Documentação de KPIs e procedimentos"],
    ],
    col_widths=[5, 2, 9]
)
add_para(doc, "Total Objetivo 3.5: 480 horas | Duração: 6 semanas", bold=True, size=10, color=BLUE)
add_para(doc, "")

# --- Objetivo 3.6 ---
add_heading_styled(doc, "Objetivo 3.6 – Ativação e Transição para Operação", 3)
add_para(doc, "Objetivo: Verificar que a arquitetura e pipelines operem de forma estável e "
              "sejam transferidos para sustentação contínua.", size=10)

add_para(doc, "Atividades:", bold=True, size=10)
for a in [
    "Execução assistida dos pipelines em produção",
    "Configuração do monitoramento inicial",
    "Ajustes de performance",
    "Transferência para times de operação",
]:
    add_bullet(doc, a)

add_para(doc, "Entregáveis:", bold=True, size=10)
for e in ["Pipelines operacionais", "Ambiente estabilizado", 
          "Backlog de evolução", "Transição para AMS concluída"]:
    add_bullet(doc, e)

add_para(doc, "Estimativa de Horas:", bold=True, size=10)
add_styled_table(doc,
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "32", "Gestão da transição e aceite formal"],
        ["ADS – Arquiteto de Dados Sr.", "52", "Validação arquitetural e revisão de backlog"],
        ["EDS – Engenheiro de Dados Sr.", "72", "Execução assistida e ajustes de pipelines"],
        ["EDO – Engenheiro DevOps", "72", "Monitoramento e ajustes de performance"],
        ["ADP – Analista de Processos", "40", "Documentação de transição e backlog"],
    ],
    col_widths=[5, 2, 9]
)
add_para(doc, "Total Objetivo 3.6: 268 horas | Duração: 4 semanas", bold=True, size=10, color=BLUE)

doc.add_page_break()

# ---------- PILAR 4 ----------
add_heading_styled(doc, "3.2  Pilar 4 – Operação, Sustentação & FinOps", 2)

# --- Objetivo 4.1 ---
add_heading_styled(doc, "Objetivo 4.1 – Desenho do Modelo Operacional", 3)
add_para(doc, "Objetivo: Definir como a operação da plataforma e da ferramenta será estruturada, "
              "alinhada à governança (Pilar 1) e à arquitetura (Pilar 3).", size=10)

add_para(doc, "Atividades:", bold=True, size=10)
for a in [
    "Definição de níveis de suporte N1, N2 e N3",
    "Definição de papéis e responsabilidades: operação da ferramenta, pipelines e gestão de incidentes",
    "Definição de processos ITSM (incidentes, problemas, requisições, mudanças, onboarding – ITIL)",
    "Definição de SLAs operacionais por severidade",
]:
    add_bullet(doc, a)

add_para(doc, "Entregáveis:", bold=True, size=10)
for e in ["Operating model de Run", "Matriz RACI para a operação", 
          "Processos ITSM definidos", "SLAs definidos"]:
    add_bullet(doc, e)

add_para(doc, "Estimativa de Horas:", bold=True, size=10)
add_styled_table(doc,
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "32", "Coordenação com governança e alinhamento organizacional"],
        ["CIO – Consultor ITSM/Operações", "144", "Operating model, RACI, processos ITSM, SLAs"],
        ["ADS – Arquiteto de Dados Sr.", "40", "Alinhamento técnico com arquitetura"],
        ["ADP – Analista de Processos", "72", "Mapeamento de processos e documentação"],
    ],
    col_widths=[5, 2, 9]
)
add_para(doc, "Total Objetivo 4.1: 288 horas | Duração: 4 semanas", bold=True, size=10, color=BLUE)
add_para(doc, "")

# --- Objetivo 4.2 ---
add_heading_styled(doc, "Objetivo 4.2 – Configuração do Modelo de Operação", 3)
add_para(doc, "Objetivo: Configurar e estruturar os mecanismos operacionais necessários "
              "para suportar o Run.", size=10)

add_para(doc, "Atividades:", bold=True, size=10)
for a in [
    "Configuração de processos (workflows de chamados, gestão de incidentes e acessos)",
    "Configuração da gestão da ferramenta de modelagem: perfis, permissões, licenças, onboarding",
    "Configuração técnica de suporte: monitoramento, logs, alertas, integração com ITSM",
    "Definição de políticas de backup e DR",
]:
    add_bullet(doc, a)

add_para(doc, "Entregáveis:", bold=True, size=10)
for e in ["Plataforma preparada para operação", "Ferramenta configurada para uso em escala",
          "Fluxos operacionais configurados", "Integração com ITSM"]:
    add_bullet(doc, e)

add_para(doc, "Estimativa de Horas:", bold=True, size=10)
add_styled_table(doc,
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "32", "Gestão de entregas e coordenação"],
        ["CIO – Consultor ITSM/Operações", "128", "Configuração de workflows e integração ITSM"],
        ["EDO – Engenheiro DevOps", "92", "Monitoramento, logs, alertas, backup/DR"],
        ["EDS – Engenheiro de Dados Sr.", "52", "Configuração técnica da ferramenta de modelagem"],
        ["ESD – Especialista em Segurança", "56", "Políticas de backup, DR, segurança operacional"],
        ["ADP – Analista de Processos", "52", "Documentação de configurações e fluxos"],
    ],
    col_widths=[5, 2, 9]
)
add_para(doc, "Total Objetivo 4.2: 412 horas | Duração: 5 semanas", bold=True, size=10, color=BLUE)
add_para(doc, "")

# --- Objetivo 4.3 ---
add_heading_styled(doc, "Objetivo 4.3 – Validação Operacional e Transição para Operação", 3)
add_para(doc, "Objetivo: Validar o modelo de operação e transferir de forma estruturada "
              "a operação para a equipe responsável, garantindo autonomia.", size=10)

add_para(doc, "Atividades:", bold=True, size=10)
for a in [
    "Validação de SLAs, tempos de resposta alvo e fluxos operacionais",
    "Execução de ajustes, se necessário",
    "Capacitação prática da equipe responsável pela operação nos processos e ferramenta ITSM",
    "Entrega de documentação (runbooks operacionais, playbooks e procedimentos de suporte)",
]:
    add_bullet(doc, a)

add_para(doc, "Entregáveis:", bold=True, size=10)
for e in ["Ajustes realizados (se necessário)", "Equipe capacitada",
          "Runbooks e Playbooks", "Documentação operacional oficialmente transferida"]:
    add_bullet(doc, e)

add_para(doc, "Estimativa de Horas:", bold=True, size=10)
add_styled_table(doc,
    ["Perfil", "Horas", "Escopo Principal"],
    [
        ["GP – Gerente de Projeto", "32", "Gestão de transição e aceite formal"],
        ["CIO – Consultor ITSM/Operações", "92", "Validação de SLAs, treinamento, ajustes finais"],
        ["EDO – Engenheiro DevOps", "52", "Suporte técnico durante validação"],
        ["ADP – Analista de Processos", "76", "Runbooks, playbooks, documentação completa"],
    ],
    col_widths=[5, 2, 9]
)
add_para(doc, "Total Objetivo 4.3: 252 horas | Duração: 4 semanas", bold=True, size=10, color=BLUE)

doc.add_page_break()

# ============================================================
# 4. ITENS ADICIONAIS RECOMENDADOS
# ============================================================
add_heading_styled(doc, "4. Itens Adicionais Recomendados", 1)

add_para(doc, "Com base na análise do escopo e nas boas práticas para projetos de transformação "
              "de dados em ambientes corporativos de grande porte, recomendamos a inclusão dos "
              "seguintes itens adicionais:")

add_heading_styled(doc, "4.1 FinOps – Gestão Financeira de Operações de Dados", 3)
add_para(doc, "O Pilar 4 menciona FinOps em seu título, porém as atividades originais não "
              "contemplam este tema. Recomendamos incluir:", size=10)
for a in [
    "Definição de modelo de custeio e chargeback por domínio de dados",
    "Dashboards de custo por workload/pipeline",
    "Políticas de otimização de custos de cloud/plataforma",
    "Recomendações de right-sizing e reservas",
]:
    add_bullet(doc, a)

add_styled_table(doc,
    ["Perfil", "Horas"],
    [
        ["CIO – Consultor ITSM/Operações", "56"],
        ["EDO – Engenheiro DevOps", "40"],
        ["ADP – Analista de Processos", "24"],
    ],
    col_widths=[8, 3]
)
add_para(doc, "Total FinOps: 120 horas", bold=True, size=10, color=BLUE)
add_para(doc, "")

add_heading_styled(doc, "4.2 Gestão de Mudança Organizacional", 3)
add_para(doc, "Para garantir a adoção efetiva das novas práticas e tecnologias:", size=10)
for a in [
    "Plano de comunicação para stakeholders e usuários impactados",
    "Gestão de resistência à mudança",
    "Workshops de engajamento por domínio de dados",
]:
    add_bullet(doc, a)

add_styled_table(doc,
    ["Perfil", "Horas"],
    [
        ["GP – Gerente de Projeto", "40"],
        ["ADP – Analista de Processos", "40"],
    ],
    col_widths=[8, 3]
)
add_para(doc, "Total Gestão de Mudança: 80 horas", bold=True, size=10, color=BLUE)
add_para(doc, "")

add_heading_styled(doc, "4.3 Buffer de Contingência (15%)", 3)
add_para(doc, "Para um projeto em ambiente corporativo de grande porte no setor de telecomunicações, "
              "recomendamos 15% de buffer sobre o total de horas para cobrir:", size=10)
for a in [
    "Complexidade não prevista durante o assessment",
    "Indisponibilidade eventual de stakeholders",
    "Ajustes de escopo menores durante a execução",
    "Curva de aprendizado em ferramentas específicas do ambiente Vivo",
]:
    add_bullet(doc, a)

add_para(doc, "Buffer de Contingência: 500 horas (≈15%)", bold=True, size=10, color=BLUE)

doc.add_page_break()

# ============================================================
# 5. RESUMO CONSOLIDADO
# ============================================================
add_heading_styled(doc, "5. Resumo Consolidado de Horas", 1)

add_heading_styled(doc, "5.1 Por Objetivo", 2)

add_styled_table(doc,
    ["#", "Objetivo", "Horas", "Duração"],
    [
        ["3.1", "Assessment do Ambiente e Arquitetura Alvo", "320", "5 sem."],
        ["3.2", "Framework de DataOps e Ciclo de Vida", "296", "4 sem."],
        ["3.3", "Arquitetura de Ingestão & Integração", "360", "5 sem."],
        ["3.4", "Pipelines DevOps (CI/CD & IaC)", "440", "6 sem."],
        ["3.5", "Observabilidade, Segurança, Performance", "480", "6 sem."],
        ["3.6", "Ativação e Transição para Operação", "268", "4 sem."],
        ["", "Subtotal Pilar 3", "2.164", ""],
        ["4.1", "Desenho do Modelo Operacional", "288", "4 sem."],
        ["4.2", "Configuração do Modelo de Operação", "412", "5 sem."],
        ["4.3", "Validação Operacional e Transição", "252", "4 sem."],
        ["", "Subtotal Pilar 4", "952", ""],
        ["", "Total Base (Pilares 3 + 4)", "3.116", ""],
        ["", "FinOps (Adicional)", "120", "3 sem."],
        ["", "Gestão de Mudança (Adicional)", "80", "2 sem."],
        ["", "Subtotal com Adicionais", "3.316", ""],
        ["", "Buffer de Contingência (15%)", "500", "—"],
        ["", "TOTAL GERAL DA PROPOSTA", "3.816", "8 meses"],
    ],
    col_widths=[1.5, 8, 2, 2.5]
)

add_para(doc, "")

add_heading_styled(doc, "5.2 Por Perfil Profissional", 2)

add_styled_table(doc,
    ["Sigla", "Perfil", "Horas Base", "Adicionais", "Total", "Meses FT*"],
    [
        ["GP", "Gerente de Projeto", "288", "40", "328", "2,1"],
        ["ADS", "Arquiteto de Dados/Soluções Sr.", "668", "—", "668", "4,2"],
        ["EDS", "Engenheiro de Dados Sênior", "628", "—", "628", "3,9"],
        ["EDO", "Engenheiro DevOps/DataOps", "652", "40", "692", "4,3"],
        ["CIO", "Consultor ITSM/Operações Sr.", "364", "56", "420", "2,6"],
        ["ESD", "Especialista em Segurança", "152", "—", "152", "1,0"],
        ["ADP", "Analista de Dados/Processos", "364", "64", "428", "2,7"],
        ["", "Total (sem buffer)", "", "", "3.316", ""],
        ["", "Buffer de Contingência (15%)", "", "", "500", ""],
        ["", "TOTAL GERAL", "", "", "3.816", ""],
    ],
    col_widths=[1.5, 5.5, 2, 2, 2, 2]
)

p = doc.add_paragraph()
run = p.add_run("* FT = Full-time equivalente, considerando 160h/mês por profissional")
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.italic = True

doc.add_page_break()

# ============================================================
# 6. CRONOGRAMA MACRO
# ============================================================
add_heading_styled(doc, "6. Cronograma Macro", 1)

add_para(doc, "O cronograma considera um período de 8 meses, com início previsto para "
              "Agosto/2026 e conclusão em Março/2027. As fases possuem dependências e "
              "sobreposições planejadas para otimizar o prazo total.")

add_heading_styled(doc, "6.1 Distribuição Mensal", 2)

cronograma = [
    ["Mês 1\nAgo/2026", "Kick-off do projeto\n3.1 Assessment completo", "GP, ADS, EDS, ADP"],
    ["Mês 2\nSet/2026", "3.2 Framework DataOps\n3.3 Ingestão (início)", "GP, ADS, EDS, EDO"],
    ["Mês 3\nOut/2026", "3.3 Ingestão (conclusão)\n3.4 Pipelines DevOps (início)\n4.1 Modelo Operacional (início)", "GP, ADS, EDS, EDO, CIO, ADP"],
    ["Mês 4\nNov/2026", "3.4 Pipelines (conclusão)\n3.5 Observabilidade (início)\n4.1 Modelo Operacional (conclusão)", "GP, ADS, EDS, EDO, CIO, ADP"],
    ["Mês 5\nDez/2026", "3.5 Observabilidade (conclusão)\n4.2 Configuração Operação (início)\nFinOps", "GP, ADS, EDS, EDO, CIO, ESD, ADP"],
    ["Mês 6\nJan/2027", "4.2 Configuração (conclusão)\nGestão de Mudança\nPreparação para ativação", "GP, EDO, CIO, ESD, ADP"],
    ["Mês 7\nFev/2027", "3.6 Ativação e Transição\n4.3 Validação Operacional (início)", "GP, ADS, EDS, EDO, CIO, ADP"],
    ["Mês 8\nMar/2027", "4.3 Validação e Transição final\nEncerramento do projeto", "GP, EDO, CIO, ADP"],
]

add_styled_table(doc,
    ["Período", "Atividades Principais", "Equipe Ativa"],
    cronograma,
    col_widths=[3, 8, 5]
)

add_para(doc, "")

add_heading_styled(doc, "6.2 Dependências entre Objetivos", 2)

deps = [
    ["3.1 → 3.2, 3.3", "O assessment é pré-requisito para definição do framework e da arquitetura de ingestão"],
    ["3.2 → 3.4", "O framework de DataOps define os padrões que serão implementados nos pipelines CI/CD"],
    ["3.3 → 3.5", "A arquitetura de ingestão deve estar definida para implementar observabilidade"],
    ["3.4, 3.5 → 3.6", "Pipelines e observabilidade devem estar prontos para ativação"],
    ["Pilar 1 → 4.1", "O modelo de governança (Pilar 1) deve estar definido para desenhar o modelo operacional"],
    ["4.1 → 4.2", "O desenho do modelo operacional precede sua configuração"],
    ["4.2, 3.6 → 4.3", "A configuração e a ativação devem preceder a validação e transição final"],
]

add_styled_table(doc,
    ["Dependência", "Justificativa"],
    deps,
    col_widths=[4, 12]
)

doc.add_page_break()

# ============================================================
# 7. ALOCAÇÃO DE EQUIPE
# ============================================================
add_heading_styled(doc, "7. Alocação de Equipe", 1)

add_styled_table(doc,
    ["Período", "GP", "ADS", "EDS", "EDO", "CIO", "ESD", "ADP", "Total"],
    [
        ["Mês 1 (Ago)", "1", "1", "1", "—", "—", "—", "1", "4"],
        ["Mês 2 (Set)", "1", "1", "1", "1", "—", "—", "—", "4"],
        ["Mês 3 (Out)", "1", "1", "1", "1", "1", "—", "1", "6"],
        ["Mês 4 (Nov)", "1", "1", "1", "1", "1", "—", "1", "6"],
        ["Mês 5 (Dez)", "1", "1", "1", "1", "1", "1", "1", "7"],
        ["Mês 6 (Jan)", "1", "—", "—", "1", "1", "1", "1", "5"],
        ["Mês 7 (Fev)", "1", "1", "1", "1", "1", "—", "1", "6"],
        ["Mês 8 (Mar)", "1", "—", "—", "1", "1", "—", "1", "4"],
    ],
    col_widths=[3, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5]
)

add_para(doc, "")
add_para(doc, "Equipe média ao longo do projeto: 5-6 profissionais simultâneos", bold=True, size=10)
add_para(doc, "Equipe máxima (pico em Dezembro): 7 profissionais", bold=True, size=10)

doc.add_page_break()

# ============================================================
# 8. PREMISSAS E RISCOS
# ============================================================
add_heading_styled(doc, "8. Premissas e Riscos", 1)

add_heading_styled(doc, "8.1 Premissas Gerais", 2)

premissas = [
    "Horas calculadas para profissionais humanos sem uso de IA generativa",
    "Jornada de 8h/dia, 20 dias úteis/mês = 160h/mês por profissional",
    "Disponibilidade dos stakeholders Vivo para workshops e validações dentro dos prazos",
    "Acesso aos ambientes e documentação existente garantidos desde o início do projeto",
    "Pilares 1 (Governança) e 2 (Modelagem de Dados) sendo executados em paralelo por equipes dedicadas",
    "Ferramentas de ITSM, CI/CD e monitoramento já licenciadas e disponíveis no ambiente",
    "A integração com a ferramenta de ITSM pressupõe que seja possível implementá-la de forma nativa",
    "Infraestrutura de cloud/on-premises disponível conforme necessidades do projeto",
    "Equipe de operação da Telefônica definida e disponível para a fase de transição",
    "Não inclui horas de gestão comercial, pré-venda ou revisões executivas",
]

for p_text in premissas:
    add_bullet(doc, p_text)

add_heading_styled(doc, "8.2 Premissas por Objetivo", 2)

premissas_obj = [
    ["3.1", "Acesso ao ambiente e documentação; envolvimento da arquitetura corporativa; direcionamento estratégico prévio"],
    ["3.2", "Tooling compatível com automação; sincronia com Pilar 1"],
    ["3.3", "Acesso às fontes de dados; disponibilidade da infraestrutura"],
    ["3.4", "Existência de plataforma de CI/CD; acesso a ambientes dev/test/prod"],
    ["3.5", "Ferramentas de monitoramento disponíveis; sincronia com Pilar 1 para KPIs; aderência a políticas corporativas"],
    ["3.6", "Sincronia com Pilar 4 para transição"],
    ["4.1", "Modelo de governança (Pilar 1) definido; arquitetura e ferramentas especificadas"],
    ["4.2", "Ferramentas disponíveis; infraestrutura pronta; integração ITSM nativa"],
    ["4.3", "Equipe definida para operação; aprovação do modelo pela Telefônica"],
]

add_styled_table(doc,
    ["Objetivo", "Premissas Específicas"],
    premissas_obj,
    col_widths=[2, 14]
)

add_para(doc, "")

add_heading_styled(doc, "8.3 Principais Riscos", 2)

riscos = [
    ["Alto", "Indisponibilidade de stakeholders para validações", "Estabelecer calendário fixo de workshops no kick-off"],
    ["Alto", "Dependência de Pilares 1 e 2 não concluídos no prazo", "Checkpoints quinzenais de alinhamento entre pilares"],
    ["Médio", "Complexidade do ambiente maior que a estimada no assessment", "Buffer de contingência de 15% incluído"],
    ["Médio", "Mudanças de escopo durante a execução", "Processo formal de change request com impacto em horas"],
    ["Médio", "Incompatibilidade de ferramentas com automação prevista", "PoC técnica no início de cada fase de implementação"],
    ["Baixo", "Rotatividade da equipe Foursys", "Documentação contínua e knowledge base do projeto"],
]

add_styled_table(doc,
    ["Severidade", "Risco", "Mitigação"],
    riscos,
    col_widths=[2.5, 6, 7.5]
)

# ============================================================
# RODAPÉ FINAL
# ============================================================
doc.add_page_break()

add_heading_styled(doc, "Considerações Finais", 1)

add_para(doc, 
    "Esta proposta foi elaborada com base nas necessidades identificadas no documento de RFI "
    "e nas boas práticas de mercado para projetos de transformação de dados em ambientes "
    "corporativos de grande porte no setor de telecomunicações.")

add_para(doc,
    "As estimativas de horas refletem o esforço necessário para uma entrega de qualidade, "
    "considerando a complexidade do ambiente Telefônica/Vivo, as múltiplas integrações "
    "envolvidas e a necessidade de transição estruturada para operação contínua.")

add_para(doc,
    "Estamos à disposição para apresentar esta proposta em detalhes e discutir eventuais "
    "ajustes de escopo, cronograma ou equipe conforme necessário.")

for _ in range(3):
    doc.add_paragraph()

add_para(doc, "_______________________________________________", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Foursys Informática Ltda.", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Junho/2026", size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# SALVAR
# ============================================================
output_path = r"C:\01. Foursys\06. BMAD Cursor\14. RFI Vivo\Proposta Tecnica COE Data AI Tools - Foursys.docx"
doc.save(output_path)
print(f"Documento gerado com sucesso: {output_path}")

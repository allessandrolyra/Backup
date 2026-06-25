from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, val in kwargs.items():
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), val.get("val", "single"))
        element.set(qn("w:sz"), val.get("sz", "4"))
        element.set(qn("w:color"), val.get("color", "000000"))
        element.set(qn("w:space"), "0")
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_color = "1F4E79"
    alt_color = "D6E4F0"

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, alt_color)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table


doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

for level in range(1, 4):
    h_style = doc.styles[f"Heading {level}"]
    h_style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

# ============================================================
# CAPA
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DOCUMENTO DE ENTENDIMENTO")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Mapeamento de Arquitetura Corporativa\ne Rastreabilidade ALM")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Referencia Metodologica: TOGAF Standard, 10th Edition\n"
                "Padrao de Modelagem: ArchiMate 3.2\n"
                "Ecossistema: Azure DevOps | Neo4j | Archi (Git-based)")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
today = datetime.date.today().strftime("%d/%m/%Y")
run = p.add_run(f"Data: {today}\nVersao: 1.0 - Entendimento Inicial")
run.font.size = Pt(10)
run.font.italic = True

doc.add_page_break()

# ============================================================
# SUMARIO (manual)
# ============================================================
doc.add_heading("Sumario", level=1)

sumario_items = [
    ("1.", "Objetivo deste Documento", 2),
    ("2.", "Resumo Executivo da Necessidade do Cliente", 2),
    ("3.", "Escopo Detalhado por Camada ArchiMate", 2),
    ("3.1", "Camada de Negocio", 3),
    ("3.2", "Camada de Aplicacao", 3),
    ("3.3", "Camada de Tecnologia", 3),
    ("3.4", "Interdependencias Criticas", 3),
    ("4.", "Integracao com Azure DevOps (ALM)", 2),
    ("5.", "Plataforma de Analise de Impacto (Neo4j)", 2),
    ("6.", "Governanca e Versionamento", 2),
    ("7.", "Entregaveis e Criterios de Aceite", 2),
    ("8.", "Perguntas Pendentes ao Cliente", 2),
    ("9.", "Riscos Identificados", 2),
    ("10.", "Proximos Passos e Cronograma Sugerido", 2),
    ("11.", "Estimativa de Perfis e Horas", 2),
]

for num, title, indent_level in sumario_items:
    p = doc.add_paragraph()
    indent = "    " if indent_level == 3 else ""
    run = p.add_run(f"{indent}{num} {title}")
    run.font.size = Pt(10)

doc.add_page_break()

# ============================================================
# 1. OBJETIVO DESTE DOCUMENTO
# ============================================================
doc.add_heading("1. Objetivo deste Documento", level=1)

doc.add_paragraph(
    "Este documento tem como objetivo consolidar o entendimento da Ordem de Servico (OS) "
    "recebida do cliente, traduzindo os requisitos tecnicos em linguagem acessivel para "
    "todas as partes interessadas. Ele serve como base para:"
)

bullets = [
    "Alinhar o entendimento entre a equipe tecnica e o cliente",
    "Identificar lacunas de informacao que precisam ser esclarecidas",
    "Orientar o planejamento e a estimativa de esforco",
    "Servir como referencia durante toda a execucao do projeto",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_paragraph()

# ============================================================
# 2. RESUMO EXECUTIVO
# ============================================================
doc.add_heading("2. Resumo Executivo da Necessidade do Cliente", level=1)

doc.add_heading("2.1 O Que o Cliente Quer (Linguagem Simples)", level=2)

doc.add_paragraph(
    "O cliente deseja criar um \"Gemeo Digital\" da sua arquitetura corporativa - essencialmente "
    "um mapa completo e interativo que mostre como todos os sistemas, processos de negocio e "
    "tecnologias da organizacao se conectam entre si."
)

doc.add_paragraph(
    "Analogia: Imagine a empresa como uma cidade. O cliente quer um mapa digital dessa cidade "
    "que mostre todos os predios (sistemas), todas as ruas (integracoes), todos os moradores "
    "(atores/times) e que, ao clicar em qualquer ponto, consiga responder: "
    "\"Se eu mudar isso aqui, o que mais e afetado?\""
)

doc.add_heading("2.2 Os 4 Pilares do Projeto", level=2)

add_styled_table(doc,
    ["Pilar", "O Que E", "Ferramenta", "Por Que Importa"],
    [
        ["Modelagem", "Desenhar toda a arquitetura em camadas (negocio, app, tech)",
         "Archi + ArchiMate 3.2", "Visao unica e padronizada da empresa"],
        ["Versionamento", "Controlar versoes do modelo como codigo-fonte",
         "Git (Archi Collaboration Plugin)", "Historico, branching, colaboracao"],
        ["Rastreabilidade ALM", "Vincular cada componente ao codigo e work items",
         "Azure DevOps", "Saber o que esta sendo desenvolvido para cada componente"],
        ["Analise de Impacto", "Consultar dependencias e impactos em tempo real",
         "Neo4j (Grafo)", "Tomada de decisao tecnica baseada em dados"],
    ],
    [4, 6, 4, 5]
)

doc.add_heading("2.3 O Que NAO Esta no Escopo", level=2)

nao_escopo = [
    "Mapeamento de infraestrutura fisica (responsabilidade do CMDB)",
    "Detalhes de capacity planning (quantidade de CPUs, memoria, storage)",
    "Migracoes ou implementacoes de sistemas",
    "Desenvolvimento de novas aplicacoes",
]
for item in nao_escopo:
    doc.add_paragraph(item, style="List Bullet")

doc.add_page_break()

# ============================================================
# 3. ESCOPO DETALHADO POR CAMADA
# ============================================================
doc.add_heading("3. Escopo Detalhado por Camada ArchiMate", level=1)

doc.add_paragraph(
    "O ArchiMate organiza a arquitetura em tres camadas principais, mais uma transversal "
    "de integracao. Abaixo detalhamos o que deve ser mapeado em cada uma:"
)

# 3.1
doc.add_heading("3.1 Camada de Negocio (Business Layer)", level=2)

add_styled_table(doc,
    ["Elemento ArchiMate", "O Que Representa", "Exemplo", "Como Descobrir"],
    [
        ["Business Capability", "O que a empresa e capaz de fazer",
         "Gestao de Sinistros, Emissao de Apolices", "Entrevistas com diretores"],
        ["Business Process", "Como a capacidade e executada (fluxo)",
         "Processo de abertura de sinistro", "Entrevistas com coordenadores"],
        ["Business Function", "Agrupamento logico de comportamentos",
         "Funcao de Atendimento ao Cliente", "Organograma funcional"],
        ["Business Actor", "Quem executa (pessoa, time, parceiro)",
         "Corretor, Segurado, Regulador", "Stakeholder mapping"],
        ["Business Service", "O que e entregue ao consumidor",
         "Servico de Cotacao Online", "Catalogo de servicos"],
    ],
    [4, 5, 5, 5]
)

p = doc.add_paragraph()
run = p.add_run("ATENCAO: ")
run.bold = True
run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
run = p.add_run("A OS exige obrigatoriamente o mapeamento de interacoes com Parceiros Externos.")

# 3.2
doc.add_heading("3.2 Camada de Aplicacao (Application Layer)", level=2)

add_styled_table(doc,
    ["Elemento ArchiMate", "O Que Representa", "Exemplo", "Como Descobrir"],
    [
        ["Application Component", "Um sistema ou modulo de software",
         "Sistema de Cotacao, ERP, CRM", "Inventario de sistemas / Azure Repos"],
        ["Application Service", "Funcionalidade exposta por um componente",
         "API de calculo de premio", "API Gateway / Swagger docs"],
        ["Application Interface", "Ponto de acesso (API, UI, arquivo)",
         "REST endpoint /api/v1/cotacao", "Documentacao tecnica"],
        ["Data Object", "Dados manipulados pelas aplicacoes",
         "Apolice, Sinistro, Cliente", "Modelo de dados / DBA"],
    ],
    [4, 5, 5, 5]
)

# 3.3
doc.add_heading("3.3 Camada de Tecnologia (Technology Layer)", level=2)

add_styled_table(doc,
    ["Elemento ArchiMate", "O Que Representa", "Exemplo", "Nivel de Detalhe"],
    [
        ["Technology Node (logico)", "Onde os componentes executam",
         "AKS Cluster, Azure App Service", "Plataforma, NAO specs fisicas"],
        ["Technology Service", "Servico de infraestrutura",
         "Azure SQL Database, Service Bus", "Tipo de servico, NAO tier/sizing"],
        ["Artifact", "Artefato deployavel",
         "Container image, WAR file", "Apenas referencia logica"],
        ["Communication Network", "Meio de comunicacao",
         "VPN, ExpressRoute, Internet", "Topologia logica"],
    ],
    [4, 5, 5, 5]
)

# 3.4
doc.add_heading("3.4 Interdependencias Criticas (Cross-Layer)", level=2)

doc.add_paragraph(
    "Este e o mapeamento mais trabalhoso e mais valioso. Cada integracao entre sistemas "
    "deve ser documentada com tipo e protocolo:"
)

add_styled_table(doc,
    ["Tipo de Integracao", "Protocolo Tipico", "Exemplo", "Prioridade de Mapeamento"],
    [
        ["API sincrona", "REST/SOAP/gRPC", "Consulta de CEP em servico externo", "ALTA"],
        ["Troca de Arquivos (Batch)", "SFTP/Azure Blob/FTP", "Retorno de pagamentos do banco", "ALTA"],
        ["Mensageria", "Service Bus/RabbitMQ/Kafka", "Evento de nova apolice emitida", "ALTA"],
        ["Banco de dados compartilhado", "JDBC/ODBC direto", "Leitura direta de tabela de outro sistema", "MEDIA"],
        ["ETL/Data Pipeline", "ADF/SSIS/Talend", "Carga de dados para DW", "MEDIA"],
    ],
    [4, 4, 5, 4]
)

doc.add_page_break()

# ============================================================
# 4. INTEGRACAO COM AZURE DEVOPS
# ============================================================
doc.add_heading("4. Integracao com Azure DevOps (ALM)", level=1)

doc.add_paragraph(
    "Este e um REQUISITO MANDATORIO da OS. O objetivo e garantir que cada linha de codigo "
    "e cada item de trabalho esteja vinculado ao modelo de arquitetura."
)

doc.add_heading("4.1 O Que Deve Ser Feito", level=2)

add_styled_table(doc,
    ["Requisito", "Descricao", "Implementacao Sugerida"],
    [
        ["Vinculo Work Items", "Cada User Story/Feature/Task deve referenciar o ID do objeto ArchiMate",
         "Campo customizado 'ArchiMate ID' no Work Item Type"],
        ["Vinculo de Repositorios", "Cada repo no Azure Repos deve estar ligado ao componente logico",
         "Arquivo ARCHITECTURE.md na raiz de cada repo com o ID"],
        ["Pipeline de Sincronizacao", "Alteracoes no modelo Git devem atualizar o Neo4j automaticamente",
         "Pipeline YAML com trigger no repo do modelo Archi"],
    ],
    [4, 6, 6]
)

doc.add_heading("4.2 Fluxo de Sincronizacao", level=2)

doc.add_paragraph(
    "1. Arquiteto edita o modelo no Archi\n"
    "2. Archi Collaboration Plugin faz push para o Azure Repos\n"
    "3. Pipeline YAML detecta a alteracao (trigger)\n"
    "4. Script de conversao transforma XML (Open Exchange Format) em comandos Cypher\n"
    "5. Comandos Cypher sao executados contra o Neo4j\n"
    "6. Grafo atualizado disponivel para consultas de impacto"
)

doc.add_page_break()

# ============================================================
# 5. NEO4J
# ============================================================
doc.add_heading("5. Plataforma de Analise de Impacto (Neo4j)", level=1)

doc.add_paragraph(
    "O Neo4j e um banco de dados orientado a grafos. A escolha e adequada porque a arquitetura "
    "corporativa e, por natureza, um grafo de relacionamentos entre componentes."
)

doc.add_heading("5.1 Consultas Obrigatorias (conforme OS)", level=2)

add_styled_table(doc,
    ["Tipo de Consulta", "Pergunta que Responde", "Beneficio"],
    [
        ["Analise de Impacto", "Se eu mudar a API X, quais processos de negocio sao afetados?",
         "Evitar quebras em cascata"],
        ["Desenvolvimentos Sombra", "Existem repositorios no ADO sem correspondente no modelo?",
         "Identificar codigo nao governado"],
        ["Arquiteturas Orfas", "Existem componentes no modelo sem codigo associado?",
         "Identificar planejado vs. implementado"],
        ["Rastreabilidade Completa", "Para a capacidade X, quais sistemas, APIs e Work Items existem?",
         "Visao fim-a-fim"],
    ],
    [4, 7, 5]
)

# ============================================================
# 6. GOVERNANCA
# ============================================================
doc.add_heading("6. Governanca e Versionamento", level=1)

add_styled_table(doc,
    ["Requisito", "Detalhe", "Beneficio"],
    [
        ["Git-based", "Modelo versionado via Archi Collaboration Plugin no Azure Repos",
         "Historico completo, branching, merge, rollback"],
        ["Portabilidade", "Entrega em Open Exchange Format (XML)",
         "Evita vendor lock-in, modelo pode ser aberto em qualquer ferramenta ArchiMate"],
        ["Taxonomia Padronizada", "Manual de convencoes para nomenclatura e relacionamentos",
         "Garante que o time interno possa continuar o trabalho"],
    ],
    [4, 7, 5]
)

doc.add_page_break()

# ============================================================
# 7. ENTREGAVEIS
# ============================================================
doc.add_heading("7. Entregaveis e Criterios de Aceite", level=1)

add_styled_table(doc,
    ["#", "Entregavel", "Formato", "Criterio de Aceite"],
    [
        ["E1", "Repositorio Git com modelo ArchiMate completo",
         "Repo Azure Repos (.archimate)", "Modelo abre no Archi sem erros; todas as camadas populadas"],
        ["E2", "Scripts de Ingestao Cypher para Neo4j",
         "Arquivos .cypher + Pipeline YAML", "Pipeline executa e popula Neo4j com dados do modelo"],
        ["E3", "Matriz de Rastreabilidade Dinamica",
         "Consultas Cypher + Dashboard", "Navegacao Business > App > Tech > Work Item funcional"],
        ["E4", "Manual de Taxonomia",
         "Documento (Word/PDF/Wiki)", "Time interno consegue adicionar novos objetos seguindo o padrao"],
    ],
    [1, 5, 4, 6]
)

doc.add_page_break()

# ============================================================
# 8. PERGUNTAS PENDENTES
# ============================================================
doc.add_heading("8. Perguntas Pendentes ao Cliente", level=1)

doc.add_paragraph(
    "As respostas abaixo sao CRITICAS para dimensionar o esforco e elaborar a proposta. "
    "Sem elas, qualquer estimativa sera imprecisa."
)

doc.add_heading("8.1 Escopo e Dimensionamento", level=2)

perguntas_escopo = [
    ["PE-01", "Quantos sistemas/aplicacoes existem no portfolio atual?",
     "CRITICA", "Define o volume de trabalho na camada de aplicacao"],
    ["PE-02", "Quantas integracoes (APIs, batch, mensageria) existem aproximadamente?",
     "CRITICA", "Integracoes sao o item mais trabalhoso do mapeamento"],
    ["PE-03", "Ja existe documentacao de arquitetura? (diagramas, wikis, Visio, planilhas)",
     "CRITICA", "Determina se partimos do zero ou de uma base existente"],
    ["PE-04", "Quantos repositorios existem no Azure Repos?",
     "ALTA", "Dimensiona o vinculo ALM"],
    ["PE-05", "Quantos Work Items ativos existem (Features/Stories)?",
     "ALTA", "Dimensiona o vinculo ALM"],
    ["PE-06", "Quantas areas de negocio/departamentos serao mapeados?",
     "CRITICA", "Define quantidade de entrevistas e amplitude do mapeamento"],
    ["PE-07", "Quem serao os pontos focais por area para entrevistas?",
     "ALTA", "Planejamento de agenda e logistica"],
]

add_styled_table(doc,
    ["ID", "Pergunta", "Prioridade", "Por Que Precisamos Saber"],
    perguntas_escopo,
    [2, 7, 2, 5]
)

doc.add_heading("8.2 Profundidade e Limites", level=2)

perguntas_profundidade = [
    ["PP-01", "Qual o limite inferior de detalhe na camada de tecnologia? (Ex: 'roda em AKS' sim, 'qual node pool' nao?)",
     "ALTA", "Evita over-engineering no mapeamento"],
    ["PP-02", "Quantos parceiros externos existem? Temos acesso a docs de integracao?",
     "ALTA", "Parceiros sao obrigatorios no escopo"],
    ["PP-03", "O que significa 'granularidade suficiente para tomada de decisao'? Exemplo de decisao?",
     "ALTA", "Define o criterio de 'pronto' do mapeamento"],
]

add_styled_table(doc,
    ["ID", "Pergunta", "Prioridade", "Por Que Precisamos Saber"],
    perguntas_profundidade,
    [2, 7, 2, 5]
)

doc.add_heading("8.3 Tecnico / Infraestrutura", level=2)

perguntas_tech = [
    ["PT-01", "O Neo4j ja existe ou precisa ser provisionado? Onde? (Azure, on-prem)",
     "ALTA", "Impacta prazo e custo de setup"],
    ["PT-02", "O Archi Collaboration Plugin ja esta em uso ou sera implantado?",
     "MEDIA", "Treinamento pode ser necessario"],
    ["PT-03", "O Azure DevOps e instancia unica ou multiplas organizacoes/projetos?",
     "ALTA", "Complexidade de integracao"],
    ["PT-04", "Temos acesso para criar pipelines no Azure DevOps? Existe governanca de aprovacao?",
     "ALTA", "Pode ser bloqueante"],
    ["PT-05", "Existe CMDB hoje? Qual ferramenta?",
     "MEDIA", "Fronteira entre CMDB e modelo ArchiMate"],
]

add_styled_table(doc,
    ["ID", "Pergunta", "Prioridade", "Por Que Precisamos Saber"],
    perguntas_tech,
    [2, 7, 2, 5]
)

doc.add_heading("8.4 Governanca e Continuidade", level=2)

perguntas_gov = [
    ["PG-01", "Quem mantera o modelo apos a entrega? Existe arquiteto corporativo no time?",
     "ALTA", "Define profundidade do Manual de Taxonomia"],
    ["PG-02", "Qual padrao de nomenclatura/ID querem? Ja tem convencao?",
     "MEDIA", "Precisa ser definido antes de comecar"],
    ["PG-03", "Quantas rodadas de validacao/aceite por entrega?",
     "MEDIA", "Impacta cronograma"],
    ["PG-04", "Qual SLA para validacao de cada entrega? (prazo para o cliente aprovar)",
     "MEDIA", "Impacta cronograma total"],
]

add_styled_table(doc,
    ["ID", "Pergunta", "Prioridade", "Por Que Precisamos Saber"],
    perguntas_gov,
    [2, 7, 2, 5]
)

doc.add_page_break()

# ============================================================
# 9. RISCOS
# ============================================================
doc.add_heading("9. Riscos Identificados", level=1)

add_styled_table(doc,
    ["ID", "Risco", "Impacto", "Probabilidade", "Mitigacao Sugerida"],
    [
        ["R01", "Falta de documentacao existente - mapeamento partindo do zero",
         "ALTO", "ALTA", "Incluir horas adicionais de discovery e entrevistas"],
        ["R02", "Indisponibilidade dos pontos focais para entrevistas",
         "ALTO", "MEDIA", "Definir agenda antecipadamente com SLA de resposta"],
        ["R03", "Complexidade de integracoes maior que o estimado",
         "ALTO", "MEDIA", "Prever buffer de 20-30% na estimativa de integracoes"],
        ["R04", "Restricoes de acesso ao Azure DevOps para criar pipelines",
         "MEDIO", "MEDIA", "Validar acessos na fase de setup"],
        ["R05", "Time interno sem experiencia em ArchiMate para continuidade",
         "MEDIO", "ALTA", "Incluir sessoes de transferencia de conhecimento"],
        ["R06", "Escopo de 'parceiros externos' indefinido - pode crescer muito",
         "ALTO", "MEDIA", "Obter lista exaustiva de parceiros antes de estimar"],
        ["R07", "Neo4j nao provisionado - dependencia de infraestrutura",
         "MEDIO", "MEDIA", "Incluir provisionamento na fase de setup"],
    ],
    [1, 6, 2, 2, 6]
)

doc.add_page_break()

# ============================================================
# 10. PROXIMOS PASSOS
# ============================================================
doc.add_heading("10. Proximos Passos e Cronograma Sugerido", level=1)

doc.add_heading("10.1 Sequencia de Atividades (Faseamento)", level=2)

add_styled_table(doc,
    ["Fase", "Atividade", "Duracao Estimada*", "Dependencia", "Entregavel"],
    [
        ["F0", "Setup do Ambiente (Archi, Git, Neo4j, Pipelines)", "1-2 semanas",
         "Acesso ao ADO confirmado", "-"],
        ["F1", "Mapeamento da Camada de Negocio", "2-3 semanas",
         "Pontos focais definidos", "Capability Map + Processos"],
        ["F2", "Mapeamento da Camada de Aplicacao", "3-4 semanas",
         "F1 concluida", "Inventario de sistemas + Integracoes"],
        ["F3", "Mapeamento da Camada de Tecnologia", "2-3 semanas",
         "F2 concluida", "Mapa de infraestrutura logica"],
        ["F4", "Vinculacao Azure DevOps (ALM)", "2-3 semanas",
         "F2 concluida (paralelo com F3)", "Work Items + Repos vinculados"],
        ["F5", "Pipeline Neo4j + Queries de Impacto", "2-3 semanas",
         "F0 + F3 concluidas", "E2 - Scripts Cypher + Pipeline"],
        ["F6", "Consolidacao e Manual de Taxonomia", "1-2 semanas",
         "Todas anteriores", "E3 + E4"],
        ["F7", "Transferencia de Conhecimento", "1 semana",
         "F6 concluida", "Sessoes com time interno"],
    ],
    [1, 6, 3, 4, 4]
)

p = doc.add_paragraph()
run = p.add_run("* Estimativas preliminares - dependem das respostas do cliente (Secao 8)")
run.italic = True
run.font.size = Pt(9)

doc.add_heading("10.2 Cronograma Visual", level=2)

doc.add_paragraph(
    "Semana:  1-2      3-5       5-8       8-10      10-12     12-14     14-15\n"
    "         |        |         |         |         |         |         |\n"
    "    [F0 Setup] [F1 Negocio] [F2 Aplicacao] [F3 Tech] [F5 Neo4j] [F6 Manual]\n"
    "                                      [F4 ADO Link ]            [F7 Handover]"
)

doc.add_paragraph(
    "Prazo total estimado: 12 a 16 semanas (dependendo do volume)\n"
    "Inicio: 10 dias apos OS aceita e assinada (conforme clausula da OS)"
)

doc.add_page_break()

# ============================================================
# 11. ESTIMATIVA DE PERFIS
# ============================================================
doc.add_heading("11. Estimativa de Perfis e Horas", level=1)

doc.add_paragraph(
    "A tabela abaixo apresenta os perfis necessarios. As horas devem ser ajustadas apos "
    "obter as respostas da Secao 8. Os valores por hora devem seguir a tabela contratual."
)

add_styled_table(doc,
    ["Perfil", "Responsabilidade Principal", "Horas Estimadas*", "Fase de Atuacao"],
    [
        ["Arquiteto Corporativo (Lead)", "Liderar o mapeamento, definir taxonomia, validar modelo",
         "A definir", "F0 a F7"],
        ["Analista de Arquitetura", "Conduzir entrevistas, documentar processos e sistemas",
         "A definir", "F1 a F4"],
        ["Engenheiro de Dados / DevOps", "Pipeline de ingestao, scripts Cypher, automacao",
         "A definir", "F0, F5"],
        ["Especialista Azure DevOps", "Configurar pipelines, campos customizados, vinculacao ALM",
         "A definir", "F0, F4"],
        ["Gerente de Projeto", "Coordenacao, cronograma, comunicacao com cliente",
         "A definir", "F0 a F7"],
    ],
    [4, 7, 3, 3]
)

p = doc.add_paragraph()
run = p.add_run(
    "* As horas serao definidas apos a reuniao de esclarecimento com o cliente. "
    "O valor-hora de cada perfil segue a tabela de precos anexada ao contrato."
)
run.italic = True
run.font.size = Pt(9)

doc.add_paragraph()

doc.add_heading("Observacoes Comerciais (conforme OS)", level=2)

obs_comerciais = [
    "Trabalho 100% off-site (remoto)",
    "Pagamento realizado a cada entrega aceita",
    "Prazo para resposta a cotacao: 5 dias uteis apos recebimento",
    "Prazo para inicio do servico: 10 dias apos OS aceita e assinada",
    "Perfis usados apenas para composicao de custo total",
]
for obs in obs_comerciais:
    doc.add_paragraph(obs, style="List Bullet")

doc.add_page_break()

# ============================================================
# GLOSSARIO
# ============================================================
doc.add_heading("Glossario de Termos", level=1)

add_styled_table(doc,
    ["Termo", "Significado"],
    [
        ["TOGAF", "The Open Group Architecture Framework - metodologia de arquitetura corporativa"],
        ["ArchiMate 3.2", "Linguagem de modelagem visual para arquitetura corporativa (padrao The Open Group)"],
        ["Archi", "Ferramenta open-source para modelagem ArchiMate"],
        ["ALM", "Application Lifecycle Management - ciclo de vida do desenvolvimento de software"],
        ["Neo4j", "Banco de dados orientado a grafos, ideal para consultas de relacionamentos"],
        ["Cypher", "Linguagem de consulta do Neo4j (equivalente ao SQL para grafos)"],
        ["CMDB", "Configuration Management Database - banco de ativos fisicos de TI"],
        ["Work Item", "Item de trabalho no Azure DevOps (User Story, Feature, Bug, Task)"],
        ["Gemeo Digital", "Representacao digital de um ativo real (neste caso, a arquitetura da empresa)"],
        ["Open Exchange Format", "Formato XML padrao para intercambio de modelos ArchiMate"],
        ["Desenvolvimentos Sombra", "Codigo existente no ADO sem correspondente no modelo de arquitetura"],
        ["Arquiteturas Orfas", "Componentes modelados sem codigo/implementacao associada"],
    ],
    [4, 14]
)

# ============================================================
# SALVAR
# ============================================================
output_path = r"c:\01. Foursys\06. BMAD Cursor\docs\entendimento-os-archimate-alm.docx"
doc.save(output_path)
print(f"Documento salvo em: {output_path}")

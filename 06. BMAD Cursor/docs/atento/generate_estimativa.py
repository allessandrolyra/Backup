from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.color.rgb = RGBColor(0, 51, 102)
    heading_style.font.name = 'Calibri'

HEADER_BG = RGBColor(0, 51, 102)
HEADER_TEXT = RGBColor(255, 255, 255)
ALT_ROW_BG = RGBColor(230, 240, 250)
HIGHLIGHT_BG = RGBColor(255, 242, 204)
ACCENT_COLOR = RGBColor(0, 102, 153)


def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), str(color))
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.name = 'Calibri'
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size


def add_styled_table(doc, headers, rows, col_widths=None, highlight_last_row=False):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, header, bold=True, color=HEADER_TEXT, size=Pt(9))

    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            is_last = highlight_last_row and row_idx == len(rows) - 1
            if is_last:
                set_cell_shading(cell, HIGHLIGHT_BG)
                set_cell_text(cell, value, bold=True, size=Pt(9))
            else:
                if row_idx % 2 == 1:
                    set_cell_shading(cell, ALT_ROW_BG)
                set_cell_text(cell, value, size=Pt(9))

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Cm(width)

    return table


def add_spacer(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(size)
    p.paragraph_format.space_after = Pt(size)


# ============================================================
# CAPA
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ESTIMATIVA DE CUSTOS AZURE')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = HEADER_BG

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Atento — ASR Cloud Solution')
run.font.size = Pt(18)
run.font.color.rgb = ACCENT_COLOR

add_spacer(doc, 20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Reconhecimento de Voz (ASR) para CNPJ Alfanumérico')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

add_spacer(doc, 40)

info_items = [
    ('Cliente', 'Atento'),
    ('Fornecedor', 'Foursys'),
    ('Cloud', 'Microsoft Azure (Brazil South)'),
    ('Data', 'Abril 2026'),
    ('Versão', 'v1.0'),
    ('Classificação', 'Confidencial — Uso Interno Foursys + Atento'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)
    run = p.add_run(value)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# ============================================================
# 1. PREMISSAS
# ============================================================
doc.add_heading('1. Premissas Base', level=1)

p = doc.add_paragraph('As estimativas abaixo consideram as seguintes premissas operacionais:')

add_styled_table(doc,
    ['Parâmetro', 'Valor'],
    [
        ['Duração média de interação ASR por chamada', '10 segundos de áudio'],
        ['Janela de operação', '12 horas/dia (07h-19h)'],
        ['Dias por mês', '22 dias úteis'],
        ['Região Azure', 'Brazil South'],
        ['Pico de carga', 'Volume/min é o pico sustentado durante a janela'],
    ],
    col_widths=[8, 8]
)

add_spacer(doc)

# ============================================================
# 2. CENÁRIOS DE VOLUME
# ============================================================
doc.add_heading('2. Cenários de Volume', level=1)

p = doc.add_paragraph('Foram definidos 4 cenários para cobrir desde o MVP até operação em larga escala:')

add_styled_table(doc,
    ['Cenário', 'Chamadas/min', 'Chamadas/hora', 'Chamadas/dia (12h)', 'Chamadas/mês', 'Horas áudio/mês', 'Sessões concorrentes'],
    [
        ['MVP', '10', '600', '7.200', '158.400', '440h', '~2'],
        ['A', '100', '6.000', '72.000', '1.584.000', '4.400h', '~17'],
        ['B', '1.000', '60.000', '720.000', '15.840.000', '44.000h', '~167'],
        ['C', '10.000', '600.000', '7.200.000', '158.400.000', '440.000h', '~1.667'],
    ],
    col_widths=[2, 2.5, 2.5, 3, 3, 2.5, 2.5]
)

add_spacer(doc)

# ============================================================
# 3. SERVIÇOS AZURE DETALHADOS
# ============================================================
doc.add_heading('3. Serviços Azure — Detalhamento por Cenário', level=1)

# 3.1 VPN Gateway
doc.add_heading('3.1 VPN Gateway', level=2)
p = doc.add_paragraph('Conectividade segura entre o datacenter Atento e a Azure via túnel IPSec Site-to-Site.')

add_styled_table(doc,
    ['Campo', 'MVP (10/min)', 'A (100/min)', 'B (1.000/min)', 'C (10.000/min)'],
    [
        ['SKU', 'VpnGw1AZ', 'VpnGw2AZ', 'VpnGw3AZ', 'VpnGw5AZ'],
        ['Tipo', 'Zone-Redundant', 'Zone-Redundant', 'Zone-Redundant', 'Zone-Redundant'],
        ['Modo', 'Active-Passive', 'Active-Active', 'Active-Active', 'Active-Active'],
        ['Throughput máximo', '650 Mbps', '1.25 Gbps', '2.5 Gbps', '10 Gbps'],
        ['Banda necessária', '~1.3 Mbps', '~13 Mbps', '~130 Mbps', '~1.3 Gbps'],
        ['Data transfer outbound', '~5 GB', '~50 GB', '~500 GB', '~5.000 GB'],
    ],
    col_widths=[4, 3, 3, 3, 3]
)

add_spacer(doc)
p = doc.add_paragraph()
run = p.add_run('Cálculo de banda: ')
run.bold = True
run.font.size = Pt(9)
run = p.add_run('chamadas/min × 128 kbps (G.711) × 2 (bidirecional) + overhead.')
run.font.size = Pt(9)

add_spacer(doc)

# 3.2 Speech Services
doc.add_heading('3.2 Azure Speech Services (STT Real-Time)', level=2)
p = doc.add_paragraph('Motor de reconhecimento de voz — componente de maior impacto no custo total.')

add_styled_table(doc,
    ['Campo', 'MVP (10/min)', 'A (100/min)', 'B (1.000/min)', 'C (10.000/min)'],
    [
        ['SKU', 'Standard (S0)', 'Standard (S0)', 'Standard (S0)', 'Standard (S0)'],
        ['Horas de áudio/mês', '440h', '4.400h', '44.000h', '440.000h'],
        ['Sessões concorrentes (pico)', '~2', '~17', '~167', '~1.667'],
        ['Custom Speech Hosting', '1 endpoint', '1 endpoint', '1 endpoint', '3 endpoints'],
        ['Preço referência', '~R$ 5,20/hora de áudio', '~R$ 5,20/hora', '~R$ 5,20/hora', '~R$ 5,20/hora'],
    ],
    col_widths=[4, 3, 3, 3, 3]
)

add_spacer(doc)

p = doc.add_paragraph()
run = p.add_run('Nota: ')
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(180, 0, 0)
run = p.add_run('Para cenários B e C, verificar limite de sessões concorrentes do S0 em Brazil South. Pode ser necessário solicitar aumento de quota à Microsoft ou usar múltiplos endpoints.')
run.font.size = Pt(9)

add_spacer(doc)

# 3.3 Compute
doc.add_heading('3.3 Compute (App Service / AKS)', level=2)

add_styled_table(doc,
    ['Campo', 'MVP (10/min)', 'A (100/min)', 'B (1.000/min)', 'C (10.000/min)'],
    [
        ['Plataforma', 'App Service', 'App Service', 'AKS', 'AKS'],
        ['SKU / VM', 'P1v3 (2vCPU, 8GB)', 'P2v3 (4vCPU, 16GB)', 'D4s_v5 (4vCPU, 16GB)', 'D8s_v5 (8vCPU, 32GB)'],
        ['Instâncias / Nodes', '2', '4 (autoscale 2-8)', '8-15 (autoscale)', '30-50 (autoscale)'],
        ['OS', 'Linux', 'Linux', 'Linux', 'Linux'],
        ['Discos gerenciados', 'N/A', 'N/A', 'Premium SSD 128GB × nós', 'Premium SSD 256GB × nós'],
        ['Node pool burst (spot)', 'N/A', 'N/A', 'N/A', 'D4s_v5 × 20 (pico)'],
    ],
    col_widths=[4, 3, 3, 3, 3]
)

add_spacer(doc)

# 3.4 API Management
doc.add_heading('3.4 API Management', level=2)

add_styled_table(doc,
    ['Campo', 'MVP (10/min)', 'A (100/min)', 'B (1.000/min)', 'C (10.000/min)'],
    [
        ['Tier', 'Standard v2', 'Standard v2', 'Standard v2', 'Premium'],
        ['Unidades', '1', '1', '2', '4+'],
        ['Chamadas/mês', '158.400', '1.584.000', '15.840.000', '158.400.000'],
    ],
    col_widths=[4, 3, 3, 3, 3]
)

add_spacer(doc)

# 3.5 Redis
doc.add_heading('3.5 Azure Cache for Redis', level=2)

add_styled_table(doc,
    ['Campo', 'MVP (10/min)', 'A (100/min)', 'B (1.000/min)', 'C (10.000/min)'],
    [
        ['Tier', 'Standard', 'Standard', 'Premium', 'Premium'],
        ['SKU', 'C1 (1 GB)', 'C2 (6 GB)', 'P2 (13 GB)', 'P4 (53 GB)'],
        ['Réplicas', '1', '1 (Zone-Redundant)', '2 (Zone-Redundant)', '3 (ZR + Cluster)'],
    ],
    col_widths=[4, 3, 3, 3, 3]
)

add_spacer(doc)

# 3.6 Monitor
doc.add_heading('3.6 Azure Monitor + Application Insights + Log Analytics', level=2)

add_styled_table(doc,
    ['Campo', 'MVP (10/min)', 'A (100/min)', 'B (1.000/min)', 'C (10.000/min)'],
    [
        ['Logs ingeridos (GB/mês)', '~1 GB', '~10 GB', '~100 GB', '~1.000 GB'],
        ['App Insights', 'Basic', 'Basic', 'Basic', 'Basic'],
        ['Retenção de logs', '90 dias', '90 dias', '90 dias', '30 dias'],
        ['Regras de alerta', '5', '10', '20', '50'],
    ],
    col_widths=[4, 3, 3, 3, 3]
)

add_spacer(doc)

p = doc.add_paragraph()
run = p.add_run('Nota: ')
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(180, 0, 0)
run = p.add_run('Log Analytics em alto volume é significativamente caro. Cenário C recomenda ingestão seletiva e uso de Basic Logs.')
run.font.size = Pt(9)

add_spacer(doc)

# 3.7 Outros
doc.add_heading('3.7 Key Vault, Private Endpoints e Bandwidth', level=2)

add_styled_table(doc,
    ['Campo', 'MVP (10/min)', 'A (100/min)', 'B (1.000/min)', 'C (10.000/min)'],
    [
        ['Key Vault (Standard)', 'R$ 20', 'R$ 50', 'R$ 50', 'R$ 50'],
        ['Private Endpoints (4 PEs)', 'R$ 150', 'R$ 200', 'R$ 500', 'R$ 2.000'],
        ['Bandwidth outbound', 'R$ 0 (free tier)', 'R$ 0 (free tier)', 'R$ 200', 'R$ 2.200'],
        ['DDoS Protection', '—', '—', '—', 'R$ 15.000'],
    ],
    col_widths=[4, 3, 3, 3, 3]
)

add_spacer(doc)

# ============================================================
# 4. TABELA COMPARATIVA CONSOLIDADA
# ============================================================
doc.add_page_break()
doc.add_heading('4. Estimativa Consolidada — Comparativo dos 4 Cenários', level=1)

p = doc.add_paragraph('Valores mensais estimados em Reais (BRL) — referência Abril/2026:')

add_styled_table(doc,
    ['Componente', 'MVP (10/min)', 'A (100/min)', 'B (1.000/min)', 'C (10.000/min)'],
    [
        ['VPN Gateway', 'R$ 1.800', 'R$ 2.800', 'R$ 5.500', 'R$ 12.000'],
        ['Speech Services STT', 'R$ 2.300', 'R$ 23.500', 'R$ 235.000', 'R$ 2.350.000'],
        ['Custom Speech Hosting', 'R$ 700', 'R$ 700', 'R$ 700', 'R$ 2.100'],
        ['Compute (App Service/AKS)', 'R$ 1.000', 'R$ 2.800', 'R$ 15.000', 'R$ 80.000'],
        ['API Management', 'R$ 1.800', 'R$ 1.800', 'R$ 3.600', 'R$ 18.000'],
        ['Redis Cache', 'R$ 400', 'R$ 800', 'R$ 3.500', 'R$ 12.000'],
        ['Monitor + Logs', 'R$ 200', 'R$ 500', 'R$ 3.000', 'R$ 15.000'],
        ['Key Vault', 'R$ 20', 'R$ 50', 'R$ 50', 'R$ 50'],
        ['Private Endpoints', 'R$ 150', 'R$ 200', 'R$ 500', 'R$ 2.000'],
        ['Bandwidth', 'R$ 0', 'R$ 0', 'R$ 200', 'R$ 2.200'],
        ['DDoS Protection', '—', '—', '—', 'R$ 15.000'],
        ['TOTAL MENSAL ESTIMADO', 'R$ 8.370', 'R$ 33.150', 'R$ 267.050', 'R$ 2.508.350'],
    ],
    col_widths=[4, 3, 3, 3, 3],
    highlight_last_row=True
)

add_spacer(doc)

# ============================================================
# 5. PROPORÇÃO DE CUSTO
# ============================================================
doc.add_heading('5. Análise de Proporção de Custo', level=1)

p = doc.add_paragraph('O Azure Speech Services é o componente dominante em todos os cenários acima do MVP:')

add_styled_table(doc,
    ['Cenário', 'Total Mensal', '% Speech Services', '% Compute', '% Rede (VPN+BW)', '% Outros'],
    [
        ['MVP (10/min)', 'R$ 8.370', '28%', '12%', '21%', '39%'],
        ['A (100/min)', 'R$ 33.150', '71%', '8%', '8%', '13%'],
        ['B (1.000/min)', 'R$ 267.050', '88%', '6%', '2%', '4%'],
        ['C (10.000/min)', 'R$ 2.508.350', '94%', '3%', '1%', '2%'],
    ],
    col_widths=[3, 3, 3, 3, 3, 3]
)

add_spacer(doc)

# ============================================================
# 6. TCO
# ============================================================
doc.add_heading('6. TCO — Custo Total de Propriedade (12 meses)', level=1)

p = doc.add_paragraph('Projeção anual considerando apenas custo de infraestrutura Azure (sem serviços profissionais):')

add_styled_table(doc,
    ['Métrica', 'MVP (10/min)', 'A (100/min)', 'B (1.000/min)', 'C (10.000/min)'],
    [
        ['Custo mensal', 'R$ 8.370', 'R$ 33.150', 'R$ 267.050', 'R$ 2.508.350'],
        ['TCO 12 meses', 'R$ 100.440', 'R$ 397.800', 'R$ 3.204.600', 'R$ 30.100.200'],
        ['Custo por chamada', 'R$ 0,053', 'R$ 0,021', 'R$ 0,017', 'R$ 0,016'],
    ],
    col_widths=[4, 3, 3, 3, 3],
    highlight_last_row=True
)

add_spacer(doc)

p = doc.add_paragraph()
run = p.add_run('Observação: ')
run.bold = True
run.font.size = Pt(9)
run = p.add_run('O custo por chamada diminui com o aumento de volume devido aos custos fixos (VPN, APIM, Redis) serem diluídos. Porém o Speech Services é pay-per-use linear, limitando a economia de escala.')
run.font.size = Pt(9)

add_spacer(doc)

# ============================================================
# 7. RECOMENDAÇÕES
# ============================================================
doc.add_heading('7. Recomendações de Otimização de Custo', level=1)

recommendations = [
    ('Enterprise Agreement / Reservas', 'Negociar com a Microsoft desconto de 30-50% no Speech Services via EA ou Reserved Capacity. Impacto significativo nos cenários B e C.'),
    ('Alternativa Self-Hosted (Cenário C)', 'Para 10.000/min, avaliar Whisper (OpenAI) self-hosted no AKS — troca custo pay-per-use do Speech por custo de compute (potencialmente menor).'),
    ('Log Analytics Seletivo', 'Cenários B e C: usar Basic Logs para dados de telemetria volumosos e reservar Analytics Logs apenas para dados de troubleshooting.'),
    ('Reserved Instances (Compute)', 'Para AKS nos cenários B e C, usar Reserved VM Instances (1 ou 3 anos) para os nodes base — economia de 30-60%.'),
    ('Spot Instances (Burst)', 'Cenário C: usar Spot VMs para o node pool de burst/pico — economia de até 80% no compute de pico.'),
    ('Volume Real do Cliente', 'A variável mais impactante é o volume real de chamadas. Solicitar dados históricos ao cliente antes de fechar proposta.'),
]

for title, desc in recommendations:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT_COLOR
    run = p.add_run(desc)
    run.font.size = Pt(10)

add_spacer(doc)

# ============================================================
# 8. CHECKLIST CALCULADORA
# ============================================================
doc.add_heading('8. Checklist para Azure Pricing Calculator', level=1)

p = doc.add_paragraph('Serviços a incluir na calculadora (azure.microsoft.com/pricing/calculator):')

checklist_items = [
    'VPN Gateway (VpnGw1AZ / 2AZ / 3AZ / 5AZ conforme cenário)',
    'Speech Services — STT Real-Time (horas de áudio: 440 / 4.400 / 44.000 / 440.000)',
    'Speech Services — Custom Speech (1-3 endpoints hosting)',
    'App Service P1v3/P2v3 (MVP e A) OU AKS com node pools (B e C)',
    'API Management Standard v2 (MVP/A/B) OU Premium (C)',
    'Azure Cache for Redis — Standard C1/C2 (MVP/A) OU Premium P2/P4 (B/C)',
    'Azure Monitor + Application Insights + Log Analytics (GB ingeridos/mês)',
    'Key Vault Standard',
    'Virtual Network — Private Endpoints (4 endpoints + dados processados)',
    'Bandwidth — Outbound Data Transfer (GB/mês)',
    'DDoS Network Protection (cenário C apenas)',
]

for item in checklist_items:
    p = doc.add_paragraph(f'☑  {item}')
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(9)

add_spacer(doc)

# ============================================================
# 9. RESSALVAS
# ============================================================
doc.add_heading('9. Ressalvas', level=1)

disclaimers = [
    'Valores estimados com base nos preços públicos Azure para a região Brazil South em Abril/2026.',
    'Preços devem ser validados na Azure Pricing Calculator no momento da proposta.',
    'Custos de serviços profissionais (implementação, operação SaaS) não estão incluídos neste documento.',
    'O volume real de chamadas do cliente é a variável de maior impacto — confirmar antes de fechar proposta.',
    'Para cenários B e C, negociar Enterprise Agreement com a Microsoft pode reduzir significativamente os custos.',
    'Custos de transferência de dados podem variar conforme o codec de áudio e overhead de protocolo.',
]

for d in disclaimers:
    p = doc.add_paragraph(d, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(9)

add_spacer(doc)

# ============================================================
# RODAPÉ
# ============================================================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('─' * 60)
run.font.color.rgb = RGBColor(180, 180, 180)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documento produzido pela Squad MEQ — Foursys')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(120, 120, 120)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Revisão: v1.0 — Abril 2026 | Classificação: Confidencial')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(120, 120, 120)

# ============================================================
# SALVAR
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), 'Estimativa Azure.docx')
doc.save(output_path)
print(f'Documento salvo em: {output_path}')

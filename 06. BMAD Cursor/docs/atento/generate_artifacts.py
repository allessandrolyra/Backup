"""
Atento ASR Cloud — Gerador de Artefatos v3
- .drawio com C4 (html=1 em todos os estilos) + Solution Design em fluxo limpo
- .docx documentação completa

v3.0 - Fix: html=1 em todos os estilos; labels limpos; Solution Design em fluxo com setas
"""

import xml.etree.ElementTree as ET
from xml.dom import minidom
import os

# ═══════════════════════════════════════════════════════════════════════════════
# ESTILOS — TODOS com html=1 para renderizar labels corretamente
# ═══════════════════════════════════════════════════════════════════════════════

S_PERSON = (
    "html=1;shape=mxgraph.c4.person2;whiteSpace=wrap;container=1;"
    "fontSize=11;fontColor=#ffffff;align=center;verticalAlign=middle;"
    "fillColor=#08427B;strokeColor=#073B6F;metaEdit=1;"
    "points=[[0.5,0,0],[1,0.5,0],[1,0.75,0],[0.75,1,0],"
    "[0.5,1,0],[0.25,1,0],[0,0.75,0],[0,0.5,0]];"
)

S_SYSTEM = (
    "html=1;rounded=1;whiteSpace=wrap;container=1;"
    "fontSize=11;fontColor=#ffffff;align=center;verticalAlign=middle;"
    "arcSize=10;fillColor=#1168BD;strokeColor=#0B4884;metaEdit=1;"
)

S_EXT = (
    "html=1;rounded=1;whiteSpace=wrap;container=1;"
    "fontSize=11;fontColor=#ffffff;align=center;verticalAlign=middle;"
    "arcSize=10;fillColor=#999999;strokeColor=#8B8B8B;metaEdit=1;dashed=1;"
)

S_CONTAINER = (
    "html=1;rounded=1;whiteSpace=wrap;container=1;"
    "fontSize=11;fontColor=#ffffff;align=center;verticalAlign=middle;"
    "arcSize=10;fillColor=#438DD5;strokeColor=#3C7FC0;metaEdit=1;"
)

S_COMPONENT = (
    "html=1;rounded=1;whiteSpace=wrap;container=1;"
    "fontSize=11;fontColor=#333333;align=center;verticalAlign=middle;"
    "arcSize=10;fillColor=#85BBF0;strokeColor=#78A8D8;metaEdit=1;"
)

S_BOUNDARY = (
    "html=1;rounded=1;fontSize=13;whiteSpace=wrap;container=1;"
    "fillColor=none;strokeColor=#444444;dashed=1;dashPattern=5 5;"
    "fontColor=#444444;fontStyle=1;align=left;verticalAlign=top;"
    "spacing=15;spacingTop=0;"
)

S_REL = (
    "html=1;endArrow=blockThin;fontSize=10;fontColor=#404040;"
    "strokeWidth=1;endFill=1;strokeColor=#828282;elbow=vertical;"
    "metaEdit=1;endSize=14;startSize=14;jumpStyle=arc;jumpSize=16;"
    "rounded=0;edgeStyle=orthogonalEdgeStyle;"
)

S_TITLE = (
    "html=1;text;align=center;verticalAlign=middle;resizable=0;points=[];"
    "autosize=1;strokeColor=none;fillColor=none;fontSize=20;fontStyle=1;"
    "fontColor=#333333;"
)

S_DEPLOY = (
    "html=1;rounded=1;fontSize=12;whiteSpace=wrap;container=1;fillColor=#F5F5F5;"
    "strokeColor=#666666;fontColor=#333333;fontStyle=1;align=left;"
    "verticalAlign=top;spacing=15;spacingTop=0;"
)

S_DEPLOY_AZ = (
    "html=1;rounded=1;fontSize=12;whiteSpace=wrap;container=1;fillColor=#E6F2FF;"
    "strokeColor=#0078D4;fontColor=#0078D4;fontStyle=1;align=left;"
    "verticalAlign=top;spacing=15;spacingTop=0;"
)

# Solution Design — boxes simples
S_FLOW_BOX = (
    "html=1;rounded=1;whiteSpace=wrap;fontSize=13;fontStyle=1;"
    "align=center;verticalAlign=middle;arcSize=12;strokeWidth=2;"
)

S_FLOW_ARROW = (
    "html=1;endArrow=blockThin;fontSize=11;strokeWidth=2;endFill=1;"
    "endSize=12;startSize=12;rounded=1;curved=1;fontColor=#333333;"
)

S_FLOW_ZONE = (
    "html=1;rounded=1;fontSize=14;whiteSpace=wrap;container=1;"
    "fontStyle=1;align=left;verticalAlign=top;spacing=15;spacingTop=0;strokeWidth=2;"
)


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

_counter = [2]

def nid():
    _counter[0] += 1
    return str(_counter[0])

def reset():
    _counter[0] = 2


def lbl(name, tag, desc=""):
    """Label C4 limpo — 3 linhas separadas, sem tags complexas."""
    lines = [f"<b>{name}</b>", f"[{tag}]"]
    if desc:
        lines.append(f"<br/>{desc}")
    return "<br/>".join(lines[:2]) + (f"<br/><br/>{desc}" if desc else "")


def mk(pw=1169, ph=827):
    m = ET.Element("mxGraphModel",
        dx="1422", dy="762", grid="1", gridSize="10", guides="1",
        tooltips="1", connect="1", arrows="1", fold="1", page="1",
        pageScale="1", pageWidth=str(pw), pageHeight=str(ph), math="0", shadow="0")
    ro = ET.SubElement(m, "root")
    ET.SubElement(ro, "mxCell", id="0")
    ET.SubElement(ro, "mxCell", id="1", parent="0")
    return m, ro


def bx(ro, cid, val, sty, x, y, w, h, par="1"):
    c = ET.SubElement(ro, "mxCell", id=cid, value=val, style=sty, vertex="1", parent=par)
    g = ET.SubElement(c, "mxGeometry", x=str(x), y=str(y), width=str(w), height=str(h))
    g.set("as", "geometry")


def ed(ro, eid, src, tgt, val="", sty=S_REL, par="1"):
    c = ET.SubElement(ro, "mxCell", id=eid, value=val, style=sty, edge="1", source=src, target=tgt, parent=par)
    g = ET.SubElement(c, "mxGeometry", relative="1")
    g.set("as", "geometry")


def ar(ro, eid, val, sty, sx, sy, tx, ty, par="1"):
    c = ET.SubElement(ro, "mxCell", id=eid, value=val, style=sty, edge="1", parent=par)
    g = ET.SubElement(c, "mxGeometry", relative="1")
    g.set("as", "geometry")
    s = ET.SubElement(g, "mxPoint", x=str(sx), y=str(sy)); s.set("as", "sourcePoint")
    t = ET.SubElement(g, "mxPoint", x=str(tx), y=str(ty)); t.set("as", "targetPoint")


def gr(ro, gid, x, y, w, h, val="", sty=S_BOUNDARY, par="1"):
    bx(ro, gid, val, sty, x, y, w, h, par)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — C4 Level 1: System Context
# ═══════════════════════════════════════════════════════════════════════════════

def p1():
    reset(); m, r = mk(1500, 1000)
    bx(r, nid(), "Atento ASR Cloud - System Context [C4 Level 1]", S_TITLE, 400, 15, 700, 35)

    p_cli = nid()
    bx(r, p_cli, lbl("Cliente Final", "Pessoa", "Liga para a central e<br/>diz o CNPJ por voz"), S_PERSON, 140, 60, 220, 200)

    p_op = nid()
    bx(r, p_op, lbl("Operador Atento", "Pessoa", "Recebe a chamada<br/>ja identificada"), S_PERSON, 1140, 60, 220, 200)

    s_asr = nid()
    bx(r, s_asr, lbl("Foursys ASR Cloud", "Software System", "Plataforma SaaS na Azure<br/>Transforma audio em texto<br/>(CNPJ) em tempo real"), S_SYSTEM, 530, 360, 360, 220)

    s_av = nid()
    bx(r, s_av, lbl("Plataforma Avaya", "Sistema Externo", "URA/IVR, PABX, SBC<br/>Telefonia on-premises"), S_EXT, 60, 400, 320, 180)

    s_bk = nid()
    bx(r, s_bk, lbl("Backend Atento", "Sistema Externo", "CRM, cadastro,<br/>validacao de CNPJ"), S_EXT, 1080, 400, 300, 180)

    s_sp = nid()
    bx(r, s_sp, lbl("Azure Speech Services", "Sistema Externo", "Motor ASR Speech-to-Text<br/>da Microsoft - pt-BR"), S_EXT, 540, 700, 340, 180)

    ed(r, nid(), p_cli, s_av, "Liga por voz [PSTN/SIP]")
    ed(r, nid(), s_av, s_asr, "Envia audio [VPN S2S / WebSocket]")
    ed(r, nid(), s_asr, s_av, "Retorna CNPJ transcrito [REST]")
    ed(r, nid(), s_asr, s_sp, "Audio para texto [Speech SDK]")
    ed(r, nid(), s_av, s_bk, "Valida CNPJ [API interna]")
    ed(r, nid(), s_av, p_op, "Transfere chamada [SIP/CTI]")
    return m


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — C4 Level 2: Container
# ═══════════════════════════════════════════════════════════════════════════════

def p2():
    reset(); m, r = mk(2000, 1300)
    bx(r, nid(), "Foursys ASR Cloud - Container Diagram [C4 Level 2]", S_TITLE, 600, 15, 700, 35)

    p1 = nid()
    bx(r, p1, lbl("Cliente Final", "Pessoa", "Diz o CNPJ por voz"), S_PERSON, 40, 100, 220, 200)

    ba = nid()
    gr(r, ba, 30, 370, 380, 260, "Atento - On-Premises", S_BOUNDARY)
    ca = nid()
    bx(r, ca, lbl("Plataforma Avaya", "Container: SBC + URA/IVR", "Captura audio e<br/>envia para ASR Cloud"),
       S_EXT.replace("fillColor=#999999", "fillColor=#666666"), 30, 50, 320, 170, par=ba)

    bz = nid()
    gr(r, bz, 500, 70, 1440, 1150, "Azure - Brazil South [Foursys ASR Cloud]",
       S_BOUNDARY.replace("strokeColor=#444444", "strokeColor=#0078D4").replace("fontColor=#444444", "fontColor=#0078D4"))

    cv = nid()
    bx(r, cv, lbl("VPN Gateway", "Container: VpnGw2AZ", "Zone-Redundant<br/>Tunel IPSec S2S"), S_CONTAINER, 40, 80, 300, 170, par=bz)

    cn = nid()
    bx(r, cn, lbl("Virtual Network", "Container: VNet 10.100.0.0/16", "Rede isolada<br/>Subnets + NSGs + PEs"), S_CONTAINER, 420, 80, 300, 170, par=bz)

    cp = nid()
    bx(r, cp, lbl("API Gateway", "Container: Azure API Management", "Autenticacao<br/>Rate limiting<br/>Roteamento"), S_CONTAINER, 40, 330, 320, 180, par=bz)

    co = nid()
    bx(r, co, lbl("ASR Orchestrator", "Container: App Service / AKS", "Orquestra fluxo completo<br/>Audio > ASR > Valida CNPJ<br/>> Retorna resultado"),
       S_CONTAINER.replace("fillColor=#438DD5", "fillColor=#2D6EB5"), 440, 330, 340, 200, par=bz)

    cr = nid()
    bx(r, cr, lbl("Session Cache", "Container: Azure Redis Cache", "Sessoes ativas<br/>CNPJs transcritos"), S_CONTAINER, 880, 80, 300, 170, par=bz)

    ck = nid()
    bx(r, ck, lbl("Secrets Store", "Container: Azure Key Vault", "Chaves, certificados<br/>Secrets"), S_CONTAINER, 880, 330, 300, 160, par=bz)

    cm = nid()
    bx(r, cm, lbl("Observability", "Container: Monitor + App Insights", "Metricas, logs, alertas<br/>Dashboards"),
       S_CONTAINER.replace("fillColor=#438DD5", "fillColor=#5B9BD5"), 880, 570, 320, 170, par=bz)

    cs = nid()
    bx(r, cs, lbl("Azure Speech Services", "External PaaS: STT Real-Time", "Custom Speech Model pt-BR<br/>CNPJ alfanumerico"), S_EXT, 380, 630, 340, 180, par=bz)

    ed(r, nid(), p1, ca, "Liga e fala [PSTN/SIP]")
    ed(r, nid(), ca, cv, "Audio stream [IPSec VPN]")
    ed(r, nid(), cv, cn, "Trafego interno")
    ed(r, nid(), cn, cp, "Roteia [Private Subnet]")
    ed(r, nid(), cp, co, "Request [HTTPS/WebSocket]")
    ed(r, nid(), co, cs, "Audio para texto [Speech SDK]")
    ed(r, nid(), co, cr, "Sessao e cache [Redis]")
    ed(r, nid(), co, ck, "Secrets [Managed Identity]")
    ed(r, nid(), co, cm, "Telemetria [App Insights]")
    return m


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — C4 Level 3: Component
# ═══════════════════════════════════════════════════════════════════════════════

def p3():
    reset(); m, r = mk(2000, 1300)
    bx(r, nid(), "ASR Orchestrator - Component Diagram [C4 Level 3]", S_TITLE, 600, 15, 700, 35)

    bo = nid()
    gr(r, bo, 300, 60, 1300, 1100, "ASR Orchestrator [Container: App Service / AKS]",
       S_BOUNDARY.replace("strokeColor=#444444", "strokeColor=#3C7FC0").replace("fontColor=#444444", "fontColor=#3C7FC0"))

    # Row 1
    c1 = nid()
    bx(r, c1, lbl("WebSocket Handler", "Component: ASP.NET", "Conexao WebSocket<br/>Stream bidirecional"), S_COMPONENT, 40, 70, 300, 150, par=bo)
    c2 = nid()
    bx(r, c2, lbl("Session Manager", "Component: Interno", "Cria/recupera sessao<br/>Correlaciona call-id"), S_COMPONENT, 420, 70, 300, 150, par=bo)
    c3 = nid()
    bx(r, c3, lbl("Config Provider", "Component: Interno", "Secrets do Key Vault<br/>Feature flags"), S_COMPONENT, 800, 70, 300, 150, par=bo)

    # Row 2
    c4 = nid()
    bx(r, c4, lbl("Audio Preprocessor", "Component: Interno", "G.711 para PCM 16kHz<br/>Voice Activity Detection"), S_COMPONENT, 40, 310, 300, 150, par=bo)
    c5 = nid()
    bx(r, c5, lbl("Speech Client", "Component: Azure Speech SDK", "Sessao STT real-time<br/>Resultados parciais e finais"), S_COMPONENT, 420, 310, 300, 150, par=bo)
    c6 = nid()
    bx(r, c6, lbl("Metrics Emitter", "Component: App Insights SDK", "Latencia, taxa sucesso<br/>Erros, chamadas"), S_COMPONENT, 800, 310, 300, 150, par=bo)

    # Row 3
    c7 = nid()
    bx(r, c7, lbl("CNPJ Parser", "Component: Interno", "Interpreta texto<br/>Extrai CNPJ<br/>Variacoes foneticas"), S_COMPONENT, 180, 560, 300, 160, par=bo)
    c8 = nid()
    bx(r, c8, lbl("CNPJ Validator", "Component: Interno", "Valida estrutura e<br/>digitos verificadores<br/>CNPJ alfanumerico"), S_COMPONENT, 580, 560, 300, 160, par=bo)

    # Row 4
    c9 = nid()
    bx(r, c9, lbl("Response Builder", "Component: Interno", "Resposta JSON:<br/>CNPJ + confianca + status"), S_COMPONENT, 180, 820, 300, 150, par=bo)
    c10 = nid()
    bx(r, c10, lbl("Retry and Fallback", "Component: Interno", "Re-prompt max 3 tentativas<br/>Fallback para operador"), S_COMPONENT, 580, 820, 300, 150, par=bo)

    # Externals
    ea = nid()
    bx(r, ea, lbl("API Gateway", "External: Azure APIM", ""), S_EXT, 40, 200, 220, 100)
    er = nid()
    bx(r, er, lbl("Session Cache", "External: Azure Redis", ""), S_EXT, 40, 400, 220, 100)
    es = nid()
    bx(r, es, lbl("Speech Services", "External: STT Real-Time", ""), S_EXT, 40, 580, 220, 100)
    ek = nid()
    bx(r, ek, lbl("Key Vault", "External: Secrets", ""), S_EXT, 1650, 130, 220, 100)
    ei = nid()
    bx(r, ei, lbl("App Insights", "External: Telemetria", ""), S_EXT, 1650, 370, 220, 100)

    ed(r, nid(), ea, c1, "WebSocket [wss://]")
    ed(r, nid(), c1, c2, "Cria sessao")
    ed(r, nid(), c2, er, "Persiste sessao")
    ed(r, nid(), c1, c4, "Chunks de audio")
    ed(r, nid(), c4, c5, "Stream PCM")
    ed(r, nid(), c5, es, "Audio para Texto")
    ed(r, nid(), c5, c7, "Texto transcrito")
    ed(r, nid(), c7, c8, "CNPJ extraido")
    ed(r, nid(), c8, c9, "CNPJ valido/invalido")
    ed(r, nid(), c9, c1, "Resposta JSON")
    ed(r, nid(), c8, c10, "CNPJ invalido")
    ed(r, nid(), c10, c1, "Nova tentativa")
    ed(r, nid(), c6, ei, "Metricas e traces")
    ed(r, nid(), c3, ek, "Secrets")
    return m


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — Deployment View
# ═══════════════════════════════════════════════════════════════════════════════

def p4():
    reset(); m, r = mk(2000, 1200)
    bx(r, nid(), "Atento ASR Cloud - Deployment View", S_TITLE, 700, 15, 500, 35)

    ga = nid()
    gr(r, ga, 30, 70, 550, 700, "Datacenter Atento [On-Prem - Sao Paulo]", S_DEPLOY)

    gv = nid()
    gr(r, gv, 25, 50, 500, 300, "Avaya Cluster [AACC / Aura]",
       S_DEPLOY.replace("fillColor=#F5F5F5", "fillColor=#EBEBEB"), par=ga)
    bx(r, nid(), lbl("Session Border Controller", "Avaya SBC", "Gateway de midia"),
       S_EXT.replace("fillColor=#999999", "fillColor=#777777"), 20, 50, 220, 130, par=gv)
    bx(r, nid(), lbl("URA / IVR", "Avaya AAEP/VXML", "Arvore de atendimento"),
       S_EXT.replace("fillColor=#999999", "fillColor=#777777"), 260, 50, 220, 130, par=gv)

    gvp = nid()
    gr(r, gvp, 25, 400, 500, 160, "VPN Appliance [Firewall/Router]",
       S_DEPLOY.replace("fillColor=#F5F5F5", "fillColor=#EBEBEB"), par=ga)
    bx(r, nid(), lbl("VPN Client", "IPSec IKEv2", "Endpoint tunel S2S"),
       S_EXT.replace("fillColor=#999999", "fillColor=#777777"), 120, 50, 260, 80, par=gvp)

    gz = nid()
    gr(r, gz, 680, 70, 1270, 1060, "Azure Brazil South [Sao Paulo Region]", S_DEPLOY_AZ)

    gn = nid()
    gr(r, gn, 25, 55, 580, 230, "Networking Layer", S_DEPLOY_AZ.replace("fillColor=#E6F2FF", "fillColor=#F0F7FF"), par=gz)
    bx(r, nid(), lbl("VPN Gateway", "VpnGw2AZ", "Zone-Redundant"), S_CONTAINER, 20, 55, 260, 130, par=gn)
    bx(r, nid(), lbl("VNet Hub", "10.100.0.0/16", "NSGs + Route Tables"), S_CONTAINER, 300, 55, 260, 130, par=gn)

    gc = nid()
    gr(r, gc, 25, 320, 580, 310, "Compute Layer [AZ 1, 2, 3]", S_DEPLOY_AZ.replace("fillColor=#E6F2FF", "fillColor=#F0F7FF"), par=gz)
    bx(r, nid(), lbl("ASR Orchestrator #1", "App Service P2v3", "AZ-1"), S_CONTAINER, 20, 55, 250, 110, par=gc)
    bx(r, nid(), lbl("ASR Orchestrator #2", "App Service P2v3", "AZ-2"), S_CONTAINER, 290, 55, 250, 110, par=gc)
    bx(r, nid(), lbl("API Management", "Standard v2", "Gateway + Policies"), S_CONTAINER, 120, 195, 280, 90, par=gc)

    gd = nid()
    gr(r, gd, 650, 55, 580, 230, "Data Layer", S_DEPLOY_AZ.replace("fillColor=#E6F2FF", "fillColor=#F0F7FF"), par=gz)
    bx(r, nid(), lbl("Redis Cache", "Standard C1", "Zone-redundant"), S_CONTAINER, 20, 55, 260, 130, par=gd)
    bx(r, nid(), lbl("Key Vault", "Standard", "Soft-delete enabled"), S_CONTAINER, 300, 55, 260, 130, par=gd)

    gp = nid()
    gr(r, gp, 650, 320, 580, 190, "PaaS Services", S_DEPLOY_AZ.replace("fillColor=#E6F2FF", "fillColor=#F0F7FF"), par=gz)
    bx(r, nid(), lbl("Speech Services", "S0 Standard", "STT + Custom Model"), S_CONTAINER, 140, 55, 300, 100, par=gp)

    go = nid()
    gr(r, go, 650, 550, 580, 230, "Observability", S_DEPLOY_AZ.replace("fillColor=#E6F2FF", "fillColor=#F0F7FF"), par=gz)
    bx(r, nid(), lbl("Azure Monitor", "", "Metricas"), S_CONTAINER.replace("fillColor=#438DD5", "fillColor=#5B9BD5"), 15, 55, 170, 110, par=go)
    bx(r, nid(), lbl("App Insights", "", "APM + Tracing"), S_CONTAINER.replace("fillColor=#438DD5", "fillColor=#5B9BD5"), 200, 55, 170, 110, par=go)
    bx(r, nid(), lbl("Log Analytics", "", "Logs + KQL"), S_CONTAINER.replace("fillColor=#438DD5", "fillColor=#5B9BD5"), 385, 55, 170, 110, par=go)

    ar(r, nid(), "IPSec Tunnel - Encrypted",
       "html=1;endArrow=blockThin;startArrow=blockThin;fontSize=13;fontStyle=1;fontColor=#CC0000;strokeWidth=3;endFill=1;startFill=1;strokeColor=#CC0000;dashed=1;dashPattern=8 8;rounded=1;",
       580, 530, 680, 200)
    return m


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 5 — Dynamic View (Sequence)
# ═══════════════════════════════════════════════════════════════════════════════

def p5():
    reset(); m, r = mk(1900, 1300)
    bx(r, nid(), "Atento ASR Cloud - Dynamic View [Fluxo de Chamada]", S_TITLE, 550, 10, 700, 35)

    parts = [
        ("Cliente", "#08427B", 50), ("PSTN", "#999999", 210), ("Avaya SBC", "#666666", 370),
        ("URA/IVR", "#666666", 540), ("VPN On-Prem", "#777777", 720), ("VPN Azure", "#0078D4", 900),
        ("API Mgmt", "#438DD5", 1080), ("ASR Orch", "#2D6EB5", 1260),
        ("Speech Svc", "#999999", 1440), ("Redis", "#438DD5", 1600), ("Backend", "#777777", 1760),
    ]

    hdr = "html=1;rounded=1;whiteSpace=wrap;fontSize=11;fontStyle=1;fontColor=#ffffff;align=center;verticalAlign=middle;arcSize=8;strokeColor=#333333;"
    life = "html=1;strokeColor=#CCCCCC;strokeWidth=1;dashed=1;dashPattern=4 4;"

    for name, col, x in parts:
        bx(r, nid(), name, hdr + f"fillColor={col};", x, 60, 130, 45)
        ar(r, nid(), "", life, x+65, 105, x+65, 1250)

    msg = "html=1;endArrow=blockThin;fontSize=9;fontColor=#333333;strokeWidth=1;endFill=1;strokeColor=#404040;"

    steps = [
        (140, "1. Liga para central", 0, 1), (175, "2. SIP inbound", 1, 2),
        (210, "3. Roteia para URA", 2, 3),
        (260, "4. Toca: Diga seu CNPJ", 3, 3),
        (300, "5. Captura audio", 3, 2),
        (350, "6. Audio stream", 2, 4), (385, "7. IPSec tunnel", 4, 5),
        (420, "8. Audio na VNet", 5, 6), (455, "9. WebSocket session", 6, 7),
        (500, "10. Cria sessao (call-id)", 7, 9),
        (545, "11. Preprocessa G.711>PCM", 7, 7),
        (585, "12. Abre STT Real-Time", 7, 8), (625, "13. Stream chunks audio", 7, 8),
        (665, "14. Resultado parcial", 8, 7), (710, "15. Resultado FINAL", 8, 7),
        (755, "16. Parser + Validator", 7, 7),
        (800, "17. Cache CNPJ", 7, 9),
        (845, "18. JSON {cnpj, valid, 0.92}", 7, 6), (880, "19. Resposta", 6, 5),
        (915, "20. Retorno tunel", 5, 4), (950, "21. CNPJ identificado", 4, 3),
        (990, "22. Valida no cadastro", 3, 10), (1025, "23. Cliente encontrado", 10, 3),
        (1065, "24. Transfere chamada", 3, 2), (1100, "25. Identificamos sua empresa", 2, 0),
    ]

    note_s = "html=1;text;align=left;verticalAlign=middle;resizable=0;points=[];autosize=1;strokeColor=#FFC107;fillColor=#FFF8E1;fontSize=9;fontColor=#333333;rounded=1;strokeWidth=1;"

    for y, lab, si, ti in steps:
        eid = nid()
        sx = parts[si][2] + 65
        tx = parts[ti][2] + 65
        if si == ti:
            bx(r, eid, lab, note_s, sx+15, y-12, 260, 28)
            continue
        if ti < si:
            sx, tx = tx, sx
        ar(r, eid, lab, msg, sx, y, tx, y)

    bx(r, nid(), "", "html=1;rounded=1;fillColor=#E8F4FD;strokeColor=#0078D4;strokeWidth=2;dashed=1;opacity=40;", 710, 330, 960, 530)
    bx(r, nid(), "FASE ASR (via VPN)", "html=1;text;fontSize=13;fontStyle=1;fontColor=#0078D4;fillColor=none;strokeColor=none;", 720, 335, 200, 25)
    return m


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 6 — SOLUTION DESIGN EXECUTIVO (fluxo limpo com setas)
# ═══════════════════════════════════════════════════════════════════════════════

def p6():
    reset(); m, r = mk(1800, 1400)

    bx(r, nid(), "DESENHO DE SOLUCAO EXECUTIVO - Atento ASR Cloud",
       S_TITLE.replace("fontSize=20", "fontSize=22"), 400, 10, 1000, 40)
    bx(r, nid(), "Foursys SaaS | Azure Brazil South | VPN S2S com Atento",
       "html=1;text;align=center;verticalAlign=middle;resizable=0;points=[];strokeColor=none;fillColor=none;fontSize=13;fontColor=#666666;", 500, 50, 800, 25)

    # ── Cores dos boxes ──
    C_BLUE   = S_FLOW_BOX + "fillColor=#1565C0;strokeColor=#0D47A1;fontColor=#ffffff;"
    C_DBLUE  = S_FLOW_BOX + "fillColor=#0078D4;strokeColor=#005A9E;fontColor=#ffffff;"
    C_LBLUE  = S_FLOW_BOX + "fillColor=#438DD5;strokeColor=#3C7FC0;fontColor=#ffffff;"
    C_GRAY   = S_FLOW_BOX + "fillColor=#616161;strokeColor=#424242;fontColor=#ffffff;"
    C_BROWN  = S_FLOW_BOX + "fillColor=#5D4037;strokeColor=#3E2723;fontColor=#ffffff;"
    C_RED    = S_FLOW_BOX + "fillColor=#C62828;strokeColor=#8B2500;fontColor=#ffffff;"
    C_PURPLE = S_FLOW_BOX + "fillColor=#6A1B9A;strokeColor=#4A148C;fontColor=#ffffff;"
    C_ORANGE = S_FLOW_BOX + "fillColor=#E65100;strokeColor=#BF360C;fontColor=#ffffff;"
    C_TEAL   = S_FLOW_BOX + "fillColor=#00695C;strokeColor=#004D40;fontColor=#ffffff;"
    C_GREEN  = S_FLOW_BOX + "fillColor=#2E7D32;strokeColor=#1B5E20;fontColor=#ffffff;"
    C_NAVY   = S_FLOW_BOX + "fillColor=#08427B;strokeColor=#073B6F;fontColor=#ffffff;"

    A_BLUE = S_FLOW_ARROW + "strokeColor=#1565C0;"
    A_RED  = S_FLOW_ARROW + "strokeColor=#C62828;"
    A_PURP = S_FLOW_ARROW + "strokeColor=#6A1B9A;"
    A_GREEN = S_FLOW_ARROW + "strokeColor=#2E7D32;dashed=1;dashPattern=8 4;"
    A_GRAY = S_FLOW_ARROW + "strokeColor=#757575;"

    # ════════════════════════════════════════
    # FLUXO PRINCIPAL (de cima para baixo, esquerda para direita)
    # ════════════════════════════════════════

    # Row 1 — Cliente e Telefonia (y=100)
    b1 = nid()
    bx(r, b1, "Cliente Final<br/>Liga para a central", C_NAVY, 50, 100, 220, 80)

    b2 = nid()
    bx(r, b2, "Rede PSTN / SIP<br/>Telefonia publica", C_GRAY, 370, 100, 220, 80)

    # Row 2 — Avaya (y=260)
    b3 = nid()
    bx(r, b3, "Avaya SBC<br/>Session Border Controller", C_BROWN, 50, 260, 220, 80)

    b4 = nid()
    bx(r, b4, "URA / IVR<br/>'Diga seu CNPJ'", C_BROWN, 370, 260, 220, 80)

    # Row 3 — VPN (y=420)
    b5 = nid()
    bx(r, b5, "VPN On-Premises<br/>Firewall / Router", C_RED, 200, 420, 220, 80)

    b6 = nid()
    bx(r, b6, "TUNEL VPN S2S<br/>IPSec IKEv2 AES-256<br/>Criptografado", C_RED, 530, 420, 260, 80)

    # Row 4 — Azure Networking (y=580)
    b7 = nid()
    bx(r, b7, "Azure VPN Gateway<br/>VpnGw2AZ Zone-Redundant", C_DBLUE, 530, 580, 260, 80)

    b8 = nid()
    bx(r, b8, "Azure Virtual Network<br/>10.100.0.0/16 | NSG + PEs", C_DBLUE, 900, 580, 280, 80)

    # Row 5 — Azure Application (y=740)
    b9 = nid()
    bx(r, b9, "Azure API Management<br/>Auth + Rate Limit + Routing", C_LBLUE, 530, 740, 260, 80)

    b10 = nid()
    bx(r, b10, "ASR Orchestrator<br/>App Service P2v3 / AKS<br/>Autoscale min=2 max=20", C_BLUE, 900, 720, 300, 100)

    # Row 6 — Azure Intelligence (y=900)
    b11 = nid()
    bx(r, b11, "Azure Speech Services<br/>STT Real-Time pt-BR", C_PURPLE, 900, 900, 300, 80)

    b12 = nid()
    bx(r, b12, "Custom Speech Model<br/>CNPJ Alfanumerico<br/>Sotaques regionais", C_PURPLE, 900, 1020, 300, 80)

    # Lado direito — Dados e Observabilidade (coluna x=1350)
    b13 = nid()
    bx(r, b13, "Azure Redis Cache<br/>Sessoes + Cache CNPJ", C_ORANGE, 1350, 720, 260, 80)

    b14 = nid()
    bx(r, b14, "Azure Key Vault<br/>Secrets + Certificados", C_ORANGE, 1350, 840, 260, 80)

    b15 = nid()
    bx(r, b15, "Azure Monitor<br/>App Insights + Log Analytics<br/>Dashboards + Alertas", C_TEAL, 1350, 960, 260, 100)

    # Retorno — Atento
    b16 = nid()
    bx(r, b16, "Backend Atento<br/>CRM + Validacao CNPJ", C_GRAY, 50, 580, 220, 80)

    b17 = nid()
    bx(r, b17, "Operador Atento<br/>Recebe chamada identificada", C_NAVY, 50, 740, 220, 80)

    # ════════════════════════════════════════
    # SETAS DO FLUXO (numeradas)
    # ════════════════════════════════════════

    ed(r, nid(), b1, b2, "1 - Liga por voz", A_BLUE)
    ed(r, nid(), b2, b3, "2 - SIP Inbound", A_BLUE)
    ed(r, nid(), b3, b4, "3 - Roteia para URA", A_BLUE)
    ed(r, nid(), b4, b5, "4 - Captura audio", A_BLUE)
    ed(r, nid(), b5, b6, "5 - Envia via VPN", A_RED)
    ed(r, nid(), b6, b7, "6 - Tunel criptografado", A_RED)
    ed(r, nid(), b7, b8, "7 - Ingress na VNet", A_BLUE)
    ed(r, nid(), b8, b9, "8 - Roteia para APIM", A_BLUE)
    ed(r, nid(), b9, b10, "9 - WebSocket session", A_BLUE)
    ed(r, nid(), b10, b11, "10 - Audio para ASR", A_PURP)
    ed(r, nid(), b11, b12, "11 - Custom Model", A_PURP)

    # Return path
    ed(r, nid(), b11, b10, "12 - CNPJ transcrito", A_GREEN)
    ed(r, nid(), b10, b9, "13 - JSON resultado", A_GREEN)
    ed(r, nid(), b9, b7, "14 - Resposta", A_GREEN)
    ed(r, nid(), b7, b6, "15 - Retorno VPN", A_GREEN)
    ed(r, nid(), b6, b5, "16 - Decrypt", A_GREEN)
    ed(r, nid(), b5, b4, "17 - CNPJ valido", A_GREEN)
    ed(r, nid(), b4, b16, "18 - Valida cadastro", A_GRAY)
    ed(r, nid(), b4, b17, "19 - Transfere chamada", A_GREEN)

    # Orchestrator side connections
    ed(r, nid(), b10, b13, "Cache sessao", A_GRAY)
    ed(r, nid(), b10, b14, "Secrets", A_GRAY)
    ed(r, nid(), b10, b15, "Telemetria", A_GRAY)

    # ════════════════════════════════════════
    # BOXES INFORMATIVOS
    # ════════════════════════════════════════

    info_s = "html=1;rounded=1;whiteSpace=wrap;fillColor=#E8F5E9;strokeColor=#2E7D32;strokeWidth=2;fontColor=#1B5E20;align=center;verticalAlign=middle;fontSize=12;"

    bx(r, nid(), "Latencia E2E<br/>300 - 900ms", info_s, 50, 1100, 200, 70)
    bx(r, nid(), "SLA 99.95%<br/>Zone-Redundant 3 AZs", info_s, 300, 1100, 220, 70)
    bx(r, nid(), "Escala: 1 a 1000 chamadas/min<br/>Autoscale automatico", info_s, 570, 1100, 280, 70)
    bx(r, nid(), "LGPD: Audio descartado<br/>apos transcricao", info_s, 900, 1100, 260, 70)
    bx(r, nid(), "Fase Futura:<br/>WhatsApp + Chatbot", info_s, 1210, 1100, 220, 70)

    # ════════════════════════════════════════
    # LEGENDA
    # ════════════════════════════════════════

    bx(r, nid(), "LEGENDA", "html=1;text;fontSize=14;fontStyle=1;fontColor=#333333;fillColor=none;strokeColor=none;", 50, 1210, 120, 30)

    leg_box = "html=1;rounded=0;whiteSpace=wrap;fontSize=11;align=left;verticalAlign=middle;spacingLeft=8;"
    bx(r, nid(), "Fluxo de ida (audio)", leg_box + "fillColor=#E3F2FD;strokeColor=#1565C0;fontColor=#1565C0;strokeWidth=2;", 50, 1250, 200, 30)
    bx(r, nid(), "Tunel VPN criptografado", leg_box + "fillColor=#FFEBEE;strokeColor=#C62828;fontColor=#C62828;strokeWidth=2;", 270, 1250, 220, 30)
    bx(r, nid(), "Processamento ASR", leg_box + "fillColor=#F3E5F5;strokeColor=#6A1B9A;fontColor=#6A1B9A;strokeWidth=2;", 510, 1250, 200, 30)
    bx(r, nid(), "Retorno (resultado)", leg_box + "fillColor=#E8F5E9;strokeColor=#2E7D32;fontColor=#2E7D32;strokeWidth=2;dashed=1;", 730, 1250, 200, 30)
    bx(r, nid(), "Dados e Observabilidade", leg_box + "fillColor=#FFF3E0;strokeColor=#E65100;fontColor=#E65100;strokeWidth=2;", 950, 1250, 220, 30)

    return m


# ═══════════════════════════════════════════════════════════════════════════════
# BUILD
# ═══════════════════════════════════════════════════════════════════════════════

def build_drawio(path):
    mxf = ET.Element("mxfile", host="app.diagrams.net", modified="2026-04-08T00:00:00.000Z",
                     agent="Lyra-Architect", version="24.0.0", type="device")
    for i, (name, fn) in enumerate([
        ("C4 L1 - System Context", p1), ("C4 L2 - Container", p2),
        ("C4 L3 - Component", p3), ("Deployment View", p4),
        ("Dynamic View", p5), ("Solution Design Executivo", p6),
    ]):
        d = ET.SubElement(mxf, "diagram", id=f"p{i+1}", name=name)
        d.append(fn())

    raw = ET.tostring(mxf, encoding="unicode")
    pretty = minidom.parseString(raw).toprettyxml(indent="  ")
    lines = [l for l in pretty.split("\n") if l.strip()]
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"[OK] .drawio: {path}  ({os.path.getsize(path):,} bytes)")


def build_docx(path):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    for _ in range(6): doc.add_paragraph("")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Atento - ASR Cloud Solution"); run.bold = True; run.font.size = Pt(28); run.font.color.rgb = RGBColor(0x0B,0x48,0x84)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Modelo C4 & Desenho de Solucao Executivo"); run.font.size = Pt(18); run.font.color.rgb = RGBColor(0x43,0x8D,0xD5)
    doc.add_paragraph("")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Reconhecimento de Voz (ASR) para CNPJ Alfanumerico"); run.font.size = Pt(14); run.font.color.rgb = RGBColor(0x66,0x66,0x66)
    for _ in range(4): doc.add_paragraph("")

    meta = [("Arquiteta","Lyra - Multicloud Solution Architect"),("Cliente","Atento"),
            ("Cloud","Microsoft Azure (Brazil South)"),("Modelo","SaaS Gerenciado Foursys"),
            ("Data","Abril 2026"),("Versao","3.0"),("Classificacao","Confidencial")]
    t = doc.add_table(rows=len(meta), cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(k,v) in enumerate(meta):
        t.rows[i].cells[0].text = k; t.rows[i].cells[1].text = v
        for rn in t.rows[i].cells[0].paragraphs[0].runs: rn.bold = True
    doc.add_page_break()

    def tbl(hdrs, rows, sn="Light Grid Accent 1"):
        t = doc.add_table(rows=len(rows)+1, cols=len(hdrs)); t.style = sn
        for j,h in enumerate(hdrs): t.rows[0].cells[j].text = h
        for i,row in enumerate(rows):
            for j,v in enumerate(row): t.rows[i+1].cells[j].text = str(v)
        return t

    doc.add_heading("1. Sumario Executivo", level=1)
    doc.add_paragraph("A partir de Julho/2026 o CNPJ sera alfanumerico. As URAs Avaya da Atento so aceitam DTMF. A solucao adiciona ASR na Azure via VPN S2S, permitindo que o chamador diga o CNPJ. Entregue como SaaS gerenciado.")
    doc.add_heading("2. ADRs", level=1)
    tbl(["ID","Decisao","Justificativa"],[("ADR-001","Azure Brazil South","Menor latencia ~2-5ms"),("ADR-002","VPN S2S Zone-Redundant","HA com 3 AZs"),("ADR-003","Speech Services STT","PaaS; pt-BR nativo"),("ADR-004","Custom Speech Model","CNPJ alfanumerico e padrao novo"),("ADR-005","App Service > AKS","Simplicidade fase 1"),("ADR-006","Private Endpoints","Zero exposicao; LGPD"),("ADR-007","WebSocket","Latencia minima")])

    for title, tab, desc in [
        ("3. C4 L1 - System Context","C4 L1 - System Context","Visao de alto nivel do sistema e relacionamentos."),
        ("4. C4 L2 - Container","C4 L2 - Container","Decomposicao em containers: VPN, APIM, Orchestrator, Redis, etc."),
        ("5. C4 L3 - Component","C4 L3 - Component","Componentes internos do ASR Orchestrator."),
        ("6. Deployment View","Deployment View","Mapeamento para infra real: Atento DC + Azure Brazil South."),
        ("7. Dynamic View","Dynamic View","Sequencia de 25 passos de uma chamada."),
        ("8. Solution Design Executivo","Solution Design Executivo","Fluxo executivo com setas numeradas e componentes Azure."),
    ]:
        doc.add_page_break(); doc.add_heading(title, level=1); doc.add_paragraph(desc)
        doc.add_paragraph(f"Diagrama: atento-asr-c4-solution.drawio > aba '{tab}'", style="Intense Quote")

    doc.add_heading("8.1 Topologia de Rede", level=2)
    tbl(["Subnet","CIDR","Proposito"],[("GatewaySubnet","10.100.0.0/27","VPN Gateway"),("Subnet-APIM","10.100.1.0/24","API Management"),("Subnet-App","10.100.2.0/24","App Service"),("Subnet-PE","10.100.3.0/24","Private Endpoints")])
    doc.add_heading("8.2 Escalabilidade", level=2)
    doc.add_paragraph("Fase 1: App Service P2v3, 2 inst, autoscale CPU>70%, max 5", style="List Bullet")
    doc.add_paragraph("Fase 2: AKS com HPA, Cluster Autoscaler 3-20 nodes, KEDA", style="List Bullet")
    doc.add_heading("8.3 HA e DR", level=2)
    tbl(["Cenario","Tipo","Estrategia"],[("Falha Instancia","HA Auto","2+ inst em AZs; <30s; 99.95%"),("Falha Zona","HA Auto","Zone-Redundant"),("Falha Regional","DR Manual","Brazil Southeast; RTO 30-60min"),("Falha VPN","HA Auto","Active-Active; <15s; BGP")])
    doc.add_heading("8.4 Seguranca", level=2)
    tbl(["Camada","Controle","Detalhe"],[("Rede","VPN IPSec","IKEv2 AES-256-GCM"),("Rede","NSG","Whitelist Atento"),("Rede","Private Endpoints","Sem internet"),("Identidade","Managed Identity","Zero credenciais"),("Dados","Encryption","AES-256 rest; TLS 1.2+ transit"),("Dados","Audio","Descartado pos-transcricao (LGPD)"),("Compliance","LGPD","Brazil South; sem transferencia")])
    doc.add_heading("8.5 Roadmap", level=2)
    tbl(["Periodo","Fase","Entregaveis"],[("2026 H2","Fase 1 - ASR URA","VPN, Speech, Middleware, SaaS MVP"),("2027 H1","Fase 2 - Escala","AKS, Custom Model, Analytics"),("2027 H2","Fase 3 - Omnichannel","WhatsApp, Chatbot, NLU")])

    doc.add_page_break(); doc.add_heading("9. BoM - Estimativa Azure", level=1)
    t = tbl(["Recurso","SKU","Mensal BRL"],[("VPN Gateway","VpnGw2AZ","~R$ 2.800"),("App Service","P2v3 Linux","~R$ 1.400"),("API Management","Standard v2","~R$ 1.800"),("Speech Services","S0 STT","~R$ 500-2.000"),("Custom Speech","Hosting","~R$ 700"),("Redis Cache","Standard C1","~R$ 400"),("Key Vault","Standard","~R$ 20"),("Monitor + AI","Pay-per-use","~R$ 300-500"),("VNet+NSG+PE","-","~R$ 200"),("TOTAL","","~R$ 8.100-9.900/mes")])
    for rn in t.rows[-1].cells[0].paragraphs[0].runs: rn.bold = True
    for rn in t.rows[-1].cells[2].paragraphs[0].runs: rn.bold = True

    doc.add_page_break(); doc.add_heading("10. Riscos", level=1)
    tbl(["#","Risco","Mitigacao"],[("R1","Acuracia ASR insuficiente","Custom Model + heuristicas"),("R2","Latencia VPN alta","Monitoramento; ExpressRoute"),("R3","Integracao Avaya complexa","PoC antes do go-live"),("R4","Escala nao suave","Load testing + AKS"),("R5","LGPD auditoria","Audio transiente"),("R6","Custo acima","Alertas + Azure Advisor")])

    doc.add_heading("11. Proximos Passos", level=1)
    tbl(["#","Acao","Responsavel","Prazo"],[("1","Info tecnica Avaya","Vendedor+Atento","1 sem"),("2","Custom Speech Brazil South","Arquiteto","3 dias"),("3","PoC ASR CNPJ","Arquiteto+Dev","2 sem"),("4","Sizing Azure Calculator","Arquiteto","1 sem"),("5","Proposta comercial SaaS","Vendedor+Fin","1 sem"),("6","Kick-off","Atento","TBD")])

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Lyra - Arquiteta Multicloud, Foursys | v3.0 Abril 2026 | Confidencial")
    run.italic = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x99,0x99,0x99)
    doc.save(path)
    print(f"[OK] .docx: {path}  ({os.path.getsize(path):,} bytes)")


if __name__ == "__main__":
    d = os.path.dirname(os.path.abspath(__file__))
    print("Atento ASR Cloud - Gerador v3.0")
    print()
    build_drawio(os.path.join(d, "atento-asr-c4-solution.drawio"))
    build_docx(os.path.join(d, "atento-asr-solution-design-v3.docx"))
    print("\nPronto!")

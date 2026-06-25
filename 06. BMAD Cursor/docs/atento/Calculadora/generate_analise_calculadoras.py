from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.color.rgb = RGBColor(0, 51, 102)
    heading_style.font.name = 'Calibri'

HEADER_BG = RGBColor(0, 51, 102)
HEADER_TEXT = RGBColor(255, 255, 255)
ALT_ROW_BG = RGBColor(230, 240, 250)
HIGHLIGHT_BG = RGBColor(255, 242, 204)
ALERT_BG = RGBColor(255, 200, 200)
GREEN_BG = RGBColor(200, 240, 200)
ACCENT_COLOR = RGBColor(0, 102, 153)
RED_COLOR = RGBColor(180, 0, 0)
GREEN_COLOR = RGBColor(0, 120, 0)
GRAY_COLOR = RGBColor(100, 100, 100)

USD_BRL = 5.80


def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), str(color))
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.name = 'Calibri'
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size


def add_styled_table(doc, headers, rows, col_widths=None, highlight_last_row=False, alert_rows=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, header, bold=True, color=HEADER_TEXT, size=Pt(9))

    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            is_last = highlight_last_row and row_idx == len(rows) - 1
            is_alert = alert_rows and row_idx in alert_rows
            if is_last:
                set_cell_shading(cell, HIGHLIGHT_BG)
                set_cell_text(cell, value, bold=True, size=Pt(9))
            elif is_alert:
                set_cell_shading(cell, ALERT_BG)
                set_cell_text(cell, value, bold=True, size=Pt(9), color=RED_COLOR)
            else:
                if row_idx % 2 == 1:
                    set_cell_shading(cell, ALT_ROW_BG)
                set_cell_text(cell, value, size=Pt(9))

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)

    return table


def spacer(size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(size)
    p.paragraph_format.space_after = Pt(size)


def add_alert(text, color=RED_COLOR):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = color


def add_note(label, text, label_color=RED_COLOR):
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = label_color
    run = p.add_run(text)
    run.font.size = Pt(9)


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.runs[0] if p.runs else p.add_run()
        run.text = ''
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        r = p.add_run(text)
        r.font.size = Pt(10)
    else:
        for run in p.runs:
            run.font.size = Pt(10)


def fmt_usd(v):
    return f'${v:,.2f}'


def fmt_brl(v):
    return f'R$ {v:,.0f}'


# ============================================================
# CAPA
# ============================================================
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ANALISE DAS CALCULADORAS AZURE')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = HEADER_BG

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Projeto Atento \u2014 ASR Cloud Solution')
run.font.size = Pt(18)
run.font.color.rgb = ACCENT_COLOR

spacer(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Valores Reais da Azure Pricing Calculator vs Arquitetura Aprovada')
run.font.size = Pt(12)
run.font.color.rgb = GRAY_COLOR

spacer(30)

info_items = [
    ('Cliente', 'Atento'),
    ('Fornecedor', 'Foursys'),
    ('Cloud', 'Microsoft Azure (Brazil South)'),
    ('Moeda', 'USD (D\u00f3lar Americano) \u2014 Convers\u00e3o: $1 = R$ 5,80'),
    ('Fonte', '4 exporta\u00e7\u00f5es da Azure Pricing Calculator (14/04/2026)'),
    ('Data da An\u00e1lise', '14 de Abril de 2026'),
    ('Vers\u00e3o', 'v1.0'),
    ('Classifica\u00e7\u00e3o', 'Confidencial \u2014 Uso Interno Foursys + Atento'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)
    run = p.add_run(value)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# ============================================================
# 1. RESUMO EXECUTIVO
# ============================================================
doc.add_heading('1. Resumo Executivo', level=1)

p = doc.add_paragraph(
    'Foram analisadas 4 estimativas exportadas da Azure Pricing Calculator da Microsoft, '
    'cobrindo cen\u00e1rios de 10 a 5.000 liga\u00e7\u00f5es/minuto para o projeto ASR Cloud da Atento. '
    'Todos os valores est\u00e3o em D\u00f3lar Americano (USD), conforme a calculadora.'
)

spacer()

add_styled_table(doc,
    ['#', 'Cen\u00e1rio', 'Total USD/m\u00eas', 'Total BRL/m\u00eas', 'Status'],
    [
        ['1', 'MVP \u2014 10 chamadas/min', '$1,009.86', 'R$ 5,857', 'SKUs inadequados'],
        ['2', '100 chamadas/min', '$5,879.36', 'R$ 34,100', 'V\u00e1lido c/ ajustes'],
        ['3', '1.000 chamadas/min', '$13,398.64', 'R$ 77,712', 'Speech subestimado 5.5x'],
        ['4', '5.000 chamadas/min', '$29,893.92', 'R$ 173,385', 'Speech subestimado 11x'],
    ],
    col_widths=[1, 5, 3, 3, 4],
    alert_rows=[2, 3]
)

spacer()

add_alert('\u26a0 ACHADO CR\u00cdTICO: Os cen\u00e1rios de 1.000 e 5.000 chamadas/min possuem horas de Speech Services '
          'gravemente subestimadas, resultando em custos 3.7x a 7.7x menores que o real.')

spacer()
doc.add_page_break()

# ============================================================
# 2. DETALHAMENTO POR CENARIO
# ============================================================
doc.add_heading('2. Detalhamento por Cen\u00e1rio (Valores Reais em USD)', level=1)

# --- 2.1 ---
doc.add_heading('2.1 Cen\u00e1rio 10 chamadas/min (MVP)', level=2)
p = doc.add_paragraph('Total mensal: $1,009.86 USD | Upfront: $11.99 (SSL)')

add_styled_table(doc,
    ['Servi\u00e7o', 'SKU / Tier', 'USD/m\u00eas'],
    [
        ['VPN Gateway', 'VpnGw1AZ (Zone-Redundant)', '$156.95'],
        ['Speech Services STT', 'Standard S0 \u2014 440h \u00e1udio + 1 Custom endpoint', '$463.65'],
        ['App Service', 'Basic B2 (2 Cores, 3.5GB) x 2 inst\u00e2ncias', '$59.86'],
        ['API Management', 'Basic \u2014 1 unidade', '$147.17'],
        ['Redis Cache', 'Basic C1 (1 GB)', '$40.15'],
        ['Azure Monitor + App Insights', '~0.4 GB/dia logs', '$128.90'],
        ['Key Vault', 'Standard', '$13.18'],
        ['TOTAL', '', '$1,009.86'],
    ],
    col_widths=[5, 7, 3],
    highlight_last_row=True
)

spacer()
add_note('PROBLEMA', 'App Service Basic e API Management Basic N\u00c3O suportam VNet Integration. '
         'Impossibilita Private Endpoints e a conex\u00e3o segura via VPN definida na arquitetura.')
add_note('A\u00c7\u00c3O', 'Refazer com App Service P1v3, APIM Standard v2 e Redis Standard C1.', label_color=ACCENT_COLOR)
spacer()

# --- 2.2 ---
doc.add_heading('2.2 Cen\u00e1rio 100 chamadas/min', level=2)
p = doc.add_paragraph('Total mensal: $5,879.36 USD')

add_styled_table(doc,
    ['Servi\u00e7o', 'SKU / Tier', 'USD/m\u00eas'],
    [
        ['VPN Gateway', 'VpnGw2AZ (Zone-Redundant) \u2014 1.2 TB', '$458.45'],
        ['Speech Services STT', 'Standard S0 \u2014 4,000h \u00e1udio + 1 Custom endpoint', '$4,023.65'],
        ['AKS', '2x D4as v4 (4 vCPUs, 16GB) + 3 discos S15', '$528.32'],
        ['API Management', 'Standard \u2014 1 unidade', '$686.71'],
        ['Redis Cache', 'Basic C1 (1 GB)', '$40.15'],
        ['Azure Monitor + App Insights', '~0.4 GB/dia logs', '$128.90'],
        ['Key Vault', 'Standard', '$13.18'],
        ['TOTAL', '', '$5,879.36'],
    ],
    col_widths=[5, 7, 3],
    highlight_last_row=True
)

spacer()
add_note('Observa\u00e7\u00e3o', 'Speech hours levemente subestimado (4,000h vs 4,400h esperado). '
         'Redis Basic inadequado para produ\u00e7\u00e3o (sem SLA, sem replica\u00e7\u00e3o).', label_color=ACCENT_COLOR)
spacer()

# --- 2.3 ---
doc.add_heading('2.3 Cen\u00e1rio 1.000 chamadas/min', level=2)
p = doc.add_paragraph('Total mensal: $13,398.64 USD')

add_styled_table(doc,
    ['Servi\u00e7o', 'SKU / Tier', 'USD/m\u00eas'],
    [
        ['VPN Gateway', 'VpnGw3AZ (Zone-Redundant) \u2014 3 TB', '$1,153.65'],
        ['Speech Services STT', 'Standard S0 \u2014 8,000h \u00e1udio + 1 Custom endpoint', '$8,023.65'],
        ['AKS', '2 clusters, 4x D8as v4 (8 vCPUs, 32GB) + 4 discos E15', '$2,076.40'],
        ['API Management', 'Standard \u2014 2 unidades', '$1,373.42'],
        ['Redis Cache', 'Standard C2 x 2 inst\u00e2ncias', '$327.04'],
        ['Azure Monitor + App Insights', '~4 GB/dia logs', '$431.30'],
        ['Key Vault', 'Standard', '$13.18'],
        ['TOTAL', '', '$13,398.64'],
    ],
    col_widths=[5, 7, 3],
    highlight_last_row=True
)

spacer()
add_alert('\u26a0 Speech hours est\u00e3o 5.5x subestimadas! Deveria ser 44,000h, n\u00e3o 8,000h.')
spacer()

# --- 2.4 ---
doc.add_heading('2.4 Cen\u00e1rio 5.000 chamadas/min', level=2)
p = doc.add_paragraph('Total mensal: $29,893.92 USD')

add_styled_table(doc,
    ['Servi\u00e7o', 'SKU / Tier', 'USD/m\u00eas'],
    [
        ['VPN Gateway', 'VpnGw4AZ (Zone-Redundant) \u2014 3 TB', '$1,774.15'],
        ['Speech Services STT', 'Standard S0 \u2014 20,000h \u00e1udio + 1 Custom endpoint', '$20,023.65'],
        ['AKS', '3 clusters, 8x D8as v4 (8 vCPUs, 32GB) + 8 discos E15', '$4,079.80'],
        ['API Management', 'Standard \u2014 4 unidades', '$2,746.84'],
        ['Redis Cache', 'Standard C3 x 2 inst\u00e2ncias', '$657.00'],
        ['Azure Monitor + App Insights', '~6 GB/dia logs', '$599.30'],
        ['Key Vault', 'Standard', '$13.18'],
        ['TOTAL', '', '$29,893.92'],
    ],
    col_widths=[5, 7, 3],
    highlight_last_row=True
)

spacer()
add_alert('\u26a0 Speech hours est\u00e3o 11x subestimadas! Deveria ser 220,000h, n\u00e3o 20,000h.')
spacer()

doc.add_page_break()

# ============================================================
# 3. PROBLEMA CRITICO: SPEECH HOURS
# ============================================================
doc.add_heading('3. Problema Cr\u00edtico: Horas de Speech Services Subestimadas', level=1)

doc.add_heading('3.1 Premissas de C\u00e1lculo', level=2)

add_styled_table(doc,
    ['Par\u00e2metro', 'Valor'],
    [
        ['Janela de opera\u00e7\u00e3o', '12 horas/dia (07h-19h)'],
        ['Dias \u00fateis/m\u00eas', '22 dias'],
        ['Dura\u00e7\u00e3o m\u00e9dia de \u00e1udio por chamada', '10 segundos'],
        ['Minutos de opera\u00e7\u00e3o/m\u00eas', '15,840 min (12h \u00d7 60min \u00d7 22d)'],
        ['Pre\u00e7o STT Real-Time Standard', '~$1.00 por hora de \u00e1udio'],
        ['Custom Speech Endpoint', '~$23.65/m\u00eas por endpoint'],
    ],
    col_widths=[7, 9]
)

spacer()

doc.add_heading('3.2 Valida\u00e7\u00e3o das Horas de \u00c1udio', level=2)

add_styled_table(doc,
    ['Cen\u00e1rio', 'Chamadas/m\u00eas', 'Horas esperadas\n(@10s)', 'Horas na\nCalculadora', 'Diferen\u00e7a', 'Status'],
    [
        ['10/min', '158,400', '440h', '440h', '0%', '\u2705 OK'],
        ['100/min', '1,584,000', '4,400h', '4,000h', '-9%', '\u26a0 Aceit\u00e1vel'],
        ['1.000/min', '15,840,000', '44,000h', '8,000h', '-82%', '\u274c 5.5x menor!'],
        ['5.000/min', '79,200,000', '220,000h', '20,000h', '-91%', '\u274c 11x menor!'],
    ],
    col_widths=[2.5, 3, 3, 3, 2, 3],
    alert_rows=[2, 3]
)

spacer()

doc.add_heading('3.3 Impacto nos Custos Totais', level=2)

p = doc.add_paragraph('Com as horas de Speech corrigidas (premissa de 10 segundos/chamada):')
spacer()

add_styled_table(doc,
    ['Cen\u00e1rio', 'Total na\nCalculadora', 'Speech\nCorrigido', 'Total\nCorrigido (USD)', 'Total\nCorrigido (BRL)', 'Diferen\u00e7a'],
    [
        ['10/min', '$1,009.86', '$463.65', '$1,009.86', 'R$ 5,857', '\u2014'],
        ['100/min', '$5,879.36', '$4,423.65', '$6,279.36', 'R$ 36,420', '+$400'],
        ['1.000/min', '$13,398.64', '$44,023.65', '$49,398.64', 'R$ 286,512', '+$36,000'],
        ['5.000/min', '$29,893.92', '$220,023.65', '$229,893.92', 'R$ 1,333,385', '+$200,000'],
    ],
    col_widths=[2.5, 2.5, 2.5, 3, 3, 2.5],
    alert_rows=[2, 3]
)

spacer()
doc.add_page_break()

# ============================================================
# 4. PROBLEMAS DE SKU
# ============================================================
doc.add_heading('4. Problemas de Compatibilidade com a Arquitetura', level=1)

doc.add_heading('4.1 Cen\u00e1rio 10/min \u2014 SKUs Inadequados', level=2)

p = doc.add_paragraph(
    'Os SKUs escolhidos na calculadora para o cen\u00e1rio MVP n\u00e3o atendem os requisitos '
    'de seguran\u00e7a definidos na arquitetura (Private Endpoints, VNet isolation):')

spacer()

add_styled_table(doc,
    ['Componente', 'Calculadora', 'Arquitetura', 'Problema'],
    [
        ['App Service', 'Basic B2', 'P1v3 / P2v3', 'Basic N\u00c3O suporta VNet Integration'],
        ['API Management', 'Basic', 'Standard v2', 'Basic N\u00c3O suporta VNet Integration'],
        ['Redis Cache', 'Basic C1', 'Standard C1', 'Basic sem SLA e sem replica\u00e7\u00e3o'],
    ],
    col_widths=[3, 3, 3, 7]
)

spacer()
add_alert('RISCO: Com SKUs Basic, o t\u00fanel VPN fica inutilizado \u2014 os servi\u00e7os ficariam '
          'expostos publicamente, violando os requisitos de seguran\u00e7a e LGPD.')
spacer()

doc.add_heading('4.2 Redis Basic nos Cen\u00e1rios 10/min e 100/min', level=2)

p = doc.add_paragraph(
    'O Redis Cache na tier Basic n\u00e3o oferece SLA, replica\u00e7\u00e3o nem zone-redundancy. '
    'Para um ambiente de produ\u00e7\u00e3o com chamadas em tempo real, o m\u00ednimo recomendado \u00e9 Standard.')

spacer()

doc.add_heading('4.3 Resumo de Conformidade por Cen\u00e1rio', level=2)

add_styled_table(doc,
    ['Cen\u00e1rio', 'Speech Hours', 'Compute', 'APIM', 'Redis', 'VPN', 'Veredito'],
    [
        ['10/min', '\u2705 OK', '\u274c Basic', '\u274c Basic', '\u274c Basic', '\u2705 OK', 'Refazer SKUs'],
        ['100/min', '\u26a0 -9%', '\u2705 AKS', '\u2705 Standard', '\u274c Basic', '\u2705 OK', 'Ajustar Redis + Speech'],
        ['1.000/min', '\u274c -82%', '\u2705 AKS', '\u2705 Standard', '\u2705 Standard', '\u2705 OK', 'Refazer Speech'],
        ['5.000/min', '\u274c -91%', '\u2705 AKS', '\u2705 Standard', '\u2705 Standard', '\u2705 OK', 'Refazer Speech'],
    ],
    col_widths=[2.5, 2.5, 2, 2.5, 2, 2, 3]
)

spacer()
doc.add_page_break()

# ============================================================
# 5. TABELA COMPARATIVA CONSOLIDADA
# ============================================================
doc.add_heading('5. Tabela Comparativa Consolidada', level=1)

doc.add_heading('5.1 Valores Originais da Calculadora (USD)', level=2)

add_styled_table(doc,
    ['Componente', '10/min', '100/min', '1.000/min', '5.000/min'],
    [
        ['VPN Gateway', '$156.95', '$458.45', '$1,153.65', '$1,774.15'],
        ['Speech Services STT', '$463.65', '$4,023.65', '$8,023.65', '$20,023.65'],
        ['Compute (App Svc / AKS)', '$59.86', '$528.32', '$2,076.40', '$4,079.80'],
        ['API Management', '$147.17', '$686.71', '$1,373.42', '$2,746.84'],
        ['Redis Cache', '$40.15', '$40.15', '$327.04', '$657.00'],
        ['Monitor + Logs', '$128.90', '$128.90', '$431.30', '$599.30'],
        ['Key Vault', '$13.18', '$13.18', '$13.18', '$13.18'],
        ['TOTAL (USD)', '$1,009.86', '$5,879.36', '$13,398.64', '$29,893.92'],
        ['TOTAL (BRL @5.80)', 'R$ 5,857', 'R$ 34,100', 'R$ 77,712', 'R$ 173,385'],
    ],
    col_widths=[5, 2.5, 2.5, 3, 3],
    highlight_last_row=True
)

spacer()

doc.add_heading('5.2 Valores Corrigidos (Speech Hours Corretas)', level=2)

add_styled_table(doc,
    ['Componente', '10/min', '100/min', '1.000/min', '5.000/min'],
    [
        ['VPN Gateway', '$156.95', '$458.45', '$1,153.65', '$1,774.15'],
        ['Speech Services STT', '$463.65', '$4,423.65', '$44,023.65', '$220,023.65'],
        ['Compute (App Svc / AKS)', '$59.86', '$528.32', '$2,076.40', '$4,079.80'],
        ['API Management', '$147.17', '$686.71', '$1,373.42', '$2,746.84'],
        ['Redis Cache', '$40.15', '$40.15', '$327.04', '$657.00'],
        ['Monitor + Logs', '$128.90', '$128.90', '$431.30', '$599.30'],
        ['Key Vault', '$13.18', '$13.18', '$13.18', '$13.18'],
        ['TOTAL CORRIGIDO (USD)', '$1,009.86', '$6,279.36', '$49,398.64', '$229,893.92'],
        ['TOTAL CORRIGIDO (BRL)', 'R$ 5,857', 'R$ 36,420', 'R$ 286,512', 'R$ 1,333,385'],
    ],
    col_widths=[5, 2.5, 2.5, 3, 3],
    highlight_last_row=True
)

spacer()

doc.add_heading('5.3 Calculadora vs Estimativa Anterior do Documento de Solu\u00e7\u00e3o', level=2)

add_styled_table(doc,
    ['Cen\u00e1rio', 'Estimativa Anterior\n(Doc. Solu\u00e7\u00e3o - BRL)', 'Calculadora Real\n(USD\u2192BRL)', 'Calculadora\nCorrigida (BRL)', 'Desvio vs\nAnterior'],
    [
        ['10/min', 'R$ 8,370', 'R$ 5,857', 'R$ 5,857*', '-30% (SKUs Basic)'],
        ['100/min', 'R$ 33,150', 'R$ 34,100', 'R$ 36,420', '+10% (alinhado)'],
        ['1.000/min', 'R$ 267,050', 'R$ 77,712', 'R$ 286,512', '+7% (alinhado)'],
        ['5.000/min', '\u2014 (n\u00e3o estimado)', 'R$ 173,385', 'R$ 1,333,385', 'Novo cen\u00e1rio'],
    ],
    col_widths=[2.5, 3.5, 3, 3, 3.5]
)

spacer()
add_note('* Nota', 'O cen\u00e1rio 10/min \u00e9 mais barato porque usa SKUs Basic que n\u00e3o atendem '
         'os requisitos. Com SKUs adequados (P1v3, Standard APIM, Standard Redis), '
         'o custo sobe para ~$1,600-1,800/m\u00eas (R$ 9,300-10,440), alinhado \u00e0 estimativa anterior.',
         label_color=ACCENT_COLOR)

spacer()
doc.add_page_break()

# ============================================================
# 6. PROPORCAO DE CUSTO
# ============================================================
doc.add_heading('6. Propor\u00e7\u00e3o de Custo por Componente (Valores Corrigidos)', level=1)

p = doc.add_paragraph(
    'O Azure Speech Services \u00e9 o componente dominante em todos os cen\u00e1rios:')
spacer()

add_styled_table(doc,
    ['Cen\u00e1rio', 'Total Mensal\n(USD)', '% Speech\nServices', '% Compute', '% Rede\n(VPN+BW)', '% Outros'],
    [
        ['10/min', '$1,010', '46%', '6%', '16%', '32%'],
        ['100/min', '$6,279', '70%', '8%', '7%', '14%'],
        ['1.000/min', '$49,399', '89%', '4%', '2%', '4%'],
        ['5.000/min', '$229,894', '96%', '2%', '1%', '1%'],
    ],
    col_widths=[2.5, 3, 2.5, 2.5, 2.5, 2.5]
)

spacer()

add_note('Conclus\u00e3o', 'A partir de 100 chamadas/min, Speech Services representa 70-96% do custo total. '
         'Qualquer otimiza\u00e7\u00e3o de custo deve focar prioritariamente neste componente '
         '(Enterprise Agreement, Reserved Capacity, ou alternativa self-hosted).',
         label_color=ACCENT_COLOR)

spacer()

# ============================================================
# 7. TCO CORRIGIDO
# ============================================================
doc.add_heading('7. TCO Corrigido \u2014 12 Meses', level=1)

p = doc.add_paragraph('Proje\u00e7\u00e3o anual com valores corrigidos (apenas infraestrutura Azure):')
spacer()

add_styled_table(doc,
    ['M\u00e9trica', '10/min', '100/min', '1.000/min', '5.000/min'],
    [
        ['Custo mensal (USD)', '$1,010', '$6,279', '$49,399', '$229,894'],
        ['Custo mensal (BRL)', 'R$ 5,857', 'R$ 36,420', 'R$ 286,512', 'R$ 1,333,385'],
        ['TCO 12 meses (USD)', '$12,118', '$75,352', '$592,784', '$2,758,727'],
        ['TCO 12 meses (BRL)', 'R$ 70,286', 'R$ 437,044', 'R$ 3,438,147', 'R$ 16,000,617'],
        ['Custo por chamada', '$0.0064', '$0.0040', '$0.0031', '$0.0029'],
    ],
    col_widths=[4, 3, 3, 3, 3],
    highlight_last_row=True
)

spacer()

add_note('Observa\u00e7\u00e3o', 'Custo por chamada diminui com escala (custos fixos se diluem), '
         'por\u00e9m Speech Services \u00e9 pay-per-use linear, limitando a economia.',
         label_color=ACCENT_COLOR)

spacer()
doc.add_page_break()

# ============================================================
# 8. VEREDITO
# ============================================================
doc.add_heading('8. Veredito Final: As Calculadoras Fazem Sentido?', level=1)

scenarios_verdict = [
    ('10/min (MVP)', 'PARCIALMENTE V\u00c1LIDO',
     'Valores corretos para os SKUs selecionados, por\u00e9m os SKUs Basic s\u00e3o '
     'inadequados para a arquitetura (n\u00e3o suportam VNet/Private Endpoints). '
     'A\u00e7\u00e3o: refazer com P1v3, Standard APIM, Standard Redis.',
     HIGHLIGHT_BG),
    ('100/min', 'V\u00c1LIDO COM AJUSTES',
     'Valores consistentes. Speech hours levemente subestimado (-9%, impacto ~$400). '
     'Redis deveria ser Standard. Mais pr\u00f3ximo da estimativa anterior.',
     GREEN_BG),
    ('1.000/min', 'INV\u00c1LIDO',
     'Speech hours 5.5x abaixo do necess\u00e1rio. Custo real: ~$49,400/m\u00eas '
     '(n\u00e3o $13,400). A\u00e7\u00e3o: refazer na calculadora com 44,000h de Speech.',
     ALERT_BG),
    ('5.000/min', 'INV\u00c1LIDO',
     'Speech hours 11x abaixo do necess\u00e1rio. Custo real: ~$229,900/m\u00eas '
     '(n\u00e3o $29,900). A\u00e7\u00e3o: refazer na calculadora com 220,000h de Speech.',
     ALERT_BG),
]

for title, status, desc, bg in scenarios_verdict:
    doc.add_heading(f'Cen\u00e1rio {title} \u2014 {status}', level=2)
    p = doc.add_paragraph(desc)
    spacer(4)

spacer()

# ============================================================
# 9. RECOMENDACOES
# ============================================================
doc.add_heading('9. Recomenda\u00e7\u00f5es', level=1)

recommendations = [
    ('1. Corrigir urgente cen\u00e1rios 1.000 e 5.000/min',
     'Refazer na Azure Pricing Calculator com as horas de Speech corretas: '
     '44,000h para 1.000/min e 220,000h para 5.000/min.'),
    ('2. Refazer cen\u00e1rio MVP (10/min)',
     'Substituir SKUs Basic por vers\u00f5es que suportem VNet Integration: '
     'App Service P1v3, API Management Standard v2, Redis Standard C1.'),
    ('3. Negociar Enterprise Agreement',
     'Para cen\u00e1rios acima de 100/min, Speech Services = 70-96% do custo. '
     'Descontos de 30-50% via EA mudam a viabilidade financeira.'),
    ('4. Avaliar alternativa self-hosted',
     'Para 1.000+ chamadas/min, Whisper/Faster-Whisper no AKS com GPUs '
     'pode ser drasticamente mais barato que Speech Services pay-per-use.'),
    ('5. Confirmar volume real com a Atento',
     'A vari\u00e1vel mais cr\u00edtica \u00e9 o volume real de chamadas/minuto. '
     'Solicitar dados hist\u00f3ricos da opera\u00e7\u00e3o. Um MVP de 10-100/min '
     '\u00e9 vi\u00e1vel financeiramente ($1K-$6K/m\u00eas).'),
    ('6. Considerar Reserved Capacity',
     'Para Compute (AKS) nos cen\u00e1rios B e C, Reserved VM Instances '
     '(1 ou 3 anos) geram economia de 30-60% nos nodes base.'),
]

for title, desc in recommendations:
    p = doc.add_paragraph()
    run = p.add_run(title + ': ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT_COLOR
    run = p.add_run(desc)
    run.font.size = Pt(10)

spacer()
doc.add_page_break()

# ============================================================
# 10. APENDICE - DADOS BRUTOS
# ============================================================
doc.add_heading('10. Ap\u00eandice \u2014 Dados Brutos dos Arquivos', level=1)

doc.add_heading('10.1 Arquivos Analisados', level=2)

add_styled_table(doc,
    ['Arquivo', 'Data de Cria\u00e7\u00e3o', 'Tamanho'],
    [
        ['[Atento] Cen\u00e1rio- MVP (10 Liga\u00e7\u00f5es min).xlsx', '14/04/2026 00:54 UTC', '~80 KB'],
        ['[Atento] Cen\u00e1rio- MVP (100 Liga\u00e7\u00f5es min).xlsx', '14/04/2026 00:57 UTC', '~80 KB'],
        ['[Atento] Cen\u00e1rio- MVP (1000 Liga\u00e7\u00f5es min).xlsx', '14/04/2026 00:56 UTC', '~80 KB'],
        ['[Atento] Cen\u00e1rio- MVP (5000 Liga\u00e7\u00f5es min).xlsx', '14/04/2026 00:58 UTC', '~55 KB'],
    ],
    col_widths=[8, 4, 3]
)

spacer()

doc.add_heading('10.2 Descri\u00e7\u00f5es Completas dos Servi\u00e7os (Calculadora)', level=2)

services_detail = [
    ('VPN Gateway (10/min)', 'VpnGw1 tier, 730 gateway hours, 1 S2S tunnel, 100 GB data transfer'),
    ('VPN Gateway (100/min)', 'VpnGw2 tier, 730 gateway hours, 1 S2S tunnel, 1,200 GB data transfer'),
    ('VPN Gateway (1.000/min)', 'VpnGw3 tier, 730 gateway hours, 1 S2S tunnel, 3 TB data transfer'),
    ('VPN Gateway (5.000/min)', 'VpnGw4 tier, 730 gateway hours, 1 S2S tunnel, 3 TB data transfer'),
    ('Speech Services', 'Azure Speech in Foundry Tools, Pay as you go, Standard, Speech to Text, '
     'Standard real-time transcription, 1 Custom endpoint hosting model'),
    ('App Service (10/min)', 'Basic B2 (2 Cores, 3.5 GB RAM, 10 GB Storage) x 730h, Linux'),
    ('AKS (100/min)', 'Standard, 1 cluster, 2x D4as v4 (4 vCPUs, 16 GB), 3 managed OS disks S15'),
    ('AKS (1.000/min)', 'Standard, 2 clusters, 4x D8as v4 (8 vCPUs, 32 GB), 4 managed OS disks E15'),
    ('AKS (5.000/min)', 'Standard, 3 clusters, 8x D8as v4 (8 vCPUs, 32 GB), 8 managed OS disks E15'),
    ('API Management (10/min)', 'Basic tier, 1 unit, 730 hours'),
    ('API Management (100/min)', 'Standard tier, 1 unit, 730 hours'),
    ('API Management (1.000/min)', 'Standard tier, 2 units, 730 hours'),
    ('API Management (5.000/min)', 'Standard tier, 4 units, 730 hours'),
]

for title, desc in services_detail:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(desc)
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY_COLOR

spacer()

doc.add_heading('10.3 Pre\u00e7os de Refer\u00eancia Azure (Brazil South, Abril 2026)', level=2)

add_styled_table(doc,
    ['Servi\u00e7o', 'M\u00e9trica', 'Pre\u00e7o USD'],
    [
        ['Speech STT Real-Time', 'Por hora de \u00e1udio', '~$1.00/h'],
        ['Custom Speech Endpoint', 'Por endpoint/m\u00eas', '~$23.65'],
        ['VPN Gateway VpnGw1AZ', 'Por hora (730h/m\u00eas)', '~$0.215/h'],
        ['VPN Gateway VpnGw2AZ', 'Por hora (730h/m\u00eas)', '~$0.528/h'],
        ['App Service Basic B2', 'Por inst\u00e2ncia/m\u00eas', '~$29.93'],
        ['AKS D4as v4', 'Por node/m\u00eas', '~$220'],
        ['AKS D8as v4', 'Por node/m\u00eas', '~$440'],
        ['Redis Basic C1', 'Por inst\u00e2ncia/m\u00eas', '~$40.15'],
    ],
    col_widths=[5, 5, 4]
)

spacer()

# ============================================================
# RODAPE
# ============================================================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\u2500' * 60)
run.font.color.rgb = RGBColor(180, 180, 180)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documento produzido pela Squad MEQ \u2014 Foursys')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(120, 120, 120)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Revis\u00e3o: v1.0 \u2014 14 Abril 2026 | Classifica\u00e7\u00e3o: Confidencial')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(120, 120, 120)

# ============================================================
# SALVAR
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Analise Calculadoras Azure - Atento.docx')
doc.save(output_path)
print(f'Documento salvo em: {output_path}')

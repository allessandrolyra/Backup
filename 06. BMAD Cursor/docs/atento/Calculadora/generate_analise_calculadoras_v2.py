from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.color.rgb = RGBColor(0, 51, 102)
    heading_style.font.name = 'Calibri'

HEADER_BG = RGBColor(0, 51, 102)
HEADER_TEXT = RGBColor(255, 255, 255)
ALT_ROW_BG = RGBColor(230, 240, 250)
HIGHLIGHT_BG = RGBColor(255, 242, 204)
ALERT_BG = RGBColor(255, 200, 200)
GREEN_BG = RGBColor(200, 235, 200)
ACCENT_COLOR = RGBColor(0, 102, 153)
RED_COLOR = RGBColor(180, 0, 0)
GREEN_COLOR = RGBColor(0, 100, 0)
GRAY_COLOR = RGBColor(100, 100, 100)

USD_BRL = 5.80


def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), str(color))
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.name = 'Calibri'
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size


def add_styled_table(doc, headers, rows, col_widths=None, highlight_last_row=False,
                     alert_rows=None, green_rows=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_BG)
        set_cell_text(cell, header, bold=True, color=HEADER_TEXT, size=Pt(9))

    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            is_last = highlight_last_row and row_idx == len(rows) - 1
            is_alert = alert_rows and row_idx in alert_rows
            is_green = green_rows and row_idx in green_rows
            if is_last:
                set_cell_shading(cell, HIGHLIGHT_BG)
                set_cell_text(cell, value, bold=True, size=Pt(9))
            elif is_alert:
                set_cell_shading(cell, ALERT_BG)
                set_cell_text(cell, value, bold=True, size=Pt(9), color=RED_COLOR)
            elif is_green:
                set_cell_shading(cell, GREEN_BG)
                set_cell_text(cell, value, bold=True, size=Pt(9), color=GREEN_COLOR)
            else:
                if row_idx % 2 == 1:
                    set_cell_shading(cell, ALT_ROW_BG)
                set_cell_text(cell, value, size=Pt(9))

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)
    return table


def spacer(size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(size)
    p.paragraph_format.space_after = Pt(size)


def add_alert(text, color=RED_COLOR):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = color


def add_note(label, text, label_color=RED_COLOR):
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = label_color
    run = p.add_run(text)
    run.font.size = Pt(9)


# ============================================================
# CAPA
# ============================================================
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AN\u00c1LISE DAS CALCULADORAS AZURE')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = HEADER_BG

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Projeto Atento \u2014 ASR Cloud Solution')
run.font.size = Pt(18)
run.font.color.rgb = ACCENT_COLOR

spacer(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Valores Reais da Azure Pricing Calculator (Vers\u00e3o Corrigida)')
run.font.size = Pt(12)
run.font.color.rgb = GRAY_COLOR

spacer(30)

info_items = [
    ('Cliente', 'Atento'),
    ('Fornecedor', 'Foursys'),
    ('Cloud', 'Microsoft Azure (Brazil South)'),
    ('Moeda', 'USD (D\u00f3lar Americano) \u2014 Convers\u00e3o: $1 = R$ 5,80'),
    ('Fonte', '4 exporta\u00e7\u00f5es da Azure Pricing Calculator (14/04/2026 \u2014 v2 corrigida)'),
    ('Data', '14 de Abril de 2026'),
    ('Vers\u00e3o', 'v2.0 \u2014 Calculadoras revisadas com corre\u00e7\u00f5es'),
    ('Classifica\u00e7\u00e3o', 'Confidencial \u2014 Uso Interno Foursys + Atento'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)
    run = p.add_run(value)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# ============================================================
# 1. RESUMO EXECUTIVO
# ============================================================
doc.add_heading('1. Resumo Executivo', level=1)

p = doc.add_paragraph(
    'Foram analisadas as 4 estimativas revisadas exportadas da Azure Pricing Calculator, '
    'cobrindo cen\u00e1rios de 10 a 5.000 liga\u00e7\u00f5es/minuto para o projeto ASR Cloud da Atento. '
    'Esta vers\u00e3o incorpora as corre\u00e7\u00f5es identificadas na an\u00e1lise anterior: '
    'horas de Speech corrigidas, SKUs adequados para VNet Integration e Reserved Instances para AKS.'
)
spacer()

add_styled_table(doc,
    ['#', 'Cen\u00e1rio', 'Total USD/m\u00eas', 'Total BRL/m\u00eas', 'Status'],
    [
        ['1', 'MVP \u2014 10 chamadas/min', '$1,712.19', 'R$ 9,931', '\u2705 V\u00e1lido'],
        ['2', '100 chamadas/min', '$6,061.19', 'R$ 35,155', '\u2705 V\u00e1lido'],
        ['3', '1.000 chamadas/min', '$48,283.49', 'R$ 280,044', '\u2705 V\u00e1lido'],
        ['4', '5.000 chamadas/min', '$227,663.62', 'R$ 1,320,449', '\u2705 V\u00e1lido'],
    ],
    col_widths=[1, 5, 3, 3, 3],
    green_rows=[0, 1, 2, 3]
)

spacer()

add_alert('\u2705 TODAS AS CALCULADORAS EST\u00c3O CONSISTENTES COM A ARQUITETURA DO PROJETO ATENTO.',
          color=GREEN_COLOR)

spacer()

p = doc.add_paragraph('Corre\u00e7\u00f5es aplicadas nesta vers\u00e3o (vs. vers\u00e3o anterior):')
corrections = [
    ('Speech Hours corrigidas', '1.000/min: 8,000h \u2192 44,000h | 5.000/min: 20,000h \u2192 220,000h | 100/min: 4,000h \u2192 4,400h'),
    ('SKUs MVP atualizados', 'App Service: Basic B2 \u2192 Premium V3 P1V3 | APIM: Basic \u2192 Standard | Redis: Basic \u2192 Standard'),
    ('Reserved Instances (AKS)', 'Todos os cen\u00e1rios com AKS agora usam Reserved 3 anos (economia ~55%)'),
]
for title, desc in corrections:
    p = doc.add_paragraph()
    run = p.add_run(f'\u2705 {title}: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = GREEN_COLOR
    run = p.add_run(desc)
    run.font.size = Pt(9)

spacer()
doc.add_page_break()

# ============================================================
# 2. DETALHAMENTO POR CENARIO
# ============================================================
doc.add_heading('2. Detalhamento por Cen\u00e1rio (Valores Reais em USD)', level=1)

# --- 2.1 ---
doc.add_heading('2.1 Cen\u00e1rio 10 chamadas/min (MVP)', level=2)

p = doc.add_paragraph()
run = p.add_run('Total mensal: $1,712.19 USD (R$ 9,931)')
run.bold = True
run = p.add_run(' | Upfront: $11.99 (SSL)')
run.font.size = Pt(9)
run.font.color.rgb = GRAY_COLOR

spacer()
add_styled_table(doc,
    ['Servi\u00e7o', 'SKU / Tier', 'USD/m\u00eas', '% Total'],
    [
        ['VPN Gateway', 'VpnGw1AZ (Zone-Redundant), 100 GB transfer', '$156.95', '9%'],
        ['Speech Services STT', 'Standard S0 \u2014 440h \u00e1udio + 1 Custom endpoint', '$463.65', '27%'],
        ['App Service', 'Premium V3 P1V3 (2 vCPU, 8GB RAM) x 1, Linux', '$162.06', '9%'],
        ['API Management', 'Standard \u2014 1 unidade', '$686.71', '40%'],
        ['Redis Cache', 'Standard C1 (1 GB)', '$100.74', '6%'],
        ['Azure Monitor + App Insights', '~0.4 GB/dia logs, 3 meses reten\u00e7\u00e3o', '$128.90', '8%'],
        ['Key Vault', 'Standard', '$13.18', '1%'],
        ['TOTAL', '', '$1,712.19', '100%'],
    ],
    col_widths=[4.5, 6, 2.5, 1.5],
    highlight_last_row=True
)

spacer()
add_note('Valida\u00e7\u00e3o',
         'Speech: 440h = 10 chamadas/min \u00d7 10s \u00d7 12h/dia \u00d7 22 dias \u2705 | '
         'App Service P1V3 suporta VNet Integration \u2705 | '
         'APIM Standard suporta VNet \u2705 | '
         'Redis Standard tem SLA 99.9% \u2705',
         label_color=GREEN_COLOR)

add_note('Observa\u00e7\u00e3o',
         'O APIM Standard representa 40% do custo neste cen\u00e1rio. '
         'Para o MVP, considerar se \u00e9 poss\u00edvel usar APIM Basic v2 (nova tier que suporta VNet) '
         'ou iniciar sem APIM usando autentica\u00e7\u00e3o direta no App Service.',
         label_color=ACCENT_COLOR)
spacer()

# --- 2.2 ---
doc.add_heading('2.2 Cen\u00e1rio 100 chamadas/min', level=2)

p = doc.add_paragraph()
run = p.add_run('Total mensal: $6,061.19 USD (R$ 35,155)')
run.bold = True

spacer()
add_styled_table(doc,
    ['Servi\u00e7o', 'SKU / Tier', 'USD/m\u00eas', '% Total'],
    [
        ['VPN Gateway', 'VpnGw2AZ (Zone-Redundant), 1.2 TB transfer', '$458.45', '8%'],
        ['Speech Services STT', 'Standard S0 \u2014 4,400h \u00e1udio + 1 Custom endpoint', '$4,423.65', '73%'],
        ['AKS', '1 cluster, 2x D4as v4 (4 vCPU, 16GB), Reserved 3 anos', '$249.56', '4%'],
        ['API Management', 'Standard \u2014 1 unidade', '$686.71', '11%'],
        ['Redis Cache', 'Standard C1 (1 GB)', '$100.74', '2%'],
        ['Azure Monitor + App Insights', '~0.4 GB/dia logs, 3 meses reten\u00e7\u00e3o', '$128.90', '2%'],
        ['Key Vault', 'Standard', '$13.18', '<1%'],
        ['TOTAL', '', '$6,061.19', '100%'],
    ],
    col_widths=[4.5, 6, 2.5, 1.5],
    highlight_last_row=True
)

spacer()
add_note('Valida\u00e7\u00e3o',
         'Speech: 4,400h = 100 chamadas/min \u00d7 10s \u00d7 12h/dia \u00d7 22 dias \u2705 | '
         'AKS Reserved 3 anos: economia de ~53% vs pay-as-you-go \u2705 | '
         'Redis Standard com SLA \u2705',
         label_color=GREEN_COLOR)

add_note('Destaque',
         'AKS com Reserved Instances 3 anos reduziu de $528/m\u00eas para $250/m\u00eas (economia de $279/m\u00eas). '
         'Speech Services j\u00e1 domina 73% do custo total neste cen\u00e1rio.',
         label_color=ACCENT_COLOR)
spacer()

# --- 2.3 ---
doc.add_heading('2.3 Cen\u00e1rio 1.000 chamadas/min', level=2)

p = doc.add_paragraph()
run = p.add_run('Total mensal: $48,283.49 USD (R$ 280,044)')
run.bold = True

spacer()
add_styled_table(doc,
    ['Servi\u00e7o', 'SKU / Tier', 'USD/m\u00eas', '% Total'],
    [
        ['VPN Gateway', 'VpnGw3AZ (Zone-Redundant), 3 TB transfer', '$1,153.65', '2%'],
        ['Speech Services STT', 'Standard S0 \u2014 44,000h \u00e1udio + 1 Custom endpoint', '$44,023.65', '91%'],
        ['AKS', '2 clusters, 4x D8as v4 (8 vCPU, 32GB), Reserved 3 anos', '$961.25', '2%'],
        ['API Management', 'Standard \u2014 2 unidades', '$1,373.42', '3%'],
        ['Redis Cache', 'Standard C2 x 2 inst\u00e2ncias', '$327.04', '1%'],
        ['Azure Monitor + App Insights', '~4 GB/dia logs', '$431.30', '1%'],
        ['Key Vault', 'Standard', '$13.18', '<1%'],
        ['TOTAL', '', '$48,283.49', '100%'],
    ],
    col_widths=[4.5, 6, 2.5, 1.5],
    highlight_last_row=True
)

spacer()
add_note('Valida\u00e7\u00e3o',
         'Speech: 44,000h = 1.000 \u00d7 10s \u00d7 15.840 min/m\u00eas \u2705 | '
         'AKS Reserved: $961 vs $2,076 pay-as-you-go (economia 54%) \u2705',
         label_color=GREEN_COLOR)

add_note('Alerta de escala',
         'Com 44,000h de Speech/m\u00eas, verificar limite de sess\u00f5es concorrentes do S0 em Brazil South. '
         '~167 sess\u00f5es simult\u00e2neas \u2014 pode ser necess\u00e1rio aumento de quota junto \u00e0 Microsoft.',
         label_color=RED_COLOR)
spacer()

# --- 2.4 ---
doc.add_heading('2.4 Cen\u00e1rio 5.000 chamadas/min', level=2)

p = doc.add_paragraph()
run = p.add_run('Total mensal: $227,663.62 USD (R$ 1,320,449)')
run.bold = True

spacer()
add_styled_table(doc,
    ['Servi\u00e7o', 'SKU / Tier', 'USD/m\u00eas', '% Total'],
    [
        ['VPN Gateway', 'VpnGw4AZ (Zone-Redundant), 3 TB transfer', '$1,774.15', '1%'],
        ['Speech Services STT', 'Standard S0 \u2014 220,000h \u00e1udio + 1 Custom endpoint', '$220,023.65', '97%'],
        ['AKS', '3 clusters, 8x D8as v4 (8 vCPU, 32GB), Reserved 3 anos', '$1,849.50', '1%'],
        ['API Management', 'Standard \u2014 4 unidades', '$2,746.84', '1%'],
        ['Redis Cache', 'Standard C3 x 2 inst\u00e2ncias', '$657.00', '<1%'],
        ['Azure Monitor + App Insights', '~6 GB/dia logs', '$599.30', '<1%'],
        ['Key Vault', 'Standard', '$13.18', '<1%'],
        ['TOTAL', '', '$227,663.62', '100%'],
    ],
    col_widths=[4.5, 6, 2.5, 1.5],
    highlight_last_row=True
)

spacer()
add_note('Valida\u00e7\u00e3o',
         'Speech: 220,000h = 5.000 \u00d7 10s \u00d7 15.840 min/m\u00eas \u2705 | '
         'AKS Reserved: $1,850 vs $4,080 pay-as-you-go (economia 55%) \u2705',
         label_color=GREEN_COLOR)

add_note('Alerta cr\u00edtico',
         'Speech Services representa 97% do custo ($220K/m\u00eas). '
         'Este cen\u00e1rio EXIGE negocia\u00e7\u00e3o Enterprise Agreement ou alternativa self-hosted (Whisper). '
         'Tamb\u00e9m verificar: ~833 sess\u00f5es concorrentes \u2014 m\u00faltiplos endpoints de Speech necess\u00e1rios.',
         label_color=RED_COLOR)
spacer()

doc.add_page_break()

# ============================================================
# 3. VALIDACAO SPEECH HOURS
# ============================================================
doc.add_heading('3. Valida\u00e7\u00e3o das Horas de Speech Services', level=1)

doc.add_heading('3.1 Premissas de C\u00e1lculo', level=2)

add_styled_table(doc,
    ['Par\u00e2metro', 'Valor'],
    [
        ['Janela de opera\u00e7\u00e3o', '12 horas/dia (07h-19h)'],
        ['Dias \u00fateis/m\u00eas', '22 dias'],
        ['Dura\u00e7\u00e3o m\u00e9dia de \u00e1udio por chamada', '10 segundos'],
        ['Minutos de opera\u00e7\u00e3o/m\u00eas', '15,840 min (12h \u00d7 60min \u00d7 22d)'],
        ['Pre\u00e7o STT Real-Time Standard', '~$1.00 por hora de \u00e1udio'],
        ['Custom Speech Endpoint', '~$23.65/m\u00eas por endpoint'],
    ],
    col_widths=[7, 9]
)

spacer()

doc.add_heading('3.2 Confer\u00eancia das Horas', level=2)

add_styled_table(doc,
    ['Cen\u00e1rio', 'Chamadas/m\u00eas', 'Horas esperadas\n(@10s)', 'Horas na\nCalculadora', 'Status'],
    [
        ['10/min', '158,400', '440h', '440h', '\u2705 Correto'],
        ['100/min', '1,584,000', '4,400h', '4,400h', '\u2705 Correto'],
        ['1.000/min', '15,840,000', '44,000h', '44,000h', '\u2705 Correto'],
        ['5.000/min', '79,200,000', '220,000h', '220,000h', '\u2705 Correto'],
    ],
    col_widths=[2.5, 3, 3, 3, 3],
    green_rows=[0, 1, 2, 3]
)

spacer()
add_alert('\u2705 Todas as horas de Speech est\u00e3o corretas e aderentes \u00e0s premissas.', color=GREEN_COLOR)
spacer()

doc.add_page_break()

# ============================================================
# 4. CONFORMIDADE COM ARQUITETURA
# ============================================================
doc.add_heading('4. Conformidade com a Arquitetura do Projeto', level=1)

p = doc.add_paragraph(
    'Valida\u00e7\u00e3o dos SKUs selecionados na calculadora contra os requisitos '
    'definidos no documento de solu\u00e7\u00e3o (C4 Model & Desenho de Solu\u00e7\u00e3o):')
spacer()

doc.add_heading('4.1 Matriz de Conformidade', level=2)

add_styled_table(doc,
    ['Requisito', '10/min', '100/min', '1.000/min', '5.000/min'],
    [
        ['VNet Integration', '\u2705 P1V3', '\u2705 AKS', '\u2705 AKS', '\u2705 AKS'],
        ['Private Endpoints', '\u2705 Suportado', '\u2705 Suportado', '\u2705 Suportado', '\u2705 Suportado'],
        ['APIM c/ VNet', '\u2705 Standard', '\u2705 Standard', '\u2705 Standard', '\u2705 Standard'],
        ['Redis c/ SLA', '\u2705 Standard', '\u2705 Standard', '\u2705 Standard', '\u2705 Standard'],
        ['Zone-Redundant VPN', '\u2705 Gw1AZ', '\u2705 Gw2AZ', '\u2705 Gw3AZ', '\u2705 Gw4AZ'],
        ['Speech Hours Corretas', '\u2705 440h', '\u2705 4,400h', '\u2705 44,000h', '\u2705 220,000h'],
        ['Reserved Instances', '\u2014 (App Svc)', '\u2705 AKS 3yr', '\u2705 AKS 3yr', '\u2705 AKS 3yr'],
    ],
    col_widths=[4, 3, 3, 3, 3],
    green_rows=[0, 1, 2, 3, 4, 5, 6]
)

spacer()

doc.add_heading('4.2 Itens Pendentes (n\u00e3o presentes na calculadora)', level=2)

p = doc.add_paragraph('Os seguintes custos n\u00e3o est\u00e3o nos exports mas devem ser considerados:')
spacer()

add_styled_table(doc,
    ['Item', 'Impacto estimado', 'Cen\u00e1rios afetados', 'Prioridade'],
    [
        ['Private Endpoints (4 PEs)', '$40\u2013100/m\u00eas', 'Todos', 'Incluir'],
        ['Bandwidth outbound', '$0 (10/min) a $200+ (5.000/min)', '1.000/min, 5.000/min', 'Incluir'],
        ['DDoS Network Protection', '~$2,944/m\u00eas', 'Somente 5.000/min', 'Avaliar'],
        ['Custom Speech Training', 'Pontual (n\u00e3o recorrente)', 'Todos', 'Informativo'],
        ['Servi\u00e7os profissionais Foursys', 'Fora do escopo Azure', 'Todos', 'Separado'],
    ],
    col_widths=[5, 3.5, 3.5, 3]
)

spacer()
add_note('Impacto total estimado',
         'Private Endpoints + Bandwidth adicionam ~$50\u2013300/m\u00eas dependendo do cen\u00e1rio. '
         'Impacto percentual < 2% nos cen\u00e1rios acima de 100/min.',
         label_color=ACCENT_COLOR)
spacer()

doc.add_page_break()

# ============================================================
# 5. COMPARATIVO CONSOLIDADO
# ============================================================
doc.add_heading('5. Tabela Comparativa Consolidada', level=1)

doc.add_heading('5.1 Valores da Calculadora Revisada (USD)', level=2)

add_styled_table(doc,
    ['Componente', '10/min', '100/min', '1.000/min', '5.000/min'],
    [
        ['VPN Gateway', '$156.95', '$458.45', '$1,153.65', '$1,774.15'],
        ['Speech Services STT', '$463.65', '$4,423.65', '$44,023.65', '$220,023.65'],
        ['Compute (App Svc / AKS)', '$162.06', '$249.56', '$961.25', '$1,849.50'],
        ['API Management', '$686.71', '$686.71', '$1,373.42', '$2,746.84'],
        ['Redis Cache', '$100.74', '$100.74', '$327.04', '$657.00'],
        ['Monitor + Logs', '$128.90', '$128.90', '$431.30', '$599.30'],
        ['Key Vault', '$13.18', '$13.18', '$13.18', '$13.18'],
        ['TOTAL (USD)', '$1,712.19', '$6,061.19', '$48,283.49', '$227,663.62'],
        ['TOTAL (BRL @5.80)', 'R$ 9,931', 'R$ 35,155', 'R$ 280,044', 'R$ 1,320,449'],
    ],
    col_widths=[5, 2.5, 2.5, 3, 3],
    highlight_last_row=True
)

spacer()

doc.add_heading('5.2 Evolu\u00e7\u00e3o: Calculadora v1 vs v2 vs Estimativa Original', level=2)

add_styled_table(doc,
    ['Cen\u00e1rio', 'Estimativa Original\n(Documento)', 'Calculadora v1\n(com erros)', 'Calculadora v2\n(corrigida)', 'Desvio\nv2 vs Original'],
    [
        ['10/min', 'R$ 8,370', 'R$ 5,857', 'R$ 9,931', '+19%'],
        ['100/min', 'R$ 33,150', 'R$ 34,100', 'R$ 35,155', '+6%'],
        ['1.000/min', 'R$ 267,050', 'R$ 77,712', 'R$ 280,044', '+5%'],
        ['5.000/min', '\u2014', 'R$ 173,385', 'R$ 1,320,449', 'Novo'],
    ],
    col_widths=[2.5, 3, 3, 3, 3]
)

spacer()

p = doc.add_paragraph('An\u00e1lise dos desvios:')
desvios = [
    ('10/min (+19%)', 'A estimativa original usava APIM Standard v2 (mais barato que Standard legacy). '
     'O Redis Standard tamb\u00e9m \u00e9 mais caro que o Basic estimado anteriormente. '
     'Desvio aceit\u00e1vel \u2014 calculadora reflete melhor o custo real.'),
    ('100/min (+6%)', 'Praticamente alinhado. A diferen\u00e7a vem do ajuste de 400h extras de Speech '
     'e Redis Standard vs Basic. Compensa\u00e7\u00e3o parcial pelo AKS Reserved.'),
    ('1.000/min (+5%)', 'Excelente ader\u00eancia. A corre\u00e7\u00e3o maci\u00e7a do Speech (8.000h \u2192 44.000h) '
     'confirmou a estimativa original. AKS Reserved compensou parcialmente.'),
]

for title, desc in desvios:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = ACCENT_COLOR
    run = p.add_run(desc)
    run.font.size = Pt(9)

spacer()
doc.add_page_break()

# ============================================================
# 6. PROPORCAO DE CUSTO
# ============================================================
doc.add_heading('6. Propor\u00e7\u00e3o de Custo por Componente', level=1)

add_styled_table(doc,
    ['Cen\u00e1rio', 'Total/m\u00eas', '% Speech', '% APIM', '% Compute', '% VPN', '% Outros'],
    [
        ['10/min', '$1,712', '27%', '40%', '9%', '9%', '14%'],
        ['100/min', '$6,061', '73%', '11%', '4%', '8%', '4%'],
        ['1.000/min', '$48,283', '91%', '3%', '2%', '2%', '2%'],
        ['5.000/min', '$227,664', '97%', '1%', '1%', '1%', '<1%'],
    ],
    col_widths=[2.5, 2.5, 2, 2, 2.5, 2, 2.5]
)

spacer()

add_note('Padr\u00e3o identificado',
         'No MVP (10/min), o APIM Standard \u00e9 o maior custo (40%). '
         'A partir de 100/min, Speech Services domina e cresce de 73% para 97%. '
         'Isso confirma que a estrat\u00e9gia de otimiza\u00e7\u00e3o deve mudar conforme a escala.',
         label_color=ACCENT_COLOR)

spacer()

# ============================================================
# 7. TCO 12 MESES
# ============================================================
doc.add_heading('7. TCO \u2014 Custo Total de Propriedade (12 meses)', level=1)

p = doc.add_paragraph('Proje\u00e7\u00e3o anual considerando apenas custo de infraestrutura Azure:')
spacer()

add_styled_table(doc,
    ['M\u00e9trica', '10/min', '100/min', '1.000/min', '5.000/min'],
    [
        ['Custo mensal (USD)', '$1,712', '$6,061', '$48,283', '$227,664'],
        ['Custo mensal (BRL)', 'R$ 9,931', 'R$ 35,155', 'R$ 280,044', 'R$ 1,320,449'],
        ['TCO 12 meses (USD)', '$20,546', '$72,734', '$579,402', '$2,731,963'],
        ['TCO 12 meses (BRL)', 'R$ 119,169', 'R$ 421,859', 'R$ 3,360,530', 'R$ 15,845,387'],
        ['Custo por chamada', '$0.0108', '$0.0038', '$0.0030', '$0.0029'],
        ['Custo por chamada (BRL)', 'R$ 0,063', 'R$ 0,022', 'R$ 0,018', 'R$ 0,017'],
    ],
    col_widths=[4, 3, 3, 3, 3],
    highlight_last_row=True
)

spacer()

add_note('Economia de escala',
         'O custo por chamada cai de R$ 0,063 (10/min) para R$ 0,017 (5.000/min) \u2014 '
         'redu\u00e7\u00e3o de 73%. Por\u00e9m, a economia \u00e9 limitada pelo Speech Services (pay-per-use linear). '
         'Os custos fixos (VPN, APIM, Redis) se diluem, mas representam apenas 3-27% do total.',
         label_color=ACCENT_COLOR)

spacer()
doc.add_page_break()

# ============================================================
# 8. ECONOMIA COM RESERVED INSTANCES
# ============================================================
doc.add_heading('8. Impacto das Reserved Instances (AKS)', level=1)

p = doc.add_paragraph(
    'As calculadoras revisadas utilizam Reserved Instances de 3 anos para os nodes AKS, '
    'gerando economia significativa no compute:')
spacer()

add_styled_table(doc,
    ['Cen\u00e1rio', 'AKS Config', 'Pay-as-you-go\n(estimado)', 'Reserved 3yr\n(calculadora)', 'Economia\nmensal', 'Economia\n3 anos'],
    [
        ['100/min', '2x D4as v4, 1 cluster', '$528', '$250', '$279 (53%)', '$10,028'],
        ['1.000/min', '4x D8as v4, 2 clusters', '$2,076', '$961', '$1,115 (54%)', '$40,151'],
        ['5.000/min', '8x D8as v4, 3 clusters', '$4,080', '$1,850', '$2,230 (55%)', '$80,290'],
    ],
    col_widths=[2.5, 3, 2.5, 2.5, 2.5, 2.5]
)

spacer()
add_note('Total economizado',
         'Reserved Instances de 3 anos geram ~$3,624/m\u00eas de economia somando os 3 cen\u00e1rios com AKS. '
         'Sobre 3 anos, isso representa ~$130,469 de economia agregada.',
         label_color=GREEN_COLOR)

add_note('Trade-off',
         'Reserved 3 anos exige compromisso de longo prazo. Se o volume de chamadas n\u00e3o se confirmar, '
         'o custo fica \"preso\". Recomenda\u00e7\u00e3o: iniciar com pay-as-you-go no MVP e migrar para Reserved '
         'ap\u00f3s validar o volume real com a Atento.',
         label_color=ACCENT_COLOR)

spacer()
doc.add_page_break()

# ============================================================
# 9. VEREDITO FINAL
# ============================================================
doc.add_heading('9. Veredito: As Calculadoras Fazem Sentido?', level=1)

p = doc.add_paragraph()
run = p.add_run('SIM \u2014 todas as 4 calculadoras est\u00e3o agora consistentes com o projeto Atento.')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = GREEN_COLOR

spacer()

verdicts = [
    ('10/min (MVP)', '\u2705 V\u00c1LIDO',
     'SKUs adequados (P1V3, Standard APIM, Standard Redis). '
     'Speech hours corretas. VNet Integration suportada. '
     'Custo mensal de $1,712 (R$ 9,931) \u00e9 razo\u00e1vel para um MVP.',
     GREEN_COLOR),
    ('100/min', '\u2705 V\u00c1LIDO',
     'Speech hours corrigidas para 4,400h. AKS com Reserved 3yr. '
     'Redis e APIM adequados. Custo de $6,061 (R$ 35,155) alinhado \u00e0 estimativa original.',
     GREEN_COLOR),
    ('1.000/min', '\u2705 V\u00c1LIDO',
     'Corre\u00e7\u00e3o cr\u00edtica aplicada: 44,000h de Speech (era 8,000h). '
     'Total de $48,283 (R$ 280,044) confirma a estimativa original. '
     'Ponto de aten\u00e7\u00e3o: quota de sess\u00f5es concorrentes do Speech.',
     GREEN_COLOR),
    ('5.000/min', '\u2705 V\u00c1LIDO (com ressalvas)',
     'Corre\u00e7\u00e3o cr\u00edtica aplicada: 220,000h de Speech (era 20,000h). '
     'Total de $227,664 (R$ 1,320,449). Cen\u00e1rio vi\u00e1vel tecnicamente, '
     'mas EXIGE Enterprise Agreement ou alternativa self-hosted para viabilidade financeira.',
     ACCENT_COLOR),
]

for title, status, desc, color in verdicts:
    doc.add_heading(f'{title} \u2014 {status}', level=2)
    p = doc.add_paragraph(desc)
    spacer(4)

spacer()

# ============================================================
# 10. RECOMENDACOES
# ============================================================
doc.add_heading('10. Recomenda\u00e7\u00f5es', level=1)

doc.add_heading('10.1 A\u00e7\u00f5es Imediatas', level=2)

actions = [
    ('Confirmar volume real com a Atento',
     'Solicitar dados hist\u00f3ricos de chamadas/minuto durante pico e m\u00e9dia. '
     'Essa \u00e9 a vari\u00e1vel de maior impacto no custo \u2014 determina qual cen\u00e1rio usar.'),
    ('Validar quota de Speech Services em Brazil South',
     'Para cen\u00e1rios 1.000/min e 5.000/min, as sess\u00f5es concorrentes (167 e 833) '
     'podem exceder o limite padr\u00e3o. Abrir request de aumento de quota \u00e0 Microsoft.'),
    ('Incluir Private Endpoints e Bandwidth',
     'Adicionar ~$50-300/m\u00eas (conforme cen\u00e1rio) para Private Endpoints e data transfer '
     'outbound na estimativa final.'),
]

for title, desc in actions:
    p = doc.add_paragraph()
    run = p.add_run(f'\u25b6 {title}: ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT_COLOR
    run = p.add_run(desc)
    run.font.size = Pt(10)

spacer()

doc.add_heading('10.2 Otimiza\u00e7\u00f5es de Custo', level=2)

opts = [
    ('Enterprise Agreement (Speech)',
     'Para cen\u00e1rios 100+/min, negociar desconto de 30-50% no Speech Services via EA. '
     'No cen\u00e1rio de 5.000/min, isso pode economizar $66K-110K/m\u00eas.'),
    ('Alternativa Self-Hosted (1.000+/min)',
     'Whisper/Faster-Whisper no AKS com GPUs (A10/T4). Troca custo pay-per-use '
     'por custo de compute fixo \u2014 potencialmente 50-70% mais barato para alto volume.'),
    ('APIM no MVP',
     'Avaliar APIM Basic v2 (se suportar VNet) ou autentica\u00e7\u00e3o direta no App Service '
     'para reduzir o custo de 40% para <10% no cen\u00e1rio 10/min.'),
    ('Reserved Instances \u2014 Timing',
     'Iniciar com pay-as-you-go. Ap\u00f3s 3-6 meses validando volume real, '
     'converter para Reserved 1yr ou 3yr nos nodes AKS.'),
]

for title, desc in opts:
    p = doc.add_paragraph()
    run = p.add_run(f'\u25b6 {title}: ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT_COLOR
    run = p.add_run(desc)
    run.font.size = Pt(10)

spacer()

doc.add_heading('10.3 Cen\u00e1rio Recomendado para Proposta Comercial', level=2)

add_styled_table(doc,
    ['Fase', 'Cen\u00e1rio', 'Custo Azure', 'Per\u00edodo', 'Objetivo'],
    [
        ['PoC', '10/min (MVP)', '$1,712/m\u00eas', '2-3 meses', 'Validar integra\u00e7\u00e3o Avaya + ASR'],
        ['Produ\u00e7\u00e3o', '100/min', '$6,061/m\u00eas', '6-12 meses', 'Atender volume inicial real'],
        ['Escala', '1.000/min', '$48,283/m\u00eas', 'Sob demanda', 'Expandir conforme crescimento'],
    ],
    col_widths=[2, 3, 3, 3, 5]
)

spacer()
add_note('Estrat\u00e9gia',
         'Come\u00e7ar pelo MVP para validar a solu\u00e7\u00e3o com investimento m\u00ednimo ($1,712/m\u00eas). '
         'Escalar para 100/min quando o volume real justificar. '
         'O cen\u00e1rio de 5.000/min serve como refer\u00eancia de teto, n\u00e3o como meta inicial.',
         label_color=ACCENT_COLOR)

spacer()
doc.add_page_break()

# ============================================================
# 11. APENDICE
# ============================================================
doc.add_heading('11. Ap\u00eandice', level=1)

doc.add_heading('11.1 Arquivos Analisados', level=2)

add_styled_table(doc,
    ['Arquivo', 'Timestamp', 'Tamanho'],
    [
        ['[Atento] Cen\u00e1rio- MVP (10 Liga\u00e7\u00f5es min).xlsx', '14/04/2026 12:37 UTC', '~80 KB'],
        ['[Atento] Cen\u00e1rio- MVP (100 Liga\u00e7\u00f5es min).xlsx', '14/04/2026 12:39 UTC', '~55 KB'],
        ['[Atento] Cen\u00e1rio- MVP (1000 Liga\u00e7\u00f5es min).xlsx', '14/04/2026 12:54 UTC', '~55 KB'],
        ['[Atento] Cen\u00e1rio- MVP (5000 Liga\u00e7\u00f5es min).xlsx', '14/04/2026 12:36 UTC', '~80 KB'],
    ],
    col_widths=[8, 4, 3]
)

spacer()

doc.add_heading('11.2 Descri\u00e7\u00f5es dos Servi\u00e7os na Calculadora', level=2)

services = [
    ('VPN Gateway (10/min)', 'VpnGw1 tier, 730 gateway hours, 1 S2S tunnel, 100 GB transfer'),
    ('VPN Gateway (100/min)', 'VpnGw2 tier, 730 gateway hours, 1 S2S tunnel, 1,200 GB transfer'),
    ('VPN Gateway (1.000/min)', 'VpnGw3 tier, 730 gateway hours, 1 S2S tunnel, 3 TB transfer'),
    ('VPN Gateway (5.000/min)', 'VpnGw4 tier, 730 gateway hours, 1 S2S tunnel, 3 TB transfer'),
    ('Speech (10/min)', '440 Standard real-time hours, 1 Custom endpoint x 440 Audio hours'),
    ('Speech (100/min)', '4,400 Standard real-time hours, 1 Custom endpoint x 440 Audio hours'),
    ('Speech (1.000/min)', '44,000 Standard real-time hours, 1 Custom endpoint x 440 Audio hours'),
    ('Speech (5.000/min)', '220,000 Standard real-time hours, 1 Custom endpoint x 440 Audio hours'),
    ('App Service (10/min)', 'Premium V3 P1V3 (2 vCPU, 8 GB RAM, 250 GB) x 730h, Linux'),
    ('AKS (100/min)', 'Standard, 1 cluster, 2x D4as v4 (4 vCPUs, 16 GB), 3yr reserved, 3 discos S15'),
    ('AKS (1.000/min)', 'Standard, 2 clusters, 4x D8as v4 (8 vCPUs, 32 GB), 3yr reserved, 4 discos E15'),
    ('AKS (5.000/min)', 'Standard, 3 clusters, 8x D8as v4 (8 vCPUs, 32 GB), 3yr reserved, 8 discos E15'),
]

for title, desc in services:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(desc)
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY_COLOR

spacer()

doc.add_heading('11.3 Pre\u00e7os de Refer\u00eancia (Brazil South, Abril 2026)', level=2)

add_styled_table(doc,
    ['Servi\u00e7o', 'M\u00e9trica', 'Pre\u00e7o USD'],
    [
        ['Speech STT Real-Time', 'Por hora de \u00e1udio', '~$1.00/h'],
        ['Custom Speech Endpoint', 'Por endpoint/m\u00eas', '~$23.65'],
        ['VPN Gateway VpnGw1AZ', 'Por hora (730h/m\u00eas)', '~$0.215/h'],
        ['VPN Gateway VpnGw2AZ', 'Por hora (730h/m\u00eas)', '~$0.528/h'],
        ['App Service P1V3 Linux', 'Por inst\u00e2ncia/m\u00eas', '~$162'],
        ['AKS D4as v4 (Reserved 3yr)', 'Por node/m\u00eas', '~$105'],
        ['AKS D8as v4 (Reserved 3yr)', 'Por node/m\u00eas', '~$210'],
        ['Redis Standard C1', 'Por inst\u00e2ncia/m\u00eas', '~$100.74'],
        ['APIM Standard', 'Por unidade/m\u00eas', '~$686.71'],
    ],
    col_widths=[5, 5, 4]
)

spacer()

# ============================================================
# RODAPE
# ============================================================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\u2500' * 60)
run.font.color.rgb = RGBColor(180, 180, 180)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documento produzido pela Squad MEQ \u2014 Foursys')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(120, 120, 120)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Revis\u00e3o: v2.0 \u2014 14 Abril 2026 | Classifica\u00e7\u00e3o: Confidencial')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(120, 120, 120)

# ============================================================
# SALVAR
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           'Analise Calculadoras Azure - Atento v2.docx')
doc.save(output_path)
print(f'Documento salvo em: {output_path}')
